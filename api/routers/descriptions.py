"""
Descriptions API router.

Provides streaming SSE endpoint for AI scene description generation
and a save endpoint to persist generated descriptions.

Routes:
  POST /api/descriptions/generate — streaming SSE description generation
  POST /api/descriptions/save     — save description to SQLite scenes table
"""

from __future__ import annotations

import json
import logging
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import anthropic
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from api.auth import CurrentUser
from api.config import get_settings
from api.database import get_db
from api.prompts import DESC_SYSTEMS, DESC_COMPILATION_SYSTEMS, STUDIO_KEY_MAP, get_prompt

_log = logging.getLogger(__name__)

router = APIRouter(prefix="/api/descriptions", tags=["descriptions"])


# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------

class DescGenRequest(BaseModel):
    studio: str                     # "FuckPassVR", "VRHush", "VRAllure", "NaughtyJOI"
    scene_id: str | None = None
    is_compilation: bool = False
    performers: str = ""
    sex_positions: str = ""
    categories: str = ""
    target_keywords: str = ""
    wardrobe: str = ""
    model_properties: str = ""      # freeform model notes
    plot: str = ""                  # existing plot for context


class DescSaveBody(BaseModel):
    scene_id: str
    description: str
    meta_title: str | None = None
    meta_description: str | None = None


class SeoGenRequest(BaseModel):
    description: str
    studio: str
    performers: str = ""
    title: str = ""


class DocxRequest(BaseModel):
    description: str
    title: str = ""
    meta_title: str | None = None
    meta_description: str | None = None


class ParagraphRegenRequest(BaseModel):
    studio: str                     # "FuckPassVR", "VRHush", "VRAllure", "NaughtyJOI"
    paragraph: str                  # Current paragraph body
    paragraph_index: int            # 0-based index, for prompt context
    performer: str = ""
    title: str = ""                 # Scene title
    plot: str = ""                  # Scene plot (truncated in prompt)
    feedback: str = ""              # Optional director-style nudge


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.post("/generate")
async def generate_description(body: DescGenRequest, user: CurrentUser):
    """
    Generate a scene description via Claude (streaming SSE).

    Uses per-studio system prompts. Compilation scenes use the
    DESC_COMPILATION_SYSTEMS prompts; regular scenes use DESC_SYSTEMS.
    """
    studio_key = STUDIO_KEY_MAP.get(body.studio)
    if not studio_key:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown studio: {body.studio}. Must be one of: {', '.join(STUDIO_KEY_MAP.keys())}",
        )

    # Read via get_prompt so admin overrides apply without a restart.
    # Bundled defaults are still the fallback when no override exists.
    prompt_key = (
        f"desc_comp.{studio_key}" if body.is_compilation else f"desc.{studio_key}"
    )
    fallback_dict = DESC_COMPILATION_SYSTEMS if body.is_compilation else DESC_SYSTEMS
    system_prompt = get_prompt(prompt_key, fallback=fallback_dict.get(studio_key, ""))
    if not system_prompt:
        raise HTTPException(
            status_code=400,
            detail=f"No description prompt found for studio key: {studio_key}",
        )

    user_prompt = _build_desc_user_prompt(body)
    settings = get_settings()

    def event_stream():
        try:
            from api.ollama_client import ollama_stream
            for delta in ollama_stream(
                "description",
                user_prompt,
                system=system_prompt,
                max_tokens=2048,
                temperature=0.7,
            ):
                yield f"data: {json.dumps({'type': 'text', 'text': delta})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as exc:
            _log.error("Description generation failed: %s", exc, exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'error': str(exc)})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/regenerate-paragraph")
async def regenerate_paragraph(body: ParagraphRegenRequest, user: CurrentUser):
    """
    Rewrite a single paragraph of a scene description. Uses the same per-studio
    system prompt as /generate so tone/voice stays consistent with the rest.

    The user types an optional feedback nudge ("make it steamier", "more POV",
    etc.); we feed it alongside the current paragraph and scene context and
    return the new text as plain JSON — no streaming, since a single paragraph
    regen is fast enough that a loading spinner is fine.
    """
    studio_key = STUDIO_KEY_MAP.get(body.studio)
    if not studio_key:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown studio: {body.studio}",
        )
    system_prompt = get_prompt(f"desc.{studio_key}", fallback=DESC_SYSTEMS.get(studio_key, ""))
    if not system_prompt:
        raise HTTPException(
            status_code=400,
            detail=f"No description prompt for studio: {studio_key}",
        )

    settings = get_settings()
    if not settings.anthropic_api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")

    nudge = f"\n\nEditor feedback: {body.feedback.strip()}" if body.feedback.strip() else ""
    user_prompt = (
        f"Rewrite ONLY paragraph {body.paragraph_index + 1} of the scene description. "
        f"Keep the same voice, tense, and POV as the surrounding description.\n\n"
        f"Scene context:\n"
        f"- Performer: {body.performer or '(unspecified)'}\n"
        f"- Title: {body.title or '(untitled)'}\n"
        f"- Plot: {(body.plot or '')[:400]}\n\n"
        f"Current paragraph:\n{body.paragraph}\n"
        f"{nudge}\n\n"
        f"Output ONLY the new paragraph body — no title, no meta, no headings. "
        f"Keep it roughly the same length."
    )

    try:
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=600,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        new_text = msg.content[0].text.strip() if msg.content else ""
    except Exception as exc:
        _log.error("Paragraph regeneration failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=502, detail=f"Claude error: {exc}") from exc

    return {"paragraph": new_text}


@router.post("/seo")
async def generate_seo_tags(body: SeoGenRequest, user: CurrentUser):
    """
    Generate SEO meta title and meta description for a scene.

    Returns plain JSON (not streaming) — fast single-call Claude response.
    """
    settings = get_settings()

    system_prompt = (
        "You are an SEO specialist for adult VR content. Given a scene description, "
        "generate two fields and return ONLY valid JSON — no markdown, no commentary:\n"
        "- meta_title: Keyword-rich, under 60 characters. Include studio + performer + key act. "
        "No ALL CAPS. Examples by studio:\n"
        '  VRHush: "Kenzie Anne Deepthroat & Creampie VR Porn - VRHush"\n'
        '  FuckPassVR: "Lana Roy POV Sex in Prague VR Porn - FuckPassVR"\n'
        '  VRAllure: "Eliza Ibarra Solo Masturbation VR - VRAllure"\n'
        '  NaughtyJOI: "Lulu Chu JOI Countdown VR Experience - NJOI"\n'
        "- meta_description: Under 155 characters. Performer name + key acts + light VR CTA. "
        "Example: \"Watch Kenzie Anne take you deep in this 8K VR porn scene — blowjob, cowgirl, and creampie finish. Exclusive on VRHush.\"\n"
        "Return exactly: {\"meta_title\": \"...\", \"meta_description\": \"...\"}"
    )

    user_parts = [f"Studio: {body.studio}"]
    if body.performers:
        user_parts.append(f"Performers: {body.performers}")
    if body.title:
        user_parts.append(f"Scene title: {body.title}")
    user_parts.append(f"\nDescription:\n{body.description[:1000]}")
    user_prompt = "\n".join(user_parts)

    try:
        from api.ollama_client import ollama_generate
        raw = ollama_generate("seo", user_prompt, system=system_prompt, max_tokens=300, temperature=0.4)
        import json as _json
        # Strip any markdown fences if present
        if raw.startswith("```"):
            raw = raw.split("```")[1].lstrip("json").strip()
        # Try to find the JSON object if there's extra text around it
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            raw = raw[start:end + 1]
        data = _json.loads(raw)
        return {
            "meta_title": str(data.get("meta_title", ""))[:60],
            "meta_description": str(data.get("meta_description", ""))[:155],
        }
    except Exception as exc:
        _log.error("SEO generation failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=503, detail=str(exc))


@router.post("/docx")
async def download_description_docx(body: DocxRequest, user: CurrentUser):
    """
    Generate and return a DOCX file containing the scene description.

    Includes meta title and meta description if provided.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor
    from fastapi.responses import Response

    doc = Document()

    # Title (if provided)
    if body.title:
        h = doc.add_heading(body.title, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Main description — split on double newlines into paragraphs
    paragraphs = [p.strip() for p in body.description.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.runs[0].font.size = Pt(11)

    # SEO section
    if body.meta_title or body.meta_description:
        doc.add_paragraph("")  # spacer
        doc.add_heading("SEO Metadata", level=2)
        if body.meta_title:
            p = doc.add_paragraph()
            p.add_run("Meta Title: ").bold = True
            p.add_run(body.meta_title)
        if body.meta_description:
            p = doc.add_paragraph()
            p.add_run("Meta Description: ").bold = True
            p.add_run(body.meta_description)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = "description.docx"
    if body.title:
        safe = body.title[:40].replace(" ", "_").replace("/", "-")
        filename = f"{safe}.docx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/save")
async def save_description(body: DescSaveBody, user: CurrentUser):
    """
    Save a description to the SQLite scenes table.

    Note: Grail write-through (for the actual site) goes through the
    approvals flow, not directly here.
    """
    with get_db() as conn:
        row = conn.execute(
            "SELECT id FROM scenes WHERE id = ?", (body.scene_id,)
        ).fetchone()

        if not row:
            raise HTTPException(status_code=404, detail="Scene not found")

        conn.execute(
            "UPDATE scenes SET has_description=1 WHERE id=?",
            (body.scene_id,),
        )

    return {"scene_id": body.scene_id, "status": "saved"}


# ---------------------------------------------------------------------------
# User prompt builder
# ---------------------------------------------------------------------------

def _build_desc_user_prompt(body: DescGenRequest) -> str:
    """Build the user-turn prompt for description generation."""
    lines = [f"Write a scene description for {body.studio} with the following details:"]

    if body.performers:
        lines.append(f"Performers: {body.performers}")
    if body.sex_positions:
        lines.append(f"Sex Positions: {body.sex_positions}")
    if body.categories:
        lines.append(f"Categories: {body.categories}")
    if body.target_keywords:
        lines.append(f"Target Keywords: {body.target_keywords}")
    if body.wardrobe:
        lines.append(f"Wardrobe: {body.wardrobe}")
    if body.model_properties:
        lines.append(f"Model Notes: {body.model_properties}")
    if body.plot:
        lines.append(f"Plot Summary: {body.plot}")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Save description as DOCX to MEGA Description/ folder
# ---------------------------------------------------------------------------

_STUDIO_TO_MEGA: dict[str, str] = {
    "FuckPassVR": "FPVR",
    "VRHush":     "VRH",
    "VRAllure":   "VRA",
    "NaughtyJOI": "NNJOI",
}


class DescSaveMegaBody(BaseModel):
    scene_id: str
    description: str
    title: str = ""
    meta_title: str | None = None
    meta_description: str | None = None


@router.post("/save-mega")
async def save_description_to_mega(body: DescSaveMegaBody, user: CurrentUser):
    """
    Build a DOCX from the description and upload it to the scene's
    MEGA Description/ subfolder via rclone. Also marks has_description=1.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor

    with get_db() as conn:
        row = conn.execute(
            "SELECT id, studio, grail_tab FROM scenes WHERE id = ?", (body.scene_id,)
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Scene not found")
        scene = dict(row)
        conn.execute("UPDATE scenes SET has_description=1 WHERE id=?", (body.scene_id,))

    mega_studio = _STUDIO_TO_MEGA.get(scene["studio"], scene.get("grail_tab", scene["studio"]))
    mega_path = f"mega:/Grail/{mega_studio}/{body.scene_id}/Description/"
    filename = f"{body.scene_id}_description.docx"

    # Build DOCX
    doc = Document()
    if body.title:
        h = doc.add_heading(body.title, level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    for para_text in [p.strip() for p in body.description.split("\n\n") if p.strip()]:
        p = doc.add_paragraph(para_text)
        p.runs[0].font.size = Pt(11)
    if body.meta_title or body.meta_description:
        doc.add_paragraph("")
        doc.add_heading("SEO Metadata", level=2)
        if body.meta_title:
            p = doc.add_paragraph()
            p.add_run("Meta Title: ").bold = True
            p.add_run(body.meta_title)
        if body.meta_description:
            p = doc.add_paragraph()
            p.add_run("Meta Description: ").bold = True
            p.add_run(body.meta_description)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    tmp_dir = tempfile.mkdtemp()
    try:
        (Path(tmp_dir) / filename).write_bytes(buf.read())
        r = subprocess.run(
            ["rclone", "copy", tmp_dir, mega_path],
            capture_output=True, text=True, timeout=120,
        )
        if r.returncode != 0:
            raise HTTPException(status_code=502, detail=f"rclone error: {r.stderr[:300]}")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return {"scene_id": body.scene_id, "mega_path": f"{mega_path}{filename}", "status": "saved"}


# ---------------------------------------------------------------------------
# Save description to Grail sheet (admin/grail-writer only)
# ---------------------------------------------------------------------------

class DescSaveGrailBody(BaseModel):
    scene_id: str
    description: str
    meta_title: str = ""
    meta_description: str = ""


@router.post("/save-grail")
async def save_description_to_grail(body: DescSaveGrailBody, user: CurrentUser):
    """
    Save a description directly to the Grail sheet. Requires grail-writer permission.
    Also marks the scene as having a description in SQLite.
    """
    if user["name"] not in {"Drew", "David", "Duc"}:
        raise HTTPException(status_code=403, detail="Grail write access required")

    with get_db() as conn:
        row = conn.execute(
            "SELECT id, grail_tab, grail_row FROM scenes WHERE id = ?",
            (body.scene_id,),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Scene not found")
        scene = dict(row)
        conn.execute("UPDATE scenes SET has_description=1 WHERE id=?", (body.scene_id,))

    # Fire-and-forget Grail write (description goes in column 8 per Grail layout)
    import threading
    threading.Thread(
        target=_write_desc_to_grail,
        args=(scene["grail_tab"], scene["grail_row"], body.description),
        daemon=True,
    ).start()

    return {"scene_id": body.scene_id, "status": "saved_to_grail"}


def _write_desc_to_grail(grail_tab: str, grail_row: int, description: str) -> None:
    """Write description to the Grail sheet (background safe)."""
    try:
        from api.sheets_client import open_grail, with_retry
        sh = open_grail()
        ws = sh.worksheet(grail_tab)
        # Description is typically in column 8 (H) in the Grail sheet
        with_retry(lambda: ws.update_cell(grail_row, 8, description))
        _log.info("Grail description write: %s row %d", grail_tab, grail_row)
    except Exception:
        _log.exception("Failed to write description to Grail: %s R%d", grail_tab, grail_row)
