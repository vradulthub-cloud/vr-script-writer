import io
import os
import re
import time
import requests as _requests
import streamlit as st
from script_writer import (
    SYSTEM_PROMPT, build_prompt, NJOI_STATIC_PLOT,
    validate_script, get_ollama_client, OLLAMA_BASE_URL, OLLAMA_MODEL,
    research_scene_trends,
)

# ── Title refinement helper ────────────────────────────────────────────────────
def _refine_treatment(treatment_name: str, png_bytes: bytes, prompt: str, title: str, seed: int):
    """Apply user style prompt to an existing rendered treatment.
    Fast path: PIL color/brightness transforms for simple keywords.
    Slow path: Ollama rewrites the treatment function for complex requests.
    """
    import re, inspect, requests as _req, random, io, types
    from PIL import Image, ImageEnhance, ImageFilter
    import numpy as np

    prompt_l = prompt.lower().strip()
    img = Image.open(io.BytesIO(png_bytes)).convert("RGBA")
    r_ch, g_ch, b_ch, a_ch = img.split()

    # ── Fast path: keyword transforms ─────────────────────────────────────────
    handled = False

    # Brightness
    if any(w in prompt_l for w in ["darker", "dark", "dim"]):
        rgb = Image.merge("RGB", (r_ch, g_ch, b_ch))
        rgb = ImageEnhance.Brightness(rgb).enhance(0.65)
        img = Image.merge("RGBA", (*rgb.split(), a_ch))
        handled = True
    elif any(w in prompt_l for w in ["brighter", "lighter", "light", "bright"]):
        rgb = Image.merge("RGB", (r_ch, g_ch, b_ch))
        rgb = ImageEnhance.Brightness(rgb).enhance(1.45)
        img = Image.merge("RGBA", (*rgb.split(), a_ch))
        handled = True

    # Saturation
    if any(w in prompt_l for w in ["vivid", "vibrant", "saturate", "saturated", "pop"]):
        rgb = Image.merge("RGB", (r_ch, g_ch, b_ch))
        rgb = ImageEnhance.Color(rgb).enhance(1.8)
        img = Image.merge("RGBA", (*rgb.split(), a_ch))
        r_ch, g_ch, b_ch, a_ch = img.split()
        handled = True
    elif any(w in prompt_l for w in ["muted", "desaturate", "faded", "washed"]):
        rgb = Image.merge("RGB", (r_ch, g_ch, b_ch))
        rgb = ImageEnhance.Color(rgb).enhance(0.35)
        img = Image.merge("RGBA", (*rgb.split(), a_ch))
        r_ch, g_ch, b_ch, a_ch = img.split()
        handled = True

    # Color tint
    COLOR_MAP = {
        "gold": (255, 200, 20), "golden": (255, 200, 20),
        "red": (255, 30, 30), "crimson": (180, 0, 30),
        "blue": (30, 100, 255), "electric blue": (0, 160, 255),
        "green": (0, 210, 60), "lime": (120, 255, 0),
        "purple": (160, 0, 255), "violet": (140, 0, 220),
        "pink": (255, 60, 180), "hot pink": (255, 20, 150),
        "orange": (255, 120, 0), "amber": (255, 160, 0),
        "cyan": (0, 220, 255), "teal": (0, 180, 180),
        "silver": (200, 205, 215), "white": (255, 255, 255),
        "black": (20, 20, 20), "yellow": (255, 230, 0),
        "rose": (255, 100, 120), "magenta": (255, 0, 200),
    }
    for color_name, tint_rgb in COLOR_MAP.items():
        if color_name in prompt_l:
            arr = np.array(img).astype(np.float32)
            alpha = arr[:, :, 3:4] / 255.0
            tint = np.array(tint_rgb, dtype=np.float32)
            strength = 0.55
            arr[:, :, :3] = arr[:, :, :3] * (1 - strength * alpha) + tint * strength * alpha
            img = Image.fromarray(np.clip(arr, 0, 255).astype(np.uint8), "RGBA")
            handled = True
            break

    if handled:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue(), treatment_name + "_refined"

    # ── Slow path: Ollama rewrites the function ────────────────────────────────
    # Strip any _refined suffix to get the original treatment name
    _base_name = re.sub(r"_(ai_)?refined$", "", treatment_name)
    fn = _cta.TREATMENTS.get(_base_name)
    if not fn:
        return None, None
    try:
        source = inspect.getsource(fn)
    except Exception:
        return None, None

    system = """You modify Python PIL graphic treatment functions based on user style requests.
Rules: keep the same function signature, keep transparent RGBA output, use rng for randomization.
Output ONLY the modified Python function — no markdown, no explanation."""

    user_msg = f"""Current function:
{source}

User wants: "{prompt}"

Rewrite the function applying these changes. Function name must stay the same."""

    try:
        resp = _req.post("http://localhost:11434/api/generate", json={
            "model": "qwen2.5:14b",
            "prompt": system + "\n\n" + user_msg,
            "stream": False,
            "options": {"temperature": 0.5, "num_predict": 1400}
        }, timeout=90)
        code = resp.json().get("response", "")
    except Exception:
        return None, None

    # Strip markdown fences
    import re as _re
    code = _re.sub(r"```python\s*", "", code)
    code = _re.sub(r"```\s*", "", code)
    m = _re.search(r"def\s+render_\w+\b", code)
    if not m:
        return None, None
    code = code[m.start():]
    nxt = _re.search(r"\ndef\s+\w", code[5:])
    if nxt:
        code = code[:nxt.start() + 5]

    # Execute the modified function in cta_generator's namespace
    try:
        ns = vars(_cta).copy()
        exec(compile(code, "<refined>", "exec"), ns)
        fn_name = _re.search(r"def\s+(render_\w+)\b", code).group(1)
        refined_fn = ns[fn_name]
        rng = random.Random(seed)
        out_img = refined_fn(title, rng)
        buf = io.BytesIO()
        out_img.save(buf, format="PNG")
        return buf.getvalue(), treatment_name + "_ai_refined"
    except Exception:
        return None, None

try:
    from call_sheet import get_budget_tabs, get_shoot_dates, generate_call_sheet, reload_script_cache
    HAS_CALL_SHEET = True
except ImportError:
    HAS_CALL_SHEET = False

from sheets_integration import (
    get_spreadsheet, month_tabs,
    rows_needing_scripts, write_script, parse_script_text,
    find_row_for_shoot, mark_talent_for_regen,
    STATUS_REGEN,
    COL_DATE, COL_STUDIO, COL_LOCATION, COL_SCENE,
    COL_FEMALE, COL_MALE, COL_PLOT, COL_THEME,
    COL_TITLE,
)

# ── Auto title generation ─────────────────────────────────────────────────────
_TITLE_GEN_SYSTEMS = {
    "VRHush": """You are a creative title writer for VRHush, a premium VR adult content studio.
Generate exactly ONE scene title. Rules:
- 2-3 words ONLY (never more than 3)
- Clever double-entendres, wordplay, or innuendo preferred
- Should hint at the scene's theme/action without being too literal
- Catchy, memorable, punchy — think movie title energy
- No performer names in the title
- No generic titles like "Hot Sex" or "Getting Laid"

Recent VRH titles for style reference: Heat By Design, Born To Breed, Under Her Spell, Intimate Renderings, Risqué Renter, Content Cutie, She Blooms on Command, Nailing the Interview, Kneading Your Pudding, Stretching Her Limits, The Honeymooners, Artists in Love

Respond with ONLY the title — no quotes, no explanation.""",

    "FuckPassVR": """You are a creative title writer for FuckPassVR, a premium VR adult travel/adventure content studio.
Generate exactly ONE scene title. Rules:
- 2-5 words (can be longer than VRH, more narrative)
- Travel/destination themes when applicable
- Clever wordplay, double-entendres preferred
- Should hint at the scene's theme or location
- No performer names in the title

Recent FPVR titles for style reference: The Grind Finale, Eager Beaver, Deep Devotion, Dirty Bunny Dancing, Reserved For Your Eyes, Pressing Dripping Matters, Fully Seated Affair, Fucking Like The Bulls, The Night is Young, Behind the Curtain, The Bouncing Layover

Respond with ONLY the title — no quotes, no explanation.""",

    "VRAllure": """You are a creative title writer for VRAllure, a premium VR solo/intimate content studio.
Generate exactly ONE scene title. Rules:
- 2-3 words ONLY
- Sensual, intimate, soft tone
- Suggestive but elegant — not crass
- Should hint at the solo/intimate nature
- No performer names in the title

Recent VRA titles for style reference: Sweet Surrender, Rise and Grind, Always on Top, Between the Sheets, A Swift Release, Potent Curves, The Wettest Seduction, Unhurried & Undressed, She Came to Play, Hovering With Intent

Respond with ONLY the title — no quotes, no explanation.""",

    "NaughtyJOI": """You are a creative title writer for NaughtyJOI, a premium VR JOI (jerk-off instruction) studio.
Generate a PAIRED title: "[Name] [action phrase]" then "[Name] then [contrasting action phrase]"
- First title: soft/teasing action (e.g. "draws you near", "sets the pace")
- Second title: intense/dominant contrasting action (e.g. "then sets you free", "then dares you to keep up")
- Use the performer's first name only

Recent NJOI titles: Lulu Chu builds you up / Lulu Chu then lets you fall, River Lynn starts with soft encouragement / River Lynn then strips away every ounce of your self-control

Respond with ONLY the two titles separated by a newline — no quotes, no explanation.""",
}


def _generate_title(studio, female, theme, plot, description=""):
    """Generate a scene title using Claude based on script content."""
    _claude_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not _claude_key:
        return None

    # Map app studio names to system prompt keys
    _studio_map = {"VRHush": "VRHush", "FuckPassVR": "FuckPassVR",
                   "VRAllure": "VRAllure", "NaughtyJOI": "NaughtyJOI"}
    _sys = _TITLE_GEN_SYSTEMS.get(_studio_map.get(studio, "VRHush"), _TITLE_GEN_SYSTEMS["VRHush"])

    _user = f"""Generate a title for this scene:

Performer: {female}
Theme: {theme}
Plot summary: {plot[:500] if plot else 'N/A'}
{f'Description: {description[:300]}' if description else ''}

Generate the title now."""

    try:
        import anthropic as _anth
        _ac = _anth.Anthropic(api_key=_claude_key)
        _resp = _ac.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=50,
            system=_sys,
            messages=[{"role": "user", "content": _user}]
        )
        return _resp.content[0].text.strip().strip('"').strip("'")
    except Exception:
        return None


def _write_title_to_grail(studio, scene_num, title):
    """Write a generated title to the Grail spreadsheet."""
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        creds = Credentials.from_service_account_file(
            os.path.join(os.path.dirname(__file__), "service_account.json"),
            scopes=["https://www.googleapis.com/auth/spreadsheets"]
        )
        gc = gspread.authorize(creds)
        _GRAIL_ID = "1Eq5G5FU6A8EqeFZCnZjrEaMYS8F1DiK5vP5tCSINeJk"
        sh = gc.open_by_key(_GRAIL_ID)
        # Map studio to Grail tab
        _tab_map = {"VRHush": "VRH", "FuckPassVR": "FPVR", "VRAllure": "VRA", "NaughtyJOI": "NNJOI"}
        tab = _tab_map.get(studio, "VRH")
        ws = sh.worksheet(tab)
        # Find row by scene number (col B = index 1)
        cells = ws.col_values(2)  # col B = scene numbers
        for i, val in enumerate(cells):
            if val.strip() == str(scene_num).strip():
                ws.update_cell(i + 1, 4, title)  # col D = title
                return True, f"Saved '{title}' → {tab} row {i + 1}"
        return False, f"Scene #{scene_num} not found in {tab}"
    except Exception as e:
        return False, f"Error: {e}"


def _write_grail_cell(grail_tab, grail_row, column_1based, value):
    """Write a single cell to the Grail sheet by tab name and row number.
    grail_tab/grail_row come from asset_tracker scene data.
    Column mapping: Title=4, Categories=6, Tags=7 (1-based)."""
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        creds = Credentials.from_service_account_file(
            os.path.join(os.path.dirname(__file__), "service_account.json"),
            scopes=["https://www.googleapis.com/auth/spreadsheets"]
        )
        gc = gspread.authorize(creds)
        _GRAIL_ID = "1Eq5G5FU6A8EqeFZCnZjrEaMYS8F1DiK5vP5tCSINeJk"
        sh = gc.open_by_key(_GRAIL_ID)
        ws = sh.worksheet(grail_tab)
        ws.update_cell(grail_row, column_1based, value)
        return True, f"Saved to {grail_tab} row {grail_row} col {column_1based}"
    except Exception as e:
        return False, f"Error: {e}"


def _write_title_to_scripts_sheet(ws_title, row_idx, title):
    """Write a generated title to the Scripts sheet (col K)."""
    try:
        ws = get_spreadsheet().worksheet(ws_title)
        ws.update_cell(row_idx, COL_TITLE + 1, title)  # COL_TITLE is 0-indexed, update_cell is 1-indexed
        return True
    except Exception as e:
        return False


# Models that are NOT for script writing — hidden from the script model picker
_NON_SCRIPT_MODELS = {"llava:7b", "llava", "qwen2.5-coder:14b", "qwen2.5-coder:7b",
                      "nomic-embed-text:latest", "nomic-embed-text"}

# Human-friendly labels for known models
_MODEL_LABELS = {
    "vr-scriptwriter":       "VR Script Writer  ★ recommended",
    "vr-scriptwriter:latest":"VR Script Writer  ★ recommended",
    "qwen2.5:14b":           "Qwen 2.5 14B  — general fallback",
    "qwen2.5:7b":            "Qwen 2.5 7B  — general (fast)",
    "dolphin-llama3:8b":     "Dolphin LLaMA3 8B  — fast",
    "dolphin-mistral:latest":"Dolphin Mistral  — fast",
    "llama3.2:latest":       "LLaMA 3.2  — general",
    "llama3.1:latest":       "LLaMA 3.1  — general",
}

@st.cache_data(ttl=300)
def _list_ollama_models() -> tuple[list[str], list[str]]:
    """Return (raw_ids, display_labels) for script-writing models only."""
    try:
        base  = OLLAMA_BASE_URL.replace("/v1", "")
        resp  = _requests.get(f"{base}/api/tags", timeout=3)
        all_m = [m["name"] for m in resp.json().get("models", [])]
    except Exception:
        all_m = ["vr-scriptwriter", OLLAMA_MODEL]
    script_models = [m for m in all_m if m not in _NON_SCRIPT_MODELS]
    if not script_models:
        script_models = [OLLAMA_MODEL]
    labels = [_MODEL_LABELS.get(m, m) for m in script_models]
    return script_models, labels

@st.cache_data
def _checkerboard_bg(w: int, h: int, sq: int = 16):
    """Build a checkerboard RGBA image, cached by (w, h)."""
    import numpy as np
    arr = np.full((h, w, 4), (40, 40, 45, 255), dtype=np.uint8)
    for y in range(0, h, sq):
        for x in range(0, w, sq):
            if (x // sq + y // sq) % 2 == 0:
                arr[y:y+sq, x:x+sq] = (55, 55, 60, 255)
    from PIL import Image as _Img
    return _Img.fromarray(arr, "RGBA")

st.set_page_config(page_title="Eclatech Hub", page_icon="🎬", layout="wide")


# ── Design System ─────────────────────────────────────────────────────────────
import hub_ui
st.markdown(hub_ui.global_css(), unsafe_allow_html=True)
_C = hub_ui.COLORS  # shorthand for inline style references

# ── Globally cached data loaders (shared across all sessions) ────────────────
# Data loads once, serves everyone for 30 min. Refresh button clears on demand.

@st.cache_data(ttl=1800, show_spinner="Loading asset data...")
def _cached_load_assets(_studios_tuple, limit):
    import asset_tracker as _at_inner
    studios = list(_studios_tuple) if _studios_tuple else None
    return _at_inner.load_asset_status(studios=studios, limit_per_studio=limit)


@st.cache_data(ttl=1800, show_spinner="Loading approvals...")
def _cached_load_approvals():
    import approval_tools as _apr_inner
    return _apr_inner.load_approvals()


@st.cache_data(ttl=1800, show_spinner="Loading tickets...")
def _cached_load_tickets():
    import ticket_tools as _tkt_inner
    return _tkt_inner.load_tickets()


# ── Ticket linking helper ────────────────────────────────────────────────────
def _ticket_linker(tab_key):
    """Show a ticket linking dropdown. Returns the linked ticket ID or None."""
    _lk_key = f"linked_ticket_{tab_key}"
    # Derive open tickets from global cache if available, otherwise quick load
    _all = _cached_load_tickets()
    _open = [t for t in _all if t["status"] not in ("Closed", "Rejected")]
    if not _open:
        return None
    _opts = ["No linked ticket"] + [f"{t['id']} — {t['title'][:40]}" for t in _open]
    _sel = st.selectbox("🔗 Link to ticket", _opts, key=_lk_key, label_visibility="collapsed")
    if _sel and _sel != "No linked ticket":
        _tid = _sel.split(" — ")[0]
        _match = next((t for t in _open if t["id"] == _tid), None)
        if _match:
            _sc = {"New": _C["blue"], "Approved": _C["green"], "In Progress": _C["amber"],
                   "In Review": _C["accent"]}.get(_match["status"], _C["muted"])
            st.markdown(
                f"<div style='font-size:0.7rem;color:{_sc}'>● Linked: {_match['status']}</div>",
                unsafe_allow_html=True
            )
        return _tid
    return None


def _ticket_progress(ticket_id, notes="", by=""):
    """Progress a linked ticket and clear cache."""
    if not ticket_id:
        return
    try:
        import ticket_tools as _tkt_p
        _tkt_p.progress_ticket(ticket_id, new_status="In Progress", notes=notes, by=by)
        st.session_state.pop("_open_tickets_cache", None)
        _cached_load_tickets.clear()
    except Exception:
        pass

# ── Authentication Gate ───────────────────────────────────────────────────────
import auth_config
from auth_config import get_user_permissions, get_allowed_tabs, is_admin
import notification_tools

if not st.user.is_logged_in:
    hub_ui.login_page()
    _lc1, _lc2, _lc3 = st.columns([1, 1, 1])
    with _lc2:
        if st.button("Sign in with Google", width="stretch"):
            st.login()
    st.stop()

_auth_email = st.user.email.lower()
_auth_user = get_user_permissions(_auth_email)

if _auth_user is None:
    hub_ui.denied_page(st.user.email)
    _dc1, _dc2, _dc3 = st.columns([1, 1, 1])
    with _dc2:
        if st.button("Sign out", width="stretch"):
            st.logout()
    st.stop()

_user_name = _auth_user["name"]
_user_is_admin = is_admin(_auth_user)
_user_can_write_grail = auth_config.is_grail_writer(_auth_user)
_user_can_manage_users = auth_config.is_user_manager(_auth_user)
_user_allowed_tabs = get_allowed_tabs(_auth_user)

# ── Notifications (cached) ────────────────────────────────────────────────────
@st.cache_data(ttl=600, show_spinner=False)
def _cached_unread_count(user_name):
    return notification_tools.get_unread_count(user_name)

@st.cache_data(ttl=600, show_spinner=False)
def _cached_notifications(user_name, limit=30):
    return notification_tools.load_notifications(recipient=user_name, limit=limit)

try:
    _unread_count = _cached_unread_count(_user_name)
except Exception:
    _unread_count = 0

# ── Header ────────────────────────────────────────────────────────────────────
_hdr1, _hdr2, _hdr3 = st.columns([5, 0.4, 1])
with _hdr1:
    hub_ui.logo_header(_user_name, unread_count=_unread_count)
with _hdr2:
    _bell_label = f"({_unread_count})" if _unread_count > 0 else ""
    if st.button(f"🔔{_bell_label}", key="notif_bell", help="Notifications"):
        st.session_state["_notif_open"] = not st.session_state.get("_notif_open", False)
        st.rerun()
with _hdr3:
    if st.button("Sign out", key="logout_btn"):
        st.logout()

# ── Notification panel (slides open below header) ────────────────────────────
if st.session_state.get("_notif_open", False):
    with st.container(border=True):
        _nc1, _nc2, _nc3 = st.columns([3, 2, 1])
        with _nc1:
            st.markdown(f"**Notifications** {'(' + str(_unread_count) + ' unread)' if _unread_count else ''}")
        with _nc2:
            if _unread_count > 0 and st.button("Mark all read", key="notif_mark_all"):
                notification_tools.mark_all_read(_user_name)
                _cached_unread_count.clear()
                _cached_notifications.clear()
                st.rerun()
        with _nc3:
            if st.button("Close", key="notif_close"):
                st.session_state["_notif_open"] = False
                st.rerun()

        try:
            _notifs = _cached_notifications(_user_name)
        except Exception:
            _notifs = []
        if _notifs:
            hub_ui.notification_panel(_notifs)
        else:
            st.caption("You're all caught up — no notifications.")

# ── Slop phrase substitutions (compiled once at module level) ─────────────────
_SLOP_SUBS = [
    (re.compile(r'\b(?:their\s+)?eyes?\s+meets?\b', re.I), "she holds his gaze"),
    (re.compile(r'\b(?:their\s+)?eyes?\s+met\b',   re.I), "she held his gaze"),
    (re.compile(r'\beyes?\s+locking\b',             re.I), "she holds his gaze"),
    (re.compile(r'\beyes?\s+locked\b',              re.I), "she held his gaze"),
    (re.compile(r'\bthe air between them\b',     re.I), "the quiet between them"),
    (re.compile(r'\bundeniable chemistry\b',     re.I), "a pull they've both felt building"),
    (re.compile(r'\bmagnetic attraction\b',      re.I), "a pull neither moves to break"),
    (re.compile(r'\bpalpable tension\b',         re.I), "a quiet charge in the room"),
    (re.compile(r'\bthe tension is palpable\b',  re.I), "the charge in the room is undeniable"),
    (re.compile(r'\bsparks fly\b',               re.I), "something shifts between them"),
    (re.compile(r'\belectric tension\b',         re.I), "a stillness charged with intent"),
    (re.compile(r'\bsuccumb(s|ed|ing)?\b',       re.I), "give in"),
    (re.compile(r'\bsteamy\b',                   re.I), "charged"),
    (re.compile(r'\bdesires intertwining\b',     re.I), "what comes next"),
    (re.compile(r'\bdesires intertwine\b',       re.I), "what comes next"),
    (re.compile(r'\bunspoken desire\b',          re.I), "what neither has said aloud"),
    (re.compile(r'\bpassion ignites\b',          re.I), "the moment breaks open"),
    (re.compile(r'\bunable to resist\b',         re.I), "past the point of stopping"),
    (re.compile(r'\bgive in to their desires\b', re.I), "act on it"),
    (re.compile(r'\bcan no longer be contained\b', re.I), "is past the point of stopping"),
    (re.compile(r'\bcharged atmosphere\b',       re.I), "the stillness in the room"),
    (re.compile(r'\blonging glances\b',          re.I), "the way she watches him"),
    (re.compile(r'\ba bottle of (?:wine|champagne|prosecco)\b', re.I), "a bottle of sparkling water"),
    (re.compile(r'\btwo wine glasses\b',          re.I), "two glasses of water"),
    (re.compile(r'\bwine glasses?\b',             re.I), "water glasses"),
    (re.compile(r'\bglass(?:es)? of (?:wine|champagne|prosecco)\b', re.I), "glass of water"),
    (re.compile(r'\bchardonnay\b',                re.I), "sparkling water"),
    (re.compile(r'\bprosecco\b',                  re.I), "sparkling water"),
    (re.compile(r'\bchampagne\b',                 re.I), "sparkling water"),
    (re.compile(r'\bred wine\b',                  re.I), "herbal tea"),
    (re.compile(r'\bwhite wine\b',                re.I), "sparkling water"),
    (re.compile(r'\brosé\b',                      re.I), "sparkling water"),
    (re.compile(r'\bcocktails?\b',                re.I), "coffee"),
    (re.compile(r'\bwhiskey\b',                   re.I), "coffee"),
    (re.compile(r'\bbourbon\b',                   re.I), "coffee"),
    (re.compile(r'\bwine rack\b',                 re.I), "bookshelf"),
    (re.compile(r'\bcooking wine\b',              re.I), "cooking broth"),
    (re.compile(r'\bwine\b',                      re.I), "sparkling water"),
    (re.compile(r'\bbeer\b',                      re.I), "sparkling water"),
    (re.compile(r'\bliquor\b',                    re.I), "coffee"),
    (re.compile(r'\brum\b',                       re.I), "coffee"),
    (re.compile(r'\bgin\b',                       re.I), "coffee"),
    (re.compile(r'\bvodka\b',                     re.I), "coffee"),
]

# ── Shared Description Config (used by Missing tab + Description tab) ────────
_DESC_STUDIO_CONFIG = {
    "FPVR": {"name": "FuckPassVR", "cta": "Watch {pronoun} on FuckPassVR now.", "prefix": "FPVR"},
    "VRH":  {"name": "VRHush",     "cta": "Taste {pronoun} on VRHush now.",     "prefix": "VRH"},
    "VRA":  {"name": "VRAllure",   "cta": "Watch {pronoun} on VRAllure now.",   "prefix": "VRA"},
    "NJOI": {"name": "NJOI",       "cta": "Watch {pronoun} on NJOI now.",       "prefix": "NJOI"},
}

_FPVR_CATEGORIES = [
    "8K", "Anal", "Asian", "Big Ass", "Big Tits", "Blonde", "Blowjob",
    "Body Cumshot", "Brunette", "Compilation", "Creampie", "Cum on Tits",
    "Curvy", "Ebony", "Facial Cumshot", "Hairy Pussy", "Handjob", "Latina",
    "MILF", "Natural Tits", "Petite", "Redhead", "Small Tits", "Threesome",
]

def _cat_slug(name):
    """'Big Ass' → 'big-ass'"""
    return name.lower().replace(" ", "-")

_STUDIO_DOMAINS = {
    "FPVR": "www.fuckpassvr.com",
    "VRH": "www.vrhush.com",
    "VRA": "www.vrallure.com",
    "NJOI": "www.naughtyjoi.com",
}

def _category_url(studio, cat_name):
    """Build category URL: https://domain/category/{slug}-vr-porn"""
    domain = _STUDIO_DOMAINS.get(studio, _STUDIO_DOMAINS["FPVR"])
    return f"https://{domain}/category/{_cat_slug(cat_name)}-vr-porn"

def _tag_url(studio, tag_name):
    """Build tag URL: https://domain/destination/tag/{slug}/"""
    domain = _STUDIO_DOMAINS.get(studio, _STUDIO_DOMAINS["FPVR"])
    slug = tag_name.lower().replace("'", "").replace(" ", "-")
    return f"https://{domain}/destination/tag/{slug}/"

def _cats_as_html(studio, cats_csv):
    """Convert comma-separated categories to HTML links."""
    cats = [c.strip() for c in cats_csv.split(",") if c.strip()]
    links = [f'<a href="{_category_url(studio, c)}">{c}</a>' for c in cats]
    return ", ".join(links)

def _tags_as_html(studio, tags_csv):
    """Convert comma-separated tags to HTML links."""
    tags = [t.strip() for t in tags_csv.split(",") if t.strip()]
    links = [f'<a href="{_tag_url(studio, t)}">{t}</a>' for t in tags]
    return ", ".join(links)

_FPVR_CATEGORY_URLS = {
    "8K": "https://www.fuckpassvr.com/category/8k-vr-porn",
    "Anal": "https://www.fuckpassvr.com/category/anal-vr-porn",
    "Asian": "https://www.fuckpassvr.com/category/asian-vr-porn",
    "Big Ass": "https://www.fuckpassvr.com/category/big-ass-vr-porn",
    "Big Tits": "https://www.fuckpassvr.com/category/big-tits-vr-porn",
    "Blonde": "https://www.fuckpassvr.com/category/blonde-vr-porn",
    "Blowjob": "https://www.fuckpassvr.com/category/blowjob-vr-porn",
    "Body Cumshot": "https://www.fuckpassvr.com/category/body-cumshot-vr-porn",
    "Brunette": "https://www.fuckpassvr.com/category/brunette-vr-porn",
    "Compilation": "https://www.fuckpassvr.com/category/compilation-vr-porn",
    "Creampie": "https://www.fuckpassvr.com/category/creampie-vr-porn",
    "Cum on Tits": "https://www.fuckpassvr.com/category/cum-on-tits-vr-porn",
    "Curvy": "https://www.fuckpassvr.com/category/curvy-vr-porn",
    "Ebony": "https://www.fuckpassvr.com/category/ebony-vr-porn",
    "Facial Cumshot": "https://www.fuckpassvr.com/category/facial-cumshot-vr-porn",
    "Hairy Pussy": "https://www.fuckpassvr.com/category/hairy-pussy-vr-porn",
    "Handjob": "https://www.fuckpassvr.com/category/handjob-vr-porn",
    "Latina": "https://www.fuckpassvr.com/category/latina-vr-porn",
    "MILF": "https://www.fuckpassvr.com/category/milf-vr-porn",
    "Natural Tits": "https://www.fuckpassvr.com/category/natural-tits-vr-porn",
    "Petite": "https://www.fuckpassvr.com/category/petite-vr-porn",
    "Redhead": "https://www.fuckpassvr.com/category/redhead-vr-porn",
    "Small Tits": "https://www.fuckpassvr.com/category/small-tits-vr-porn",
    "Threesome": "https://www.fuckpassvr.com/category/threesome-vr-porn",
}

_FPVR_TAGS = "8K, African, American, Anal Creampie, Analized, Arab, Argentinian, Asia, Asian, Ass, Ass Bounce, Ass Eating, Ass Fucking, Ass to Mouth, Ass Worship, Athletic, ATM, Australian, Average, Babe, Bald Pussy, Ball Sucking, Bangkok, Beauty Pageant, Belarusian, Belgian, Belgium, Berlin, Big Boobs, Big Fake Tits, Big Natural Tits, Big Oiled Tits, Black, Black Eyes, Blonde, Blue Eyes, Boobjob, Braces, Brazilian, British, Brown Eyes, Brunette, Brunette Fuck, Budapest, Bush, Butt Fuck, Butt Plug, Canadian, Caucasian, Chair Grind, Chile, Chilean, Chinese, Chubby, Clean Pussy, Clit Piercing, Close-up, Colombian, Columbian, Compilation, Cooking, Cowgirl, Creampie, Croatian, Cuban, Cum Eating, Cum in Mouth, Cum In Pussy, Cum on Ass, Cum on Face, Cum Play, Cum Swallow, Cumplay, Czech, Czech Republic, Dancing, Deep Anal, Deep Throat, Deep Throating, Deepthroat, Dick Sucking, Doggy, Doggy Style, Doggystyle, Dutch, Eating Ass, Ebony Babe, Escort, Euro Babe, European, Face Fucking, Face in Camera, Facefuck, Facial, Fake Tits, Farmer's Daughter, FFM porn, Filipina, Filipino, Finger Play, Fingering, Finnish, Fishnet Stockings, Foot Job, Footjob, Freckles, French, German, GFE, Glasses, Green Eyes, Grey Eyes, Grinding, Hair Pulling, Hairy Pussy, Hand Job, Hazel Eyes, Hispanic, Hot Tub, Humping, Hungarian, Intimate, Italian, Italy, Japanese, Jerk to Pop, Jizz Shot, Kenyan, Kissing, Landing Strip, Lap Dance, Latin, Latin Pussy, Latvian, Lingerie, Long Hair, Long Legs, Maid, Maltese, Massage, Massage Oil, Masseuse, Masturbation, Mexican, Mexico, Middle Eastern, Milf Porn, Missionary, Moldovan, Natural Boobs, Natural Tits, Navel Piercing, Nipple Piercing, Nipple Piercings, Nipple Play, Oil, Oil Massage, Oiled Tits, Oral Creampie, Panty Sniffing, Peruvian, Petite, PHAT Ass, Pierce Nipples, Pierced Clit, Pierced Nipples, Pierced Pussy, Pole Dancing, Polish, POV, POV BJ, Puerto Rican, Pull Out Cumshot, Pullout Cumshot, Pussy Eating, Pussy Fingering, Pussy Licking, Pussy Play, Pussy Spread, Pussy to Mouth, Pussy Worship, Redhead, Reverse Cowgirl, Rimjob, Roller Skates, Russian, Saudi, Saudi Arabian, Secretary, Septum Piercing, Sexy Asian, Sexy Blonde, Sexy Brunette, Sexy Ebony, Sexy Latina, Sexy Raven, Sexy Redhead, Sexy Redheads, Shaved, Shaved Pussy, Short Skirt, Sixty-nine, Slim, Sloppy Blowjob, Slovakian, Small Boobs, Small Natural Tits, South America, Spanish, Squirter, Standing Doggy, Standing Missionary, Stockings, Stripper, Stripper Pole, Stripping, Sucking Tits, Swallow, Syrian, Tall, Tattoo, Tattooed, Tattoos, Tease, Thailand, Tight Ass, Titjob, Tits in Face, Titty Fuck, Titty Sucking, Tittyjob, Toys, Trimmed, Trimmed Pussy, True 8K, Turkish, Turkish Babe, Twerking, Ukraine, Ukrainian, United States, Venezuelan, Vibrator, Wrestling"

_VRH_CATEGORIES = [
    "8K", "Anal", "Asian", "Big Ass", "Big Tits", "Blonde", "Blowjob",
    "Body Cumshot", "Brunette", "Compilation", "Creampie", "Cum in Mouth",
    "Cum on Tits", "Cumshot", "Curvy", "Ebony", "Facial Cumshot",
    "Hairy Pussy", "Handjob", "Hardcore", "Latina", "MILF", "Natural Tits",
    "Oral Creampie", "Petite", "Redhead", "Shaved Pussy", "Small Tits",
    "Threesome",
]
_VRH_TAGS = "8K, American, Anal, Analized, Arab, Asian, Ass, Ass Bounce, Ass Fucking, Ass Worship, Ass to Mouth, Athletic, ATM, Average, Bald Pussy, Ball Sucking, Big Boobs, Big Fake Tits, Big Natural Tits, Black, Blue Eyes, Brazilian, Brown Eyes, Brunette Fuck, Bush, Butt Fuck, Butt Plug, Canadian, Caucasian, Chair Grind, Chinese, Chubby, Clean Pussy, Clit Piercing, Close-up, Compilation, Cowgirl, Creampie, Cuban, Cum Eating, Cum in Mouth, Cum In Pussy, Cum on Ass, Cum on Face, Cum on Tits, Cum Play, Cum Swallow, Czech, Dancing, Deep Anal, Deep Throat, Deepthroat, Dick Sucking, Doggy Style, Doggystyle, Dutch, Ebony, Escort, Euro Babe, European, Facial, Fake Tits, Filipina, Filipino, Fingering, Fishnet Stockings, Footjob, Freckles, GFE, Glasses, Green Eyes, Grey Eyes, Hairy Pussy, Hazel Eyes, Hispanic, Intimate, Italian, Japanese, Jerk to Pop, Kissing, Landing Strip, Latin, Lingerie, Maltese, Massage, Massage Oil, Masseuse, Masturbation, Middle Eastern, Milf Porn, Missionary, Natural Boobs, Natural Tits, Navel Piercing, Oral Creampie, PHAT Ass, POV, POV BJ, Pierced Clit, Pierced Nipples, Polish, Pull Out Cumshot, Pullout Cumshot, Puerto Rican, Pussy Eating, Pussy Fingering, Pussy Licking, Pussy Play, Pussy Spread, Pussy Worship, Reverse Cowgirl, Rimjob, Roleplay, Romanian, Russian, Sexy Asian, Sexy Blonde, Sexy Brunette, Sexy Ebony, Sexy Latina, Sexy Raven, Sexy Redhead, Shaved, Shaved Pussy, Short Skirt, Sixty-nine, Slim, Slovakian, Small Boobs, Small Natural Tits, Spanish, Standing Missionary, Stockings, Stripper, Stripper Pole, Stripping, Syrian, Tattoos, Threesome, Tight Ass, Titjob, Tits in Face, Titty Fuck, Titty Sucking, Tittyjob, Toys, Trimmed, Trimmed Pussy, True 8K, Twerking, Ukrainian, Vibrator"

_VRA_CATEGORIES = [
    "Anal", "Asian", "Big Ass", "Big Tits", "Blonde", "Blowjob",
    "Brunette", "Compilation", "Curvy", "Ebony", "Hairy Pussy", "Handjob",
    "Latina", "Masturbation", "MILF", "Natural Tits", "Petite", "Redhead",
    "Sex Toys", "Shaved Pussy", "Small Tits",
]
_VRA_TAGS = "8K, American, Anal, Arab, Asian, Ass, Ass Bounce, Ass Spread, Ass Worship, Athletic, Australian, Average, Bald Pussy, Big Boobs, Big Fake Tits, Big Natural Tits, Black, Blue Eyes, Brazilian, Brown Eyes, Brunette Fuck, Bush, Butt Plug, Canadian, Caucasian, Chinese, Chubby, Clean Pussy, Clit Piercing, Close-up, Colombian, Compilation, Cowgirl, Creampie, Cuban, Curvy, Dancing, Deep Anal, Dick Sucking, Dildo Penetration, Doggy Style, Ebony, Euro Babe, European, Fake Tits, Filipina, Filipino, Fingering, Fishnet Stockings, Footjob, French, GFE, Glasses, Green Eyes, Grey Eyes, Hairy Pussy, Hazel Eyes, Intimate, Italian, Japanese, Jerk Off Instructions, Ken Doll, Kissing, Landing Strip, Latin, Latina, Lingerie, Maltese, Masturbation, Middle Eastern, Milf Porn, Missionary, Mongolian, Natural Boobs, Natural Tits, Navel Piercing, Outdoors, PHAT Ass, POV, POV BJ, Pierced Clit, Pierced Nipples, Polish, Puerto Rican, Pussy Eating, Pussy Fingering, Pussy Licking, Pussy Play, Pussy Spread, Pussy Worship, Pussylick, Reverse Cowgirl, Russian, Saudi, Sexy Asian, Sexy Blonde, Sexy Brunette, Sexy Ebony, Sexy Latina, Sexy Raven, Sexy Redhead, Shaved, Shaved Pussy, Sixty-nine, Slim, Small Boobs, Small Natural Tits, Solo, Spanish, Standing Missionary, Stockings, Stripper, Stripper Pole, Stripping, Syrian, Tattoos, Teens, Tight Ass, Titjob, Titty Fuck, Toys, Trimmed, Trimmed Pussy, True 8K, Twerking, Vibrator, Voyeur"

# Unified lookup for all studios
_STUDIO_CATEGORIES = {"FPVR": _FPVR_CATEGORIES, "VRH": _VRH_CATEGORIES, "VRA": _VRA_CATEGORIES}
_STUDIO_TAGS = {"FPVR": _FPVR_TAGS, "VRH": _VRH_TAGS, "VRA": _VRA_TAGS}

# ── Full system prompts per studio ────────────────────────────────────────────
_DESC_SYSTEMS_FULL = {}

_DESC_SYSTEMS_FULL["FPVR"] = """# PERSONALITY:
You are an expert adult copywriter specializing in crafting sexual, filthy, and deeply arousing scene descriptions for a Virtual Reality (VR) porn site called FuckPassVR. Your writing blends raw sexual energy with emotional depth and sensory immersion to create content that transports users into hyper-realistic, intimate encounters, making them feel as though they are part of the action. Your descriptions are optimized for search engines with a focus on VR adult content, captivating both human users and search algorithms.

# MAIN GOAL
The goal is to generate a sexually engaging and SEO optimized scene descriptions.

# WRITING STANDARDS:
1. Use active voice and powerful, visceral verbs to convey action and desire (e.g., "thrust," "devour," "crave"), enhancing the user's sense of agency in the VR space.
2. Incorporate erotic figurative language (metaphors, similes, personification) to elevate raw acts into poetic seduction, tailored for VR immersion (e.g., "her gaze pierces through the virtual haze, pulling you into her world").
4. Vary sentence structure for rhythmic intensity—short, sharp sentences for urgency; longer, flowing ones for sensual exploration—mimicking the ebb and flow of a VR encounter.
5. Focus on "show don't tell" to reveal desire through actions, physical reactions, and sensory cues, intensifying the first-person VR perspective (e.g., "your pulse races as her fingers graze your virtual skin").
6. Use language that feels authentic to the heat of the moment—raw, dirty, or tender as the scene demands—while weaving in VR-focused SEO-friendly terms naturally (e.g., "dive into a steamy VR sex fantasy with untamed passion").
7. Avoid any dialogue or spoken words, focusing solely on descriptive narration, internal user thoughts, and physical expressions to convey emotion and intent within the VR environment.
8. Maintain a seductive, provocative narrative voice that aligns with the tone of the work, whether raw and filthy or sensual and poetic, ensuring brand consistency for VR site SEO.
9. Balance raw eroticism with marketability, tailoring content to target VR porn audiences and optimizing for high-traffic VR adult keywords.

# EXAMPLES:
1. Ebony VR Porn: Chicago's Finest Gets Down and Dirty!
Is having a sultry ebony goddess grinding her big ass right in your face while performing a private dance one of your ultimate fantasies? Well, get ready because we've finally secured the queen of Chicago's underground scene - Ameena Green - for an exclusive VR porn experience that'll have you gripping your headset from start to finish! And trust us, this isn't just any private dance. This ebony VR porn masterpiece will show you exactly why Ameena's reputation for turning successful businessmen into drooling messes is well-earned. The moment she locks eyes with you in that exclusive club, you know you're in for a night that'll ruin all other VR experiences! Watch in stunning 8K VR as this chocolate goddess works her magic, those natural tits bouncing while she teases you with moves that should be illegal. Her wicked smile and seductive whispers are just the beginning - wait until you see what happens when the private room curtain closes and this cum-hungry queen shows you what she really does for her favorite clients!

Creampie VR Porn: When Private Dances Turn Extra Nasty!
How wild does it get? Let's just say that this creampie VR porn scene pushes boundaries you didn't even know existed! Inside this members-only paradise, Ameena transforms from sophisticated dancer to insatiable cock queen. Watch as she drops to her knees, treating your dick to a POV BJ that'll have you seeing stars. But that's just the warm-up! This ebony VR goddess takes control, mounting you in reverse cowgirl and working that hairy pussy on your cock like she's trying to earn a lifetime membership to your wallet. From intense standing missionary against the velvet walls to savage doggy style action that has her big ass clapping, every position proves why she's Chicago's best-kept secret. And when this cum-hungry goddess begs for you to flood her pussy? Well, let's just say resistance is futile! So grab your VR headset and dive into this exclusive ebony VR experience. After all, we're talking about the kind of private dance that makes every dollar spent in that club worth it - especially when it ends with Ameena's pussy dripping with your hot load! Don't miss out on the nastiest night Chicago has to offer!

2. Rouge Rendezvous: When Success Meets Seduction in Lyon
While we all know that celebrating a big business deal usually involves expensive champagne and fancy dinners, your colleagues in Lyon have something far more exciting in mind! In this big tits VR masterpiece, you'll find yourself in the city's most exclusive gentlemen's club, where the mesmerizing Anissa Kate is about to turn your victory celebration into an unforgettable private encounter. This French goddess, with her natural boobs and devilish smile, isn't your typical dancer - she's an artist of seduction who performs purely for the thrill of it. Watch in stunning 8K VR as she transforms your private dance into an intimate confession of desire. The moment she drops to her knees, it's clear this is no ordinary lap dance - her POV BJ skills prove she's mastered more than just stage moves, her skilled mouth and expert hands working in perfect harmony to drive you absolutely wild.

From Private Show to Passionate Creampie VR Porn
Inside this steamy creampie VR porn scene, you'll experience why French women have such a legendary reputation! Anissa takes complete control, mounting you in reverse cowgirl with an ass bounce that would make Paris proud. Her big tits sway hypnotically as she grinds against you, each movement building more intensity than the last. From deep standing missionary against the club's velvet walls to wild cowgirl rides that test the furniture's durability, every position showcases why she's Europe's most sought-after VR porn star. The passion reaches its peak in an intense doggy style session before she begs for that final cum in pussy finish. And trust us - when Anissa Kate demands a creampie, you don't say no! So grab your VR headset and prepare for the kind of private dance that makes every euro spent in Lyon worth it. After all, some business celebrations are better kept private, especially when they involve a French goddess who knows exactly how to make your success feel even sweeter!

3. Miami Heat: When Blonde VR Porn Dreams Ignite
Even though FuckPassVR specializes in creating mind-blowing virtual reality experiences, sometimes the hottest scenes come from the most unexpected situations. What does this mean for you? Well, you're about to discover how crashing on your old friend's couch in Miami turns into an unforgettable encounter with the stunning Thea Summers, a sexy blonde who's been harboring secret desires since your school days! Welcome to Tropic Like It's Hot - our latest 8K VR porn video that proves sometimes the best laid plans are the ones that aren't planned at all. After a night of vivid dreams about you, Thea brings you morning coffee only to find you sleeping naked on her couch. Watch as she seizes the moment, her small tits and toned body on display as she treats you to a POV BJ that'll make you forget all about that coffee getting cold.

Tropical Paradise: A Sizzling VR Porn Video Fantasy
This blonde VR porn scene explodes with raw passion as Thea takes control, mounting you reverse cowgirl with an ass bounce that defies gravity. Her sexy body becomes a work of art in motion as she spins around to face you, riding cowgirl style with an intensity that matches Miami's heat. The standing missionary position proves this fit beauty can handle whatever you give her, but it's when she gets on all fours that things really heat up. Watch her shaved pussy take every inch in doggy style before the intimate close-up missionary gives you the perfect view of what's to come. The grand finale sees her back on top, working you in cowgirl position until you flood her needy pussy with a hot load, leaving her dripping and satisfied. Who knew getting crashing on the couch could lead to such a wild ride? Let Thea Summers show you why sometimes the best plans are no plans at all in this stunning 8K VR porn experience.

4. Sinister Touches: When Yoga Meets Raw Desire in 8K VR Porn
Do we have any fitness enthusiasts among our VR porn viewers - or more precisely, someone who's dreamed of their yoga instructor taking things to the next level? Well, get ready because Maya Sinn is about to show you positions that definitely aren't in any traditional yoga manual. In this 8K VR porn scene, what starts as a typical training session quickly evolves into something far more enticing when your FuckPassVR passport catches Maya's attention. This European beauty might have started the day as your instructor, but she's about to become your personal sexual guru, trading meditation for pure, raw pleasure. Watch as she drops to her knees, taking your throbbing cock deep in her skilled mouth while incorporating that yoga ball in ways its manufacturers never intended, her POV BJ skills proving that flexibility isn't just for downward dog.

Hardcore Positions: A Cumshot VR Porn Masterclass
The real workout begins as Maya moves through positions that would make any yogi blush. She gets on all fours, that tight pussy begging to be filled as you pound her doggy style on the massage table. This cumshot VR porn scene showcases every inch of her sexual prowess as she takes you deep in missionary before mounting you in both cowgirl positions, her small tits bouncing with each thrust. The standing missionary proves this flexible vixen can handle an intense pounding, her shaved pussy gripping your cock until you're ready to explode. For the grand finale, Maya drops to her knees one last time, eager to earn her facial cumshot VR reward. Her pretty face becomes your canvas as you paint it with cum, proving some workouts are better done naked. Time to grab your VR headset and discover why this cum on face finish makes Sinister Touches an unforgettable session.

5. Your Ultimate Power Fantasy Awaits with Gizelle Blanco
Think we'd skip the billionaire's debauchery dream? Not a chance! At FuckPassVR, we turn boardroom triumphs into brunette VR porn paradise. Strap on your headset and become that tycoon celebrating Hawaii's biggest deal. Hidden behind velvet ropes in Hilo's most exclusive club, VR porn star Gizelle Blanco awaits - a raven-haired bombshell with big boobs that defy gravity and a big ass that rewrites temptation. Her pink lingerie glistens under moody lights as she whispers, "This private dance will ruin you for anyone else." Watch her electric striptease unravel, feel her skin under your roaming hands, then gasp as she drops to her knees. Her POV BJ engulfs your cock: sloppy dick sucking, throat-deep hunger, and eyes locked on yours like you're her last meal.

Cum-Worthy Finale in Jaw-Dropping 8K
This sexy brunette doesn't tease - she conquers. Mounting you in cowgirl, she rides hard, big boobs bouncing in your face while moans echo off soundproof walls. Then she spins, showcasing hypnotic reverse cowgirl action - that legendary ass bounce taunting you with every thrust. Against the stripper pole, standing missionary turns primal as she claws your back, taking reckless pumps. Doggystyle on all four on the chair? Her back arches, ass high, taking every punishing drive. When missionary shatters her into screaming orgasms, Gizelle makes gets you on the edge. The last cowgirl ride, makes you errupt on command! Kneeling before you, she strokes your cock until volcanic jizz shot erupts across her big, inviting tits - blasts of cum cascading over perfect curves in crystal 8K VR porn videos. This is how deals get sealed. Claim your filthy reward now! ONLY on FuckPassVR!"""

_DESC_SYSTEMS_FULL["VRH"] = """# PERSONALITY:
You are an expert adult copywriter specializing in crafting punchy, high-impact scene descriptions for VRHush, a premium VR porn studio. Your writing is raw, kinetic, and wastes zero words. Every sentence pushes the action forward. No scene-setting, no backstory - you drop the reader straight into the heat.

# MAIN GOAL
Generate a short, punchy, action-packed scene description optimized for VRHush's brand style.

# WRITING STANDARDS:
1. Single paragraph only. 100-140 words. No subheadings. No bold titles.
2. Open with the female performer doing something physical - no backstory, no "imagine," no setup.
3. Move through positions fast, one sentence each maximum.
4. Visceral, kinetic language: bouncing, slamming, gripping, moaning, dripping.
5. 2nd-person POV ("you") throughout. The male talent IS the viewer - NEVER refer to the male by name.
6. Mention wardrobe only if notable (lingerie, stockings, etc.).
7. Close with a one-liner: "[descriptor] in [resolution] VR porn. [CTA]"
8. Do NOT invent positions not in the plot.
9. No asterisks, bullet points, or markdown formatting.
10. No dialogue.
11. CRITICAL: In BG scenes, the male talent is YOU (the POV). Only the female performer gets named.

# EXAMPLES:
1. Kenzie Anne drops to her knees the second you walk through the door, wrapping those glossy lips around your cock like she's been starving for it. This blonde bombshell doesn't waste time - she's deepthroating you with sloppy, wet precision before climbing on top for a reverse cowgirl ride that puts her perfect ass on full display. She spins around, tits bouncing in your face as she grinds in cowgirl, then bends over the couch for doggy that has her screaming into the cushions. Standing missionary pins her against the wall, every thrust harder than the last. She finishes on her back in missionary, legs spread wide, taking every inch until you pull out and paint her stomach with a thick load. Pure filth in 8K VR porn. Taste her on VRHush now.

2. Liz Jordan's tight body is already on display when she peels off that lace set and drops into your lap. Her mouth finds your cock instantly - wet, messy, and eager. She mounts you reverse cowgirl, ass clapping with every bounce, then flips around for cowgirl with those perky tits pressed against you. Standing missionary has her pinned, moaning with each deep stroke. She gets on all fours for doggy, back arched, taking it hard and fast. The finale hits in missionary - her legs locked around you as you empty inside her with a deep creampie that leaves her trembling. Raw, unfiltered heat in 8K VR porn. Taste her on VRHush now.

3. Freya Parker greets you wearing nothing but a mischievous grin, and within seconds she's on her knees worshipping your cock with that signature sloppy enthusiasm. This petite stunner mounts up reverse cowgirl, her tiny frame bouncing impossibly fast on your shaft. Cowgirl brings those natural small tits right to your face as she rides with desperate urgency. She braces against the headboard for standing missionary, each thrust making her gasp. Doggy on the bed has her gripping the sheets, ass up, taking every punishing stroke. Missionary wraps it up - legs wide, eyes locked on yours as you pull out and blast a thick load across her pretty face. Unforgettable in 8K VR porn. Taste her on VRHush now."""

_DESC_SYSTEMS_FULL["VRA"] = """# PERSONALITY:
You are a sensual copywriter for VRAllure, a premium VR studio specializing in intimate solo and softcore content. Your writing is warm, tender, and deeply sensory - like a whispered confession. You focus on breath, touch, closeness, eye contact, and the electricity of being watched.

# MAIN GOAL
Generate a short, intimate, sensory-rich scene description for VRAllure's solo/intimate style.

# WRITING STANDARDS:
1. Single paragraph only. 60-90 words. No subheadings.
2. Intimate, whisper-close tone - not aggressive, not crude.
3. Focus on sensation: skin warmth, breath, fingertips, fabric, light.
4. These are typically solo/masturbation scenes - honor that intimacy.
5. 2nd-person POV ("you") - the viewer is a silent, invited observer.
6. Mention toys/props if in the plot.
7. Close with: "This [resolution] VR experience from VRAllure [sensory closing]."
8. Do NOT invent acts not in the plot.
9. No asterisks, bullet points, or markdown formatting.
10. No dialogue.

# EXAMPLES:
1. Skylar Vox settles onto silk sheets, sunlight tracing the curve of her waist as her fingers drift across her stomach. She takes her time, exploring herself with slow, deliberate touches - eyes half-closed, lips parting with each exhale. Her back arches as her hand slides between her thighs, hips rolling gently against her own rhythm. Every breath deepens, every movement more purposeful, until her whole body tightens and releases in a wave of quiet bliss. This 8K VR experience from VRAllure pulls you close enough to feel the warmth radiating from her skin. Watch her on VRAllure now.

2. Eliza Ibarra stretches across the bed in sheer white, letting the fabric pool around her hips as she starts to touch. Her fingertips trace circles on her inner thigh before slipping beneath the lace. The rhythm is unhurried - a slow burn that builds in the rise and fall of her chest. A vibrator hums softly as she presses it lower, her body responding with a shiver that travels from her toes to her parted lips. This 8K VR experience from VRAllure is pure intimacy captured in crystalline detail. Watch her on VRAllure now.

3. Lily Larimar lies back on the daybed, golden hour light pooling across her bare shoulders. She peels away a silk robe with no hurry, revealing her petite frame inch by inch. Her fingers find themselves, tracing slow paths across her small tits before dipping lower. Eyes flutter closed as her touch becomes more insistent, hips lifting gently off the cushion. The room is quiet except for the soft sounds of her breathing growing faster. This 8K VR experience from VRAllure lets you witness every shiver, every sigh. Watch her on VRAllure now."""

_DESC_SYSTEMS_FULL["NJOI"] = """# PERSONALITY:
You are a bold, teasing copywriter for NaughtyJOI (NJOI), a VR studio specializing in jerk-off instruction content. Your writing captures the push-pull dynamic of JOI - the performer talks directly to the viewer, guiding, teasing, commanding. You balance playfulness with intensity and always include at least one short performer quote.

# MAIN GOAL
Generate a short, JOI-focused scene description that captures the tease-build-release rhythm.

# WRITING STANDARDS:
1. Single paragraph only. 60-90 words. No subheadings.
2. Must include at least one short performer quote in double quotes.
3. JOI rhythm: tease, build, countdown, release.
4. Describe what she's wearing and removing.
5. Mention her voice, eye contact, and how she controls you.
6. 2nd-person POV ("you") throughout.
7. Playful, teasing, commanding tone.
8. Close with the studio CTA.
9. Do NOT invent acts not in the plot.
10. No asterisks, bullet points, or markdown.

# EXAMPLES:
1. Lulu Chu appears in a cropped tank and tiny shorts, eyes locked on yours with a knowing smile. She peels the tank away slowly, revealing her small natural tits as she whispers, "You don't get to touch - not yet." Her hand slides into her shorts, teasing herself while she counts you down, each number making you grip tighter. The shorts come off, and she spreads her legs wide, matching your pace stroke for stroke. "Faster," she commands, and you obey. The release hits like a wave when she finally says the word. Watch her on NJOI now.

2. River Lynn greets you in black lace, twirling for you before settling into the chair with her legs crossed. She uncrosses them slowly, giving you a peek before pulling back. "Think you can keep up?" She unclasps her bra, letting it fall as she begins to touch, guiding your rhythm with her voice. Faster, slower, stop - she controls every stroke. When the lace panties finally come off and she starts her countdown from ten, every second feels electric. Watch her on NJOI now.

3. Hazel Moore walks in wearing an oversized button-down, nothing underneath. She undoes each button like she's unwrapping a gift for you, maintaining eye contact the entire time. "I want you to go slow," she says, settling onto the bed and letting her hands wander. She mirrors your movements, building intensity until her breathing gets ragged. The countdown starts at five - short, urgent, breathless. When she hits zero, you both let go at the same time. Watch her on NJOI now."""


# ── Compilation system prompts (different style per studio) ───────────────────
_DESC_SYSTEMS_COMPILATION = {}

_DESC_SYSTEMS_COMPILATION["FPVR"] = """# PERSONALITY:
You are an expert adult copywriter for FuckPassVR writing compilation/best-of scene descriptions. Compilations are promotional "greatest hits" — you sell the CATEGORY, not a single narrative.

# WRITING STANDARDS:
1. Two paragraphs with bold subheadings. 200-300 words total.
2. Paragraph 1: Hook the category with a bold thesis. Name the series brand ("FuckPassVR Best [X] Adventures Volume [N]"). Sell the TYPE of performers collectively — archetypes, not individual stories. Superlative-heavy, promotional tone.
3. Paragraph 2: Tease what viewers will experience — the variety, the intensity, the production quality in 8K VR. Reference the number of performers. Build excitement for the collection without walking through individual scenes.
4. NEVER describe specific positions or scene-by-scene action. This is a highlight reel, not a walkthrough.
5. Reference 8K VR porn naturally. End with the studio CTA.
6. No dialogue, no asterisks, no bullet points.

# EXAMPLES:
1. **Best American Blonde Adventures: Your Ultimate Fantasy Lineup**
Think blondes have more fun? In 8K VR porn, they absolutely do — and FuckPassVR Best American Blonde Adventures Volume 1 is here to prove it. This compilation brings together the hottest blonde bombshells from across the United States, each one more irresistible than the last. From sun-kissed California babes to fiery East Coast stunners, every performer in this lineup was hand-picked to deliver the kind of raw, uninhibited passion that makes FuckPassVR the gold standard in virtual reality adult entertainment.

**8K VR Porn Blondes Who Redefine the Fantasy**
Whether your type is a petite spinner with a wicked smile or a curvy goddess who takes control, this compilation has your dream blonde waiting. Every encounter is captured in crystal-clear 8K, putting you inches away from the action as these American beauties work through an unforgettable range of positions and finishes. This isn't just a collection — it's the definitive blonde experience in VR porn. Watch them on FuckPassVR now."""

_DESC_SYSTEMS_COMPILATION["VRH"] = """# PERSONALITY:
You are an expert adult copywriter for VRHush writing compilation/best-of descriptions. VRHush compilations do a rapid per-performer walkthrough of highlights.

# WRITING STANDARDS:
1. Single paragraph. 120-180 words (longer than regular VRH, shorter than FPVR).
2. Open with a punchy hook for the category.
3. Walk through each performer with ONE sentence each — name + their standout moment/position.
4. Keep the kinetic VRH energy: fast, visceral, no fluff.
5. Close with: "[category descriptor] in 8K VR porn. [CTA]"
6. No asterisks, bullet points, or markdown.

# EXAMPLES:
1. The best curvy girls on VRHush, stacked into one relentless 8K VR compilation. Anissa Kate's teasing blowjob and reverse cowgirl ride kick off the action with those legendary natural tits bouncing in your face. Kitana Montana then takes over with intense standing missionary that puts her thick curves on full display. Mona Azar drops to her knees for a sloppy deepthroat before mounting you cowgirl, her big ass slamming down with every thrust. Karla Kush delivers a tight doggy session followed by a messy facial that leaves her grinning. Natasha Nice wraps it up with a slow-building ride that ends in the thickest creampie of the set. Five performers, five different flavors of curves — all captured in stunning 8K VR porn. Taste them on VRHush now."""

_DESC_SYSTEMS_COMPILATION["VRA"] = """# PERSONALITY:
You are a sensual copywriter for VRAllure writing compilation/best-of descriptions. Keep the intimate VRA tone but applied to a collection.

# WRITING STANDARDS:
1. Single paragraph. 80-120 words.
2. Sell the mood and sensation of the collection — breath, warmth, closeness across multiple performers.
3. Name each performer with one sensory detail each.
4. Close with the VRA-style ending.
5. No asterisks, bullet points, or markdown."""

_DESC_SYSTEMS_COMPILATION["NJOI"] = """# PERSONALITY:
You are a teasing copywriter for NaughtyJOI writing compilation/best-of descriptions.

# WRITING STANDARDS:
1. Single paragraph. 80-120 words.
2. Tease the variety of JOI styles — different voices, different commands, different paces.
3. Name each performer with their signature move or quote.
4. Build the tease-release rhythm across the whole collection.
5. Close with the NJOI CTA.
6. No asterisks, bullet points, or markdown."""


def _is_compilation(title):
    """Check if a scene title indicates a compilation."""
    if not title:
        return False
    return bool(re.search(r'\bVol\.?\s*\d|\bVolume\b|\bBest\s+Of\b|\bBest\s+\w|\bCompilation\b', title, re.I))


def _parse_desc_output(text):
    """Parse generated description into paragraphs + meta fields."""
    result = {"paragraphs": [], "meta_title": "", "meta_description": "", "raw": text}
    lines = text.strip().split("\n")

    current_title = ""
    current_body_lines = []
    in_meta = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check for meta fields (handles: "Meta Title:", "**Meta Title:**", "**Meta Title: **" etc.)
        _clean_lower = stripped.replace("*", "").replace("#", "").strip().lower()
        if _clean_lower.startswith("meta title"):
            # Save any pending paragraph
            if current_title or current_body_lines:
                result["paragraphs"].append({
                    "title": current_title,
                    "body": " ".join(current_body_lines).strip()
                })
                current_title = ""
                current_body_lines = []
            _mt_val = stripped.replace("*", "").strip()
            if ":" in _mt_val:
                _mt_val = _mt_val.split(":", 1)[1].strip()
            result["meta_title"] = _mt_val.strip('"').strip("'")
            in_meta = True
            continue
        if _clean_lower.startswith("meta description"):
            _md_val = stripped.replace("*", "").strip()
            if ":" in _md_val:
                _md_val = _md_val.split(":", 1)[1].strip()
            result["meta_description"] = _md_val.strip('"').strip("'")
            in_meta = True
            continue
        if in_meta:
            # Additional meta content on next line
            if not result["meta_description"]:
                result["meta_description"] = stripped
            continue

        in_meta = False

        # Check if line is a bold title (common patterns: **Title**, or short line followed by paragraph)
        is_title = False
        clean = stripped.strip("*").strip()
        if stripped.startswith("**") and stripped.endswith("**"):
            is_title = True
            clean = stripped.strip("*").strip()
        elif len(stripped) < 80 and not stripped.endswith(".") and not stripped.endswith("!") and len(stripped.split()) < 12:
            # Short line that doesn't end with period - likely a title
            # But only if we already have a paragraph or it's the first line
            if not current_body_lines:
                is_title = True
                clean = stripped

        if is_title:
            # Save previous paragraph if exists
            if current_title or current_body_lines:
                result["paragraphs"].append({
                    "title": current_title,
                    "body": " ".join(current_body_lines).strip()
                })
                current_body_lines = []
            current_title = clean
        else:
            current_body_lines.append(stripped)

    # Save last paragraph
    if current_title or current_body_lines:
        result["paragraphs"].append({
            "title": current_title,
            "body": " ".join(current_body_lines).strip()
        })

    # Post-processing: extract meta fields embedded inline in paragraph bodies
    import re as _re_parse
    for _p in result["paragraphs"]:
        _body = _p["body"]
        # Look for **Meta Title:** or Meta Title: inline
        _mt_match = _re_parse.search(r'[\u2014\-—]*\s*\*{0,2}Meta\s*Title:?\*{0,2}\s*[:\s]*(.+?)(?:\*{0,2}Meta\s*Desc|\Z)', _body, _re_parse.IGNORECASE)
        if _mt_match and not result["meta_title"]:
            _mt_raw = _mt_match.group(1).strip().rstrip("*").strip().strip('"').strip("'")
            result["meta_title"] = _mt_raw
        _md_match = _re_parse.search(r'\*{0,2}Meta\s*Description:?\*{0,2}\s*[:\s]*(.+)', _body, _re_parse.IGNORECASE)
        if _md_match and not result["meta_description"]:
            result["meta_description"] = _md_match.group(1).strip().rstrip("*").strip().strip('"').strip("'")
        # Remove meta fields from body text
        _body_clean = _re_parse.sub(r'[\u2014\-—]*\s*\*{0,2}Meta\s*Title:?\*{0,2}\s*[:\s]*.+?(?=\*{0,2}Meta\s*Desc|\Z)', '', _body, flags=_re_parse.IGNORECASE)
        _body_clean = _re_parse.sub(r'\*{0,2}Meta\s*Description:?\*{0,2}\s*[:\s]*.+', '', _body_clean, flags=_re_parse.IGNORECASE)
        _p["body"] = _body_clean.strip().rstrip("—").rstrip("-").strip()

    return result


def _reassemble_desc(parsed):
    """Reassemble parsed description back into text."""
    parts = []
    for p in parsed.get("paragraphs", []):
        if p["title"]:
            parts.append(f"**{p['title']}**")
        parts.append(p["body"])
        parts.append("")
    if parsed.get("meta_title"):
        parts.append(f'Meta Title: "{parsed["meta_title"]}"')
    if parsed.get("meta_description"):
        parts.append(f"Meta Description: {parsed['meta_description']}")
    return "\n".join(parts).strip()


def _build_scene_prompt(studio, cfg, title, female, male, plot, categories, model_props, sex_positions, target_keywords, resolution, wardrobe):
    """Build the full scene prompt for description generation. Studio-aware structure."""
    female_names = [n.strip() for n in female.split(",") if n.strip()]
    male_names = [n.strip() for n in male.split(",") if n.strip()]
    pronoun = "her" if not male_names else "them"
    cta = cfg["cta"].format(pronoun=pronoun)

    # POV rule: in BG scenes, male talent IS the viewer ("you") — never refer to male by name
    if male_names:
        pov_rule = f"""CRITICAL POV RULE: This is a VR porn scene shot from the male's point of view. The male talent ({', '.join(male_names)}) IS the viewer — he is "you". NEVER refer to the male talent by name or as a third person. The viewer IS the male. Only refer to the female performer(s) by name."""
        model_line = ", ".join(female_names) + f" (with you as the POV)"
    else:
        pov_rule = "This is a solo scene. The viewer watches as a silent, intimate observer."
        model_line = ", ".join(female_names) if female_names else "Unknown"

    # ── Studio-specific structure instructions ───────────────────────────────
    if studio == "FPVR":
        structure = f"""Structure: Create two paragraphs, each with a bolded, enticing title.
Paragraph 1 (The Reveal): Focus on the story setup, the surprise entrance/encounter, and the initial intimate contact. Build anticipation.
Paragraph 2 (The Fantasy): Describe the core VR experience. Chronologically flow through the key sex acts, emphasizing the visceral, close-up details that {resolution} VR provides. Build to the intense finish.
Word Count: 350-400 words total.

Output Format:
Paragraph 1 Title: [Bolded, Engaging Title Here]
Paragraph 1: [Final polished text]
Paragraph 2 Title: [Bolded, Engaging Title Here]
Paragraph 2: [Final polished text]
Meta Title: "{title or 'Scene Title'} VR Porn | {cfg['name']}"
Meta Description: (160 chars max) [Compelling description with primary keyword and CTA]"""

    elif studio == "VRH":
        structure = f"""Structure: Single paragraph only. 100-140 words. NO subheadings, NO bold titles.
Open with the performer doing something physical — no backstory, no "imagine," no setup.
Move through positions fast, one sentence each maximum.
Visceral, kinetic language: bouncing, slamming, gripping, moaning, dripping.
Close with a one-liner format: "[descriptor] in {resolution} VR porn. {cta}"
Word Count: 100-140 words. Single paragraph.

Output Format:
[Single paragraph — no title, no subheadings]
Meta Title: "{title or 'Scene Title'} VR Porn | {cfg['name']}"
Meta Description: (160 chars max) [Compelling description]"""

    elif studio == "VRA":
        structure = f"""Structure: Single paragraph only. 60-90 words. NO subheadings, NO bold titles.
Intimate, whisper-close tone — not aggressive. Focus on sensation: skin, warmth, breath, fingertips.
Close with: "This {resolution} VR experience from VRAllure [sensory closing]. {cta}"
Word Count: 60-90 words. Single paragraph.

Output Format:
[Single paragraph — no title, no subheadings]
Meta Title: "{title or 'Scene Title'} VR Porn | {cfg['name']}"
Meta Description: (160 chars max) [Compelling description]"""

    elif studio == "NJOI":
        structure = f"""Structure: Single paragraph only. 60-90 words. NO subheadings, NO bold titles.
JOI rhythm: tease, build, countdown, release. Must include at least one short performer quote in double quotes.
Describe what she's wearing and removing. Mention her voice, eye contact, and how she controls you.
Close with the CTA: "{cta}"
Word Count: 60-90 words. Single paragraph.

Output Format:
[Single paragraph — no title, no subheadings]
Meta Title: "{title or 'Scene Title'} VR Porn | {cfg['name']}"
Meta Description: (160 chars max) [Compelling description]"""

    else:
        structure = f"""Structure: Single paragraph, 100-150 words.
Word Count: 100-150 words.
End with: "{cta}"

Output Format:
[Single paragraph]
Meta Title: "{title or 'Scene Title'} VR Porn | {cfg['name']}"
Meta Description: (160 chars max)"""

    prompt = f"""Role: You are a senior copywriter for {cfg['name']}.

{pov_rule}

Core Scene Data:
Scene Title: {title or 'Untitled'}
Model: {model_line}
Plot: {plot}
Categories: {categories}
Model Properties: {model_props or 'N/A'}
Target Keywords: {target_keywords or f'{resolution} VR porn, {cfg["name"]}'}
Sex positions: {sex_positions or 'See plot above'}{f'''
Wardrobe: {wardrobe}''' if wardrobe else ''}

Writing Instructions:
Tone & Perspective: 2nd-person POV ("You") throughout. Intensely immersive.
SEO: Naturally integrate the target keywords. The site name "{cfg['name']}" must appear at least once.
Do NOT invent positions or acts not described in the plot/sex positions.
No asterisks, bullet points, or markdown formatting in the description body.

{structure}

Write the description now."""

    return prompt


# ── Main tabs (dynamic based on user permissions) ────────────────────────────
_visible_tabs = _user_allowed_tabs  # list of (key, label) tuples
if _visible_tabs:
    _tab_objs = st.tabs([label for _, label in _visible_tabs])
    _tab_map = {key: obj for (key, _), obj in zip(_visible_tabs, _tab_objs)}
else:
    _tab_map = {}
    st.info("No tabs available for your account. Contact Drew for access.")
# Use a no-op container for tabs the user can't see — the `with` block still
# executes syntactically but we immediately skip its body with a flag check.
_noop = st.container()
tab_tickets = _tab_map.get("Tickets", _noop)
tab_research = _tab_map.get("Model Research", _noop)
tab_scripts = _tab_map.get("Scripts", _noop)
tab_callsheet = _tab_map.get("Call Sheets", _noop)
tab_titles = _tab_map.get("Titles", _noop)
tab_desc = _tab_map.get("Descriptions", _noop)
tab_comp = _tab_map.get("Compilations", _noop)
_has_tab = _tab_map.__contains__

# ── TAB 1: Scripts (Manual + From Sheet, single + batch) ─────────────────────
with tab_scripts:
    if _has_tab("Scripts"):
        hub_ui.section("Script Generator")

        # ── Mode toggle ───────────────────────────────────────────────────────────
        mode = st.segmented_control("Mode", ["✏️ Manual", "📋 From Sheet"], default="✏️ Manual",
                                    key="sc_mode", label_visibility="collapsed")

        # ─────────────────────────────────────────────────────────────────────────
        # Shared generation helpers (used by both modes)
        # ─────────────────────────────────────────────────────────────────────────
        # _SLOP_SUBS is defined at module level (above tabs) to avoid recompiling 45 regexes every rerun

        def _post_process(raw):
            raw = re.sub(r'\s*[\u2014\u2013]\s*', ' ', raw)
            raw = re.sub(r'(\d+[A-Za-z]?)-(\d)', r'\1 \2', raw)
            raw = re.sub(r'(\w)-(\w)', r'\1 \2', raw)
            raw = re.sub(r'-', ' ', raw)
            # Replace banned slop phrases with better prose equivalents
            for pattern, replacement in _SLOP_SUBS:
                raw = pattern.sub(replacement, raw)
            return raw

        def _parse_and_validate(raw, female="", male=""):
            pt = raw if re.search(r'(?i)^THEME\s*:', raw.lstrip()) else "THEME: " + raw
            f = parse_script_text(pt)
            v = validate_script(f, female=female, male=male)
            return f, v


        def _build_parsed(studio, female, male, scene_type_val, destination, theme_hint):
            is_vra  = studio == "VRAllure"
            is_njoi = studio == "NaughtyJOI"
            if is_vra:
                return {"studio": "VRAllure", "female": female,
                        "vra_scene_type": scene_type_val,
                        "theme_hint": theme_hint or None}, scene_type_val
            if is_njoi:
                return {"studio": "NaughtyJOI", "female": female}, "JOI"
            scene_norm = "BGCP" if scene_type_val and ("CP" in scene_type_val.upper() or "CREAMPIE" in scene_type_val.upper()) else "BG"
            return {
                "studio": studio,
                "destination": (destination or "").strip() or None,
                "scene_type": scene_norm,
                "female": female,
                "male": male or "",
                "theme_hint": theme_hint or None,
            }, scene_norm

        def _do_save(ws_target, row_target, fields):
            write_script(ws_target, row_target,
                         theme=fields.get("theme", ""),
                         plot=fields.get("plot", ""),
                         wardrobe_female=fields.get("wardrobe_female", ""),
                         wardrobe_male=fields.get("wardrobe_male", ""),
                         set_design=fields.get("set_design", ""),
                         props=fields.get("props", ""))
            # Bust the row + sorted cache for the saved month so the ✓ column updates
            _saved_month = ws_target.title if ws_target else ""
            for _k in list(st.session_state.keys()):
                if _k in (f"sc_rows_{_saved_month}", f"sc_sorted_{_saved_month}"):
                    del st.session_state[_k]

        # ─────────────────────────────────────────────────────────────────────────
        # Result display (single script) — shared by both modes
        # ─────────────────────────────────────────────────────────────────────────
        def _show_single_result():
            if "last_script" not in st.session_state:
                return
            _s             = st.session_state["last_script"]
            full_text      = _s["full_text"]
            fields         = _s["fields"]
            violations     = _s["violations"]
            parsed         = _s["parsed"]
            _saved_studio  = _s["studio"]
            _saved_female  = _s["female"]
            _saved_male    = _s["male"]
            _saved_research= _s["research_context"]
            _saved_ws_title= _s["ws_title"]
            _saved_row_idx = _s["row_idx"]
            _saved_scene   = _s["scene_norm"]

            st.divider()

            # ── Status + Metrics row ─────────────────────────────────────────────
            _viol_count = len(violations)
            _has_plot = bool(fields.get("plot", ""))
            _has_theme = bool(fields.get("theme", ""))
            _has_wardrobe = bool(fields.get("wardrobe_female", ""))
            _script_assets = sum([_has_plot, _has_theme, _has_wardrobe, not bool(violations)])
            _script_total = 4
            _script_pct = int((_script_assets / _script_total) * 100)
            _pct_color = _C["green"] if _script_pct == 100 else (_C["amber"] if _script_pct >= 50 else _C["red"])

            _sm1, _sm2, _sm3 = st.columns(3)
            _sm1.metric("Status", "Pass" if not violations else f"{_viol_count} Issue(s)")
            _sm2.metric("Studio", _saved_studio)
            _sm3.metric("Scene", _saved_scene)

            # ── Script summary card ───────────────────────────────────────────────
            with st.container(border=True):
                _dest_line = (f"<tr><td style='color:{_C['muted']};padding:2px 12px 2px 0;font-size:0.82rem'>Destination</td>"
                              f"<td style='font-size:0.88rem'>{fields['destination']}</td></tr>"
                              if fields.get("destination") else "")
                _male_w    = fields.get("wardrobe_male", "")
                _male_line = (f"<tr><td style='color:{_C['muted']};padding:2px 12px 2px 0;font-size:0.82rem'>Male</td>"
                              f"<td style='font-size:0.88rem'>{_male_w}</td></tr>"
                              if _male_w else "")

                # Header with studio color bar
                _studio_colors_sc = {"VRHush": _C.get("vrh","#8b5cf6"), "FuckPassVR": _C.get("fpvr","#3b82f6"),
                                     "VRAllure": _C.get("vra","#ec4899"), "NaughtyJOI": _C.get("njoi","#f97316")}
                _sc_color = _studio_colors_sc.get(_saved_studio, _C["accent"])
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:10px;margin-bottom:8px'>"
                    f"<div style='width:4px;height:24px;border-radius:2px;background:{_sc_color}'></div>"
                    f"<span style='font-size:0.85rem;font-weight:700;color:{_C['text']}'>{_saved_female}</span>"
                    f"<span style='font-size:0.72rem;color:{_C['muted']};margin-left:auto'>"
                    f"{_saved_studio} · {_saved_scene}</span>"
                    f"</div>",
                    unsafe_allow_html=True
                )

                st.markdown(
                    f"<table style='border-collapse:collapse;width:100%'>"
                    f"{_dest_line}"
                    f"<tr><td style='color:{_C['muted']};padding:2px 12px 2px 0;font-size:0.82rem'>Theme</td>"
                    f"<td style='font-size:0.88rem;font-weight:600'>{fields.get('theme','—')}</td></tr>"
                    f"<tr><td style='color:{_C['muted']};padding:2px 12px 2px 0;font-size:0.82rem'>Female</td>"
                    f"<td style='font-size:0.88rem'>{fields.get('wardrobe_female','—')}</td></tr>"
                    f"{_male_line}"
                    f"</table>",
                    unsafe_allow_html=True
                )

                # ── Plot (editable) ────────────────────────────────────────────────
                _plot_text = fields.get("plot", "")
                if _plot_text:
                    _edited_plot = st.text_area(
                        "Plot", value=_plot_text, height=130,
                        key="sc_plot_edit", label_visibility="collapsed",
                    )
                    _s["fields"]["plot"] = _edited_plot
                    st.session_state["last_script"] = _s

                # ── Set / Props inline ─────────────────────────────────────────────
                _set  = fields.get("set_design", "")
                _prop = fields.get("props", "")
                if _set or _prop:
                    _sp_parts = []
                    if _set:  _sp_parts.append(f"<b style='color:{_C['muted']}'>Set</b> {_set}")
                    if _prop: _sp_parts.append(f"<b style='color:{_C['muted']}'>Props</b> {_prop}")
                    st.markdown(
                        f"<p style='font-size:0.82rem;color:{_C['text']};margin:6px 0'>"
                        + " &nbsp;·&nbsp; ".join(_sp_parts) + "</p>",
                        unsafe_allow_html=True)

            if violations:
                with st.expander(f"Rule violations ({_viol_count})", expanded=True):
                    for v in violations:
                        st.markdown(
                            f"<div style='background:{_C['red_dim']};border-left:3px solid {_C['red']};"
                            f"border-radius:4px;padding:6px 10px;margin:4px 0;font-size:0.82rem;"
                            f"color:{_C['text']}'>{v}</div>",
                            unsafe_allow_html=True)

            # ── Title generation ──────────────────────────────────────────────────
            with st.container(border=True):
                st.markdown(f"<div class='sh'>Title</div>", unsafe_allow_html=True)
                _tc1, _tc2, _tc3 = st.columns([2, 4, 1])
                with _tc1:
                    if st.button("Generate Title", width="stretch", key="sc_gen_title"):
                        with st.spinner("Generating title…"):
                            _gen_title = _generate_title(
                                _saved_studio, _saved_female,
                                fields.get("theme", ""), fields.get("plot", ""))
                            if _gen_title:
                                st.session_state["generated_title"] = _gen_title
                with _tc2:
                    _cur_title = st.text_input("Title", value=st.session_state.get("generated_title", ""),
                                               key="sc_title_edit", label_visibility="collapsed",
                                               placeholder="Generated title appears here…")
                with _tc3:
                    if _cur_title and st.button("Save", width="stretch", key="sc_save_title"):
                        if _saved_ws_title and _saved_row_idx:
                            _write_title_to_scripts_sheet(_saved_ws_title, _saved_row_idx, _cur_title)
                            st.success(f"Title saved: {_cur_title}")
                            st.session_state["generated_title"] = _cur_title

            # ── Actions ──────────────────────────────────────────────────────────
            _dl_col, _ = st.columns([1, 3])
            with _dl_col:
                st.download_button("Download Script", data=full_text,
                                   file_name=f"{_saved_studio}_{_saved_female.replace(' ','_')}_{_saved_scene}.txt",
                                   mime="text/plain", width="stretch")

            _a1, _a2, _a3 = st.columns([2, 2, 3])
            with _a1:
                _accept_label = "Accept & Save" if _user_is_admin else "Submit for Approval"
                if st.button(_accept_label, width="stretch", type="primary", key="sc_accept"):
                    if _user_is_admin:
                        # Admin: write directly
                        saved = False
                        try:
                            if _saved_ws_title and _saved_row_idx:
                                _ws = get_spreadsheet().worksheet(_saved_ws_title)
                                _do_save(_ws, _saved_row_idx, fields)
                                st.success(f"Saved → {_saved_ws_title}, row {_saved_row_idx}")
                                saved = True
                            else:
                                ws_found, row_idx = find_row_for_shoot(_saved_female, _saved_studio)
                                if ws_found and row_idx:
                                    _do_save(ws_found, row_idx, fields)
                                    st.success(f"Saved → {ws_found.title}, row {row_idx}")
                                    saved = True
                                else:
                                    st.info("No matching row — not saved to sheet.")
                        except Exception as _e:
                            st.error(f"Save failed: {_e}")
                        try:
                            from training_data import save_accepted
                            save_accepted(_saved_studio, parsed, fields, _saved_research)
                        except Exception as _te:
                            st.warning(f"Training data not saved: {_te}")
                        if saved:
                            del st.session_state["last_script"]
                            st.rerun()
                    else:
                        # Non-admin: submit for approval
                        try:
                            import json as _json_mod
                            _scene_id = f"{_saved_studio}{_saved_scene}"
                            _target = f"Scripts:{_saved_ws_title}:{_saved_row_idx}" if _saved_ws_title and _saved_row_idx else "Scripts:unknown:0"
                            _preview = f"Theme: {fields.get('theme','')}\nPlot: {fields.get('plot','')[:150]}"
                            _linked = st.session_state.get("scripts_linked_ticket", "")
                            import approval_tools as _apr_sc
                            _apr_id = _apr_sc.submit_for_approval(
                                submitted_by=_user_name,
                                content_type="script",
                                scene_id=_scene_id,
                                studio=_saved_studio,
                                content_preview=_preview,
                                content_json=_json_mod.dumps(fields),
                                target_sheet=_target,
                                linked_ticket=_linked,
                            )
                            try:
                                notification_tools.notify_approval_submitted(
                                    _apr_id, _scene_id, "script", _user_name)
                                _cached_unread_count.clear()
                            except Exception:
                                pass
                            st.success(f"Submitted for approval: **{_apr_id}**")
                            try:
                                from training_data import save_accepted
                                save_accepted(_saved_studio, parsed, fields, _saved_research)
                            except Exception:
                                pass
                            del st.session_state["last_script"]
                            st.rerun()
                        except Exception as _e:
                            st.error(f"Approval submission failed: {_e}")
            with _a2:
                feedback = st.text_input("Director's note", placeholder="What to change…",
                                         key="sc_regen_feedback", label_visibility="collapsed")
                if feedback:
                    st.session_state["director_note_override"] = feedback
            with _a3:
                if st.button("↩ Regenerate", width="stretch", key="sc_reject"):
                    try:
                        from training_data import save_rejected
                        save_rejected(_saved_studio, parsed, fields, feedback=feedback or "")
                    except Exception as _te:
                        st.warning(f"Training data not saved: {_te}")
                    del st.session_state["last_script"]
                    st.rerun()

        # ─────────────────────────────────────────────────────────────────────────
        # Shared: generate a script (non-streaming) — used by both single & batch
        # ─────────────────────────────────────────────────────────────────────────
        def _generate_script_core(parsed, female, male, research_context=""):
            """Generate a script via Claude or Ollama. Returns (full_text, fields, violations)."""
            prompt = build_prompt(parsed, research_context=research_context)
            full_text = ""
            _claude_key = os.environ.get("ANTHROPIC_API_KEY", "")
            if _claude_key:
                import anthropic as _anthropic
                _ac = _anthropic.Anthropic(api_key=_claude_key)
                _resp = _ac.messages.create(
                    model="claude-sonnet-4-6", max_tokens=600,
                    system=SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": prompt}],
                )
                full_text = _resp.content[0].text or ""
            else:
                _ollama = get_ollama_client()
                _model = st.session_state.get("selected_model", OLLAMA_MODEL)
                _resp = _ollama.chat.completions.create(
                    model=_model, max_tokens=800, temperature=0.82,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user",   "content": prompt},
                    ], stream=False,
                )
                full_text = _resp.choices[0].message.content or ""
            full_text = _post_process(full_text)
            fields, violations = _parse_and_validate(full_text, female=female, male=male)
            return full_text, fields, violations

        # ─────────────────────────────────────────────────────────────────────────
        # Core: research + stream + store result (single mode with streaming UI)
        # ─────────────────────────────────────────────────────────────────────────
        def _run_single_generation(parsed, studio, female, male, scene_norm,
                                    selected_ws, selected_row_idx, theme_hint):
            is_njoi = studio == "NaughtyJOI"

            if is_njoi:
                st.markdown(f"**NaughtyJOI Plot:**\n\n{NJOI_STATIC_PLOT}")
                try:
                    if selected_ws and selected_row_idx:
                        write_script(selected_ws, selected_row_idx, theme="JOI",
                                     plot=NJOI_STATIC_PLOT, wardrobe_female="", wardrobe_male="")
                        st.success(f"✅ Saved → {selected_ws.title}, row {selected_row_idx}")
                    else:
                        ws_f, row_f = find_row_for_shoot(female, studio)
                        if ws_f and row_f:
                            write_script(ws_f, row_f, theme="JOI",
                                         plot=NJOI_STATIC_PLOT, wardrobe_female="", wardrobe_male="")
                except Exception as _e:
                    st.warning(f"Could not save: {_e}")
                return

            ollama       = get_ollama_client()
            model        = st.session_state.get("selected_model", OLLAMA_MODEL)
            output_area  = st.empty()
            full_text    = ""

            # Research step
            research_context = ""
            _was_cached = False
            with st.spinner(f"Researching {female}…"):
                from script_writer import cache_get
                research_context = cache_get(female.strip()) or ""
                if research_context:
                    _was_cached = True
                else:
                    research_context = research_scene_trends(female.strip(), ollama_client=ollama, model=model) or ""

            if research_context:
                _cache_label = "📊 Platform Research (cached)" if _was_cached else "📊 Platform Research"
                with st.expander(_cache_label, expanded=False):
                    st.text(research_context)
            else:
                st.caption(f"ℹ️ No research found for {female} — writing without.")

            # Build banned scenario list from recent plots
            _recent_themes  = st.session_state.get("recent_themes", [])
            _recent_plots   = st.session_state.get("recent_plots", [])
            _banned_keywords = set()
            _SCENARIO_KEYWORDS = {
                "neighbor": ["neighbor", "across the hall", "apartment complex", "moving in", "grocery bags", "boxes scattered"],
                "gym":      ["gym", "personal trainer", "workout", "resistance bands", "exercise equipment"],
                "office":   ["office", "coworker", "boss", "boardroom"],
                "massage":  ["massage therapist", "massage therapy", "spa therapist", "massage table"],
            }
            for _plot in _recent_plots[-4:]:
                _plot_lower = _plot.lower()
                for _concept, _keys in _SCENARIO_KEYWORDS.items():
                    if any(_k in _plot_lower for _k in _keys):
                        _banned_keywords.add(_concept)

            # Permanently overused scenarios — always banned regardless of session history
            _banned_keywords.add("neighbor")
            _banned_keywords.add("gym")
            _banned_keywords.add("massage")
            _banned_keywords.add("art")

            _avoid_hint = "STRICTLY BANNED — do NOT write any version of these scenarios:\n"
            _avoid_hint += f"- BANNED SCENARIO TYPES: {', '.join(sorted(_banned_keywords))}\n"
            _avoid_hint += "  (neighbor/moving-in, gym/personal trainer, massage/spa, art/gallery = ALWAYS banned)\n"
            if _recent_themes:
                _avoid_hint += f"- BANNED THEME TITLES (already used): {', '.join(_recent_themes[-5:])}\n"
            _avoid_hint += "You MUST pick a completely different scenario type — something with a real job, skill, or service relationship.\n"

            # Inject banned list INTO the parsed dict so build_prompt places it
            # BEFORE the final "Start your response with THEME:" instruction.
            existing_hint = parsed.get("theme_hint") or ""
            parsed["theme_hint"] = (_avoid_hint + ("\n" + existing_hint if existing_hint else "")).strip()

            prompt = build_prompt(parsed, research_context=research_context)

            with st.spinner(f"Writing script…"):
                try:
                    import anthropic as _anthropic
                    _claude_key = os.environ.get("ANTHROPIC_API_KEY", "")
                    if _claude_key:
                        # Use Claude Sonnet for writing quality — research already done by Ollama above
                        _ac = _anthropic.Anthropic(api_key=_claude_key)
                        with _ac.messages.stream(
                            model="claude-sonnet-4-6",
                            max_tokens=600,
                            system=SYSTEM_PROMPT,
                            messages=[{"role": "user", "content": prompt}],
                        ) as _stream:
                            for _chunk in _stream.text_stream:
                                full_text += _chunk
                                output_area.markdown(full_text)
                    else:
                        # Fallback to Ollama if no Claude key
                        stream = ollama.chat.completions.create(
                            model=model, max_tokens=800, temperature=0.82,
                            messages=[
                                {"role": "system", "content": SYSTEM_PROMPT},
                                {"role": "user",   "content": prompt},
                            ],
                            stream=True,
                        )
                        for chunk in stream:
                            delta = chunk.choices[0].delta.content or ""
                            full_text += delta
                            output_area.markdown(full_text)
                except Exception as api_err:
                    st.error(f"Generation error: {api_err}")
                    return

            full_text = _post_process(full_text)
            output_area.empty()  # Clear streaming text — formatted result shown below
            fields, violations = _parse_and_validate(full_text, female=female, male=male)

            # Track recent themes + plots to avoid repeats
            _theme = fields.get("theme", "").strip()
            _plot  = fields.get("plot",  "").strip()
            if _theme:
                _rt = st.session_state.get("recent_themes", [])
                _rt.append(_theme)
                st.session_state["recent_themes"] = _rt[-6:]
            if _plot:
                _rp = st.session_state.get("recent_plots", [])
                _rp.append(_plot)
                st.session_state["recent_plots"] = _rp[-6:]

            st.session_state["last_script"] = {
                "full_text": full_text, "fields": fields, "violations": violations,
                "parsed": parsed, "studio": studio, "female": female, "male": male or "",
                "research_context": research_context,
                "ws_title": selected_ws.title if selected_ws else None,
                "row_idx": selected_row_idx, "scene_norm": scene_norm,
            }

        # =========================================================================
        # MANUAL MODE
        # =========================================================================
        if mode == "✏️ Manual":
            with st.container(border=True):
                # ── Studio + Scene Type ───────────────────────────────────────────────
                _sa, _sb = st.columns([3, 1])
                with _sa:
                    studio = st.pills("Studio", ["VRHush", "FuckPassVR", "VRAllure", "NaughtyJOI"],
                                      selection_mode="single", default="VRHush", key="sc_studio")
                studio = studio or "VRHush"

                destination    = None
                scene_type_val = None
                male           = None

                with _sb:
                    if studio == "VRAllure":
                        scene_type_val = st.pills("Scene Type", ["Waking Up", "Pornstar Experience"],
                                                  selection_mode="single", default="Waking Up", key="sc_type_vra")
                        scene_type_val = scene_type_val or "Waking Up"
                    elif studio == "NaughtyJOI":
                        scene_type_val = "JOI"
                        st.caption("JOI — fixed plot")
                    elif studio == "FuckPassVR":
                        destination = st.text_input("Destination", placeholder="e.g. Paris, France",
                                                    key="sc_dest", label_visibility="collapsed")
                    else:
                        _raw = st.pills("Scene Type", ["BG", "BGCP"],
                                        selection_mode="single", default="BG", key="sc_type_bg")
                        _raw = _raw or "BG"
                        scene_type_val = "BGCP" if "BGCP" in _raw else "BG"

                # ── Talent ────────────────────────────────────────────────────────────
                if studio in ("VRAllure", "NaughtyJOI"):
                    female = st.text_input("Female Talent", placeholder="e.g. Lucy Lotus", key="sc_female_solo")
                else:
                    _fc, _mc = st.columns(2)
                    with _fc:
                        female = st.text_input("Female Talent", placeholder="e.g. Lucy Lotus", key="sc_female")
                    with _mc:
                        male = st.text_input("Male Talent", placeholder="e.g. Danny Steele", key="sc_male")

                # ── Director's note + Generate on same row ────────────────────────────
                _default_note = st.session_state.pop("director_note_override", "")
                _nc, _bc = st.columns([5, 1])
                with _nc:
                    theme_hint = st.text_input("Director's Note", value=_default_note,
                                               placeholder="e.g. They meet at a hotel bar…",
                                               key="sc_note")
                with _bc:
                    st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
                    _gen_clicked = st.button("Generate", type="primary",
                                             width="stretch", key="sc_manual_gen")

            if _gen_clicked:
                errors = []
                if not (female or "").strip():
                    errors.append("Female talent is required.")
                if studio not in ("VRAllure", "NaughtyJOI") and not (male or "").strip():
                    errors.append("Male talent is required.")
                for e in errors:
                    st.error(e)
                if not errors:
                    parsed, scene_norm = _build_parsed(
                        studio, female.strip(), (male or "").strip(),
                        scene_type_val, destination, theme_hint.strip() or None,
                    )
                    _run_single_generation(parsed, studio, female.strip(), (male or "").strip(),
                                           scene_norm, None, None, theme_hint)

            _show_single_result()

            # System Check — developer only
            with st.expander("⚙️ Developer Tools", expanded=False):
                if st.button("Run System Check", key="sc_selfcheck"):
                    _sc_model = st.session_state.get("selected_model", OLLAMA_MODEL)
                    try:
                        _base = OLLAMA_BASE_URL.replace("/v1", "")
                        _resp = _requests.get(f"{_base}/api/tags", timeout=4)
                        _avail = [m["name"] for m in _resp.json().get("models", [])]
                        if _sc_model in _avail:
                            st.caption(f"Ollama: **{_sc_model}** loaded")
                        else:
                            st.caption(f"Ollama: **{_sc_model}** not found — available: {', '.join(_avail[:4])}")
                    except Exception as _e:
                        st.caption(f"Ollama unreachable: {_e}")
                    try:
                        _sh2 = get_spreadsheet()
                        _tabs2 = month_tabs(_sh2)
                        st.caption(f"Sheet: connected — {len(_tabs2)} tab(s)")
                    except Exception as _e:
                        st.caption(f"Sheet error: {_e}")

        # =========================================================================
        # FROM SHEET MODE
        # =========================================================================
        else:
            # ── Control bar: month + filter + refresh ─────────────────────────────
            _cc1, _cc2, _cc3, _cc4 = st.columns([2, 4, 1, 1])

            # Load tab list once, cache in session_state
            if "sc_tab_list" not in st.session_state:
                try:
                    sh = get_spreadsheet()
                    _all_tabs = month_tabs(sh)
                    st.session_state["sc_tab_list"]  = _all_tabs
                    st.session_state["sc_tab_names"] = [ws.title for ws in _all_tabs]
                except Exception as _e:
                    st.error(f"Could not connect to sheet: {_e}")

            _all_tabs  = st.session_state.get("sc_tab_list",  [])
            _tab_names = st.session_state.get("sc_tab_names", [])

            with _cc1:
                _chosen_month = st.selectbox("Month", _tab_names, key="sc_month",
                                             label_visibility="collapsed") if _tab_names else None
            with _cc2:
                _ftext = st.text_input("Filter", placeholder="Filter by name, studio, scene…",
                                       key="sc_filter", label_visibility="collapsed")
            with _cc3:
                if st.button("↺", key="sc_refresh_tabs", width="stretch",
                             help="Refresh sheet tabs"):
                    try:
                        sh = get_spreadsheet()
                        _all_tabs = month_tabs(sh)
                        st.session_state["sc_tab_list"]  = _all_tabs
                        st.session_state["sc_tab_names"] = [ws.title for ws in _all_tabs]
                        for _k in list(st.session_state.keys()):
                            if _k.startswith("sc_rows_"):
                                del st.session_state[_k]
                        st.rerun()
                    except Exception as _e:
                        st.error(f"Refresh failed: {_e}")

            with _cc4:
                with st.popover("Regen"):
                    _regen_name = st.text_input("Talent name", placeholder="e.g. Lucy Lotus",
                                                key="sc_regen_talent")
                    if st.button("Mark for Regen", key="sc_regen_btn", width="stretch"):
                        if _regen_name.strip():
                            _marked = mark_talent_for_regen(_regen_name.strip())
                            if _marked:
                                st.success(f"Marked {len(_marked)} row(s).")
                            else:
                                st.warning("No rows found.")
                        else:
                            st.error("Enter a name.")

            if not _chosen_month:
                st.info("No sheet tabs found — check connection.")
            else:
                _ws = next((t for t in _all_tabs if t.title == _chosen_month), None)
                if _ws:
                    # Only fetch rows when the month changes — cache to avoid API quota hits
                    _row_cache_key = f"sc_rows_{_chosen_month}"
                    if _row_cache_key not in st.session_state:
                        with st.spinner("Loading sheet…"):
                            try:
                                st.session_state[_row_cache_key] = _ws.get_all_values()
                            except Exception as _e:
                                st.error(f"Could not load rows: {_e}")
                                st.session_state[_row_cache_key] = []
                    _all_rows = st.session_state[_row_cache_key]

                    # Process + sort once per month (cached), then filter per keystroke
                    _sorted_cache_key = f"sc_sorted_{_chosen_month}"
                    if _sorted_cache_key not in st.session_state:
                        _scene_opts_sorted = []
                        for _i, _row in enumerate(_all_rows[1:], start=2):
                            while len(_row) <= max(COL_STUDIO, COL_FEMALE, COL_SCENE, COL_MALE, COL_LOCATION, COL_DATE, COL_PLOT):
                                _row.append("")
                            _s = _row[COL_STUDIO].strip()
                            _f = _row[COL_FEMALE].strip()
                            if not _s or not _f:
                                continue
                            _m      = _row[COL_MALE].strip()
                            _sc     = _row[COL_SCENE].strip()
                            _loc    = _row[COL_LOCATION].strip()
                            _dt     = _row[COL_DATE].strip()
                            _pltxt  = _row[COL_PLOT].strip()
                            _thtxt  = _row[COL_THEME].strip() if len(_row) > COL_THEME else ""
                            _titxt  = _row[COL_TITLE].strip() if len(_row) > COL_TITLE else ""
                            _hp     = bool(_pltxt)
                            _scene_opts_sorted.append({
                                "Select": False,
                                "✓": "✓" if _hp else "",
                                "Female": _f, "Male": _m, "Studio": _s,
                                "Scene": _sc, "Date": _dt,
                                "_row_idx": _i, "_studio": _s, "_female": _f,
                                "_male": _m, "_scene": _sc, "_loc": _loc, "_hp": _hp,
                                "_plot": _pltxt, "_theme": _thtxt, "_title": _titxt,
                            })
                        from datetime import datetime as _dt_sort2
                        def _date_sort_key(o):
                            d = o.get("Date", "").strip()
                            for fmt in ("%m/%d/%Y", "%m/%d/%y", "%Y-%m-%d", "%d/%m/%Y", "%m-%d-%Y"):
                                try:
                                    return (_dt_sort2.strptime(d, fmt), o["Female"].lower())
                                except ValueError:
                                    continue
                            return (_dt_sort2.max, o["Female"].lower())
                        _scene_opts_sorted.sort(key=_date_sort_key)
                        st.session_state[_sorted_cache_key] = _scene_opts_sorted

                    _scene_opts = st.session_state[_sorted_cache_key]

                    # Apply filter (lightweight — only runs the list comprehension)
                    if _ftext.strip():
                        _q = _ftext.strip().lower()
                        _scene_opts = [o for o in _scene_opts
                                       if _q in f"{o['Female']} {o['Studio']} {o['Scene']} {o['Date']}".lower()]

                    if not _scene_opts:
                        st.info("No scenes found — try a different filter or month.")
                    else:
                        # ── Stats bar + batch toggle ──────────────────────────────
                        _total   = len(_scene_opts)
                        _done    = sum(1 for o in _scene_opts if o["_hp"])
                        _pending = _total - _done
                        _pct     = int(_done / _total * 100) if _total else 0
                        _sb1, _sb2 = st.columns([5, 1])
                        with _sb1:
                            st.markdown(
                                f"<div style='background:{_C['surface']};border-radius:6px;padding:6px 12px;margin:4px 0;"
                                f"display:flex;align-items:center;gap:14px'>"
                                f"<div style='flex:1;height:4px;background:{_C['elevated']};border-radius:2px;overflow:hidden'>"
                                f"<div style='width:{_pct}%;height:100%;background:{_C['green']};border-radius:2px'></div></div>"
                                f"<span style='color:{_C['text']};font-size:0.8rem;white-space:nowrap'>"
                                f"<b>{_done}</b><span style='color:{_C['subtle']}'>/{_total}</span> scripted"
                                f"</span></div>",
                                unsafe_allow_html=True
                            )
                        with _sb2:
                            _batch_mode = st.toggle("Batch", key="sc_batch_toggle", value=False)

                        # Clear selection when filter/month changes
                        _state_sig = f"{_chosen_month}|{_ftext}"
                        if st.session_state.get("sc_state_sig") != _state_sig:
                            st.session_state["sc_state_sig"]   = _state_sig
                            st.session_state["sc_selected_idx"] = None
                            for _k in list(st.session_state.keys()):
                                if _k.startswith("sc_chk_"):
                                    del st.session_state[_k]

                        # ── Scene card list ───────────────────────────────────────
                        _sel_idx = st.session_state.get("sc_selected_idx")

                        for _ci, _opt in enumerate(_scene_opts):
                            _is_sel  = (_sel_idx == _ci)
                            _card_bg = _C["green_dim"] if _opt["_hp"] else _C["amber_dim"]
                            _border  = _C["green"] if _opt["_hp"] else _C["amber"]
                            _sel_bg  = _C["blue_dim"]
                            if _is_sel:
                                _card_bg, _border = _sel_bg, _C["blue"]

                            _male_part  = (f"<span style='color:{_C['muted']};font-size:0.85rem'> · {_opt['Male']}</span>"
                                           if _opt.get("Male") else "")
                            _scene_part = (f"<span style='background:{_C['blue_dim']};border-radius:3px;"
                                           f"padding:1px 7px;font-size:0.78rem;margin-left:6px'>{_opt['Scene']}</span>"
                                           if _opt.get("Scene") else "")
                            _done_badge      = f"&nbsp;&nbsp;<span style='color:{_C['green']};font-size:0.8rem'>✓ scripted</span>" if _opt["_hp"] else ""
                            _done_badge_mini = f"&nbsp;&nbsp;<span style='color:{_C['green']};font-size:0.8rem'>✓</span>" if _opt["_hp"] else ""
                            _studio_pill = (f"<span style='background:{_C['blue_dim']};border-radius:3px;padding:1px 7px;"
                                            f"font-size:0.78rem;margin-left:8px'>{_opt['Studio']}</span>")
                            _date_span   = (f"<span style='color:{_C['muted']};font-size:0.78rem;width:55px;"
                                            f"display:inline-block'>{_opt['Date']}</span>")
                            _name_span   = f"<b style='font-size:0.95rem'>{_opt['Female']}</b>"
                            _card_inner  = f"{_date_span}{_name_span}{_male_part}{_studio_pill}{_scene_part}"
                            # For scripted rows: show theme + title snippet as a second line
                            _theme_line = ""
                            if _opt["_hp"]:
                                _th = _opt.get("_theme", "")
                                _ti = _opt.get("_title", "")
                                _parts2 = []
                                if _th: _parts2.append(f"<span style='color:{_C['green']}'>{_th}</span>")
                                if _ti: _parts2.append(f"<span style='color:{_C['muted']};font-style:italic'>{_ti}</span>")
                                if _parts2:
                                    _theme_line = f"<div style='font-size:0.72rem;margin-top:3px;padding-left:55px'>{'&nbsp;·&nbsp;'.join(_parts2)}</div>"

                            if _batch_mode:
                                _cb_col, _card_col = st.columns([1, 14])
                                with _cb_col:
                                    st.checkbox("", key=f"sc_chk_{_ci}", label_visibility="collapsed")
                                with _card_col:
                                    st.markdown(
                                        f"<div style='background:{_card_bg};border-left:3px solid {_border};"
                                        f"border-radius:5px;padding:7px 12px;margin:1px 0'>"
                                        f"{_card_inner}{_done_badge}{_theme_line}</div>",
                                        unsafe_allow_html=True
                                    )
                            else:
                                _card_col, _btn_col = st.columns([8, 3])
                                with _card_col:
                                    st.markdown(
                                        f"<div style='background:{_card_bg};border-left:3px solid {_border};"
                                        f"border-radius:5px;padding:7px 12px;margin:1px 0'>"
                                        f"{_card_inner}{_done_badge_mini}{_theme_line}</div>",
                                        unsafe_allow_html=True
                                    )
                                with _btn_col:
                                    _blabel = "▼ Selected" if _is_sel else ("View / Edit" if _opt["_hp"] else "Generate →")
                                    if st.button(_blabel, key=f"sc_sel_{_ci}",
                                                 width="stretch",
                                                 type="primary" if _is_sel else "secondary"):
                                        if _is_sel:
                                            st.session_state["sc_selected_idx"] = None
                                        else:
                                            st.session_state["sc_selected_idx"] = _ci
                                        st.rerun()

                        # ── Action panel ──────────────────────────────────────────
                        if _batch_mode:
                            _sel_rows = [_scene_opts[_ci] for _ci in range(len(_scene_opts))
                                         if st.session_state.get(f"sc_chk_{_ci}", False)]
                            _n = len(_sel_rows)
                            _dry = st.checkbox("Dry run (preview only, don't write to sheet)", value=False, key="sc_dry")
                        else:
                            _sel_idx = st.session_state.get("sc_selected_idx")
                            if _sel_idx is not None and _sel_idx < len(_scene_opts):
                                _ch = _scene_opts[_sel_idx]
                                st.markdown("<div style='margin-top:12px'></div>", unsafe_allow_html=True)

                                # ── If scripted and no fresh result yet, show existing content ──
                                if _ch["_hp"] and "last_script" not in st.session_state:
                                    _ex_plot  = _ch.get("_plot", "")
                                    _ex_theme = _ch.get("_theme", "")
                                    _ex_title = _ch.get("_title", "")
                                    if _ex_theme:
                                        st.caption(f"**Theme:** {_ex_theme}" + (f"  ·  **Title:** {_ex_title}" if _ex_title else ""))
                                    _edited_existing = st.text_area(
                                        "Existing script — edit and save, or regenerate below",
                                        value=_ex_plot, height=160,
                                        key=f"sc_existing_{_sel_idx}",
                                        label_visibility="visible",
                                    )
                                    _sa1, _sa2 = st.columns(2)
                                    with _sa1:
                                        if st.button("💾 Save edits", key="sc_save_existing", width="stretch", type="primary"):
                                            try:
                                                _ws.update_cell(_ch["_row_idx"], COL_PLOT + 1, _edited_existing)
                                                st.success("Saved to sheet")
                                                del st.session_state[f"sc_rows_{_chosen_month}"]
                                                st.rerun()
                                            except Exception as _se:
                                                st.error(f"Save failed: {_se}")
                                    with _sa2:
                                        st.markdown("")  # spacer
                                    st.divider()

                                _default_note = st.session_state.pop("director_note_override", "")
                                _note = st.text_area("Director's note (optional)", value=_default_note,
                                                     placeholder="e.g. They meet at a hotel bar…",
                                                     height=70, key="sc_sheet_note")
                                _gen_label = "Regenerate Script" if _ch["_hp"] else "Generate Script"
                                if st.button(_gen_label, type="primary" if not _ch["_hp"] else "secondary",
                                             width="stretch", key="sc_sheet_single"):
                                    _parsed, _snorm = _build_parsed(
                                        _ch["_studio"], _ch["_female"], _ch["_male"],
                                        _ch["_scene"], _ch["_loc"], _note.strip() or None,
                                    )
                                    _run_single_generation(
                                        _parsed, _ch["_studio"], _ch["_female"], _ch["_male"],
                                        _snorm, _ws, _ch["_row_idx"], _note,
                                    )
                                _show_single_result()
                            else:
                                _show_single_result()
                            _sel_rows = []
                            _n = 0

                        # ── Batch generate ────────────────────────────────────────
                        if _batch_mode:
                            _dry = st.session_state.get("sc_dry", False)
                            if st.button(f"Generate {_n} Scripts", type="primary",
                                         width="stretch", key="sc_batch_gen",
                                         disabled=(_n == 0)):
                                _ollama  = get_ollama_client()
                                _bmodel  = st.session_state.get("selected_model", OLLAMA_MODEL)
                                _prog    = st.progress(0)
                                _results = []

                                for _bi, _row in enumerate(_sel_rows):
                                    _bstudio = _row["_studio"]
                                    _bfemale = _row["_female"]
                                    _bmale   = _row["_male"]
                                    _bloc    = _row["_loc"]
                                    _bscene  = _row["_scene"]
                                    _blow    = _bstudio.lower()

                                    if _dry:
                                        _prog.progress((_bi + 1) / _n)
                                        _results.append({"label": f"{_bfemale} ({_bstudio})", "dry_run": True})
                                        continue

                                    if _blow in ("naughtyjoi", "njoi"):
                                        write_script(_ws, _row["_row_idx"], theme="JOI",
                                                     plot=NJOI_STATIC_PLOT, wardrobe_female="", wardrobe_male="")
                                        _results.append({"label": f"{_bfemale} (NaughtyJOI)",
                                                         "auto_saved": True, "ws": _ws, "row_idx": _row["_row_idx"]})
                                        _prog.progress((_bi + 1) / _n)
                                        continue

                                    if _blow in ("vrallure", "vra"):
                                        _bvra = "Pornstar Experience" if "pornstar" in _bscene.lower() else \
                                                "Waking Up" if "waking" in _bscene.lower() else None
                                        _bparsed = {"studio": "VRAllure", "female": _bfemale, "vra_scene_type": _bvra}
                                        _bsnorm  = _bvra or "VRAllure"
                                    else:
                                        _bsnorm  = "BGCP" if "CP" in _bscene.upper() or "CREAMPIE" in _bscene.upper() else "BG"
                                        _bparsed = {"studio": _bstudio, "destination": _bloc or None,
                                                    "scene_type": _bsnorm, "female": _bfemale, "male": _bmale}

                                    with st.spinner(f"Writing: {_bfemale}…"):
                                        from script_writer import cache_get
                                        _brc = cache_get(_bfemale.strip()) or ""
                                        _btext, _bfields, _bviols = _generate_script_core(
                                            _bparsed, _bfemale, _bmale, research_context=_brc
                                        )

                                    _results.append({
                                        "label": f"{_bfemale} — {_bstudio}",
                                        "ws": _ws, "row_idx": _row["_row_idx"],
                                        "studio": _bstudio, "scene": _bscene,
                                        "parsed": _bparsed, "fields": _bfields,
                                        "full_text": _btext, "research": "",
                                        "violations": _bviols,
                                    })
                                    _prog.progress((_bi + 1) / _n)

                                st.session_state["batch_results"]   = _results
                                st.session_state["batch_decisions"] = {}
                                st.rerun()

                            # ── Batch review queue ────────────────────────────────
                            if st.session_state.get("batch_results"):
                                _bresults   = st.session_state["batch_results"]
                                _bdecisions = st.session_state.setdefault("batch_decisions", {})

                                # Find the single next pending script
                                _next_idx = next(
                                    (j for j, r in enumerate(_bresults)
                                     if not r.get("auto_saved") and not r.get("dry_run")
                                     and j not in _bdecisions), None)
                                _total_reviewable = sum(1 for r in _bresults
                                                        if not r.get("auto_saved") and not r.get("dry_run"))
                                _reviewed = sum(1 for i, r in enumerate(_bresults)
                                                if not r.get("auto_saved") and not r.get("dry_run")
                                                and i in _bdecisions)

                                st.divider()

                                # Progress strip (CSS dots)
                                _dot = lambda c: f"<span style='display:inline-block;width:10px;height:10px;border-radius:50%;background:{c};margin:0 2px'></span>"
                                _prog_dots = ""
                                for _pi, _pr in enumerate(_bresults):
                                    if _pr.get("auto_saved"):  _prog_dots += _dot(_C["green"])
                                    elif _pr.get("dry_run"):   _prog_dots += _dot(_C["subtle"])
                                    elif _pi in _bdecisions:
                                        _prog_dots += _dot(_C["green"] if _bdecisions[_pi] == "accepted" else _C["red"])
                                    elif _pi == _next_idx:     _prog_dots += _dot(_C["blue"])
                                    else:                      _prog_dots += _dot(_C["elevated"])
                                st.markdown(
                                    f"<div style='display:flex;align-items:center;gap:8px;margin:0 0 8px'>"
                                    f"<div>{_prog_dots}</div>"
                                    f"<span style='font-size:0.8rem;color:{_C['muted']}'>"
                                    f"<b style='color:{_C['text']}'>{_reviewed}</b>/{_total_reviewable} reviewed</span></div>",
                                    unsafe_allow_html=True)

                                if _next_idx is None:
                                    st.success("All scripts reviewed.", icon="✅")
                                    if st.button("Clear results", key="sc_batch_clear"):
                                        del st.session_state["batch_results"]
                                        del st.session_state["batch_decisions"]
                                        st.rerun()
                                else:
                                    _bi2  = _next_idx
                                    _bres = _bresults[_bi2]
                                    _bf   = _bres["fields"]
                                    _bv   = _bres["violations"]

                                    # Compact script card
                                    st.markdown(
                                        f"<div class='hub-card hub-card-accent' style='border-left-color:"
                                        f"{_C['amber'] if _bv else _C['green']}'>"
                                        f"<span style='font-size:1.05rem;font-weight:700'>{_bres['label']}</span>"
                                        f"<span style='color:{_C['muted']};margin-left:10px;font-size:0.85rem'>"
                                        f"{_bf.get('theme','')}</span><br>"
                                        f"<span style='color:{_C['muted']};font-size:0.82rem'>"
                                        f"👗 {_bf.get('wardrobe_female','—')}</span>"
                                        + (f"&emsp;<span style='color:{_C['muted']};font-size:0.82rem'>"
                                           f"👔 {_bf.get('wardrobe_male','')}</span>"
                                           if _bf.get('wardrobe_male') else "")
                                        + f"</div>",
                                        unsafe_allow_html=True)

                                    if _bv:
                                        st.warning(" · ".join(_bv), icon="⚠️")

                                    st.text_area("Plot", value=_bf.get("plot", ""),
                                                 height=90, key=f"sc_plot_{_bi2}", disabled=True)

                                    _bca, _bcb = st.columns(2)
                                    with _bca:
                                        _b_acc_label = "Accept & Save" if _user_is_admin else "Submit for Approval"
                                        if st.button(_b_acc_label, key=f"sc_acc_{_bi2}",
                                                     width="stretch", type="primary"):
                                            if _user_is_admin:
                                                try:
                                                    _do_save(_bres["ws"], _bres["row_idx"], _bf)
                                                    try:
                                                        from training_data import save_accepted
                                                        save_accepted(_bres["parsed"]["studio"],
                                                                      _bres["parsed"], _bf, _bres["research"])
                                                    except Exception as _te:
                                                        st.warning(f"Training data not saved: {_te}")
                                                    _bdecisions[_bi2] = "accepted"
                                                    st.session_state["batch_decisions"] = _bdecisions
                                                    st.rerun()
                                                except Exception as _e:
                                                    st.error(f"Save failed: {_e}")
                                            else:
                                                try:
                                                    import json as _json_mod
                                                    _b_studio = _bres.get("studio", _bres["parsed"].get("studio", ""))
                                                    _b_scene = _bres.get("scene", "")
                                                    _b_sid = f"{_b_studio}{_b_scene}" if _b_scene else _b_studio
                                                    _b_target = f"Scripts:{_bres['ws'].title}:{_bres['row_idx']}"
                                                    _b_preview = f"Theme: {_bf.get('theme','')}\nPlot: {_bf.get('plot','')[:150]}"
                                                    import approval_tools as _apr_b
                                                    _b_apr_id = _apr_b.submit_for_approval(
                                                        submitted_by=_user_name,
                                                        content_type="script",
                                                        scene_id=_b_sid,
                                                        studio=_b_studio,
                                                        content_preview=_b_preview,
                                                        content_json=_json_mod.dumps(_bf),
                                                        target_sheet=_b_target,
                                                    )
                                                    try:
                                                        notification_tools.notify_approval_submitted(
                                                            _b_apr_id, _b_sid, "script", _user_name)
                                                    except Exception:
                                                        pass
                                                    _bdecisions[_bi2] = "accepted"
                                                    st.session_state["batch_decisions"] = _bdecisions
                                                    st.rerun()
                                                except Exception as _e:
                                                    st.error(f"Approval failed: {_e}")
                                    with _bcb:
                                        if st.button("👎 Skip", key=f"sc_skip_{_bi2}",
                                                     width="stretch"):
                                            try:
                                                from training_data import save_rejected
                                                save_rejected(_bres["parsed"]["studio"],
                                                              _bres["parsed"], _bf)
                                            except Exception:
                                                pass
                                            _bdecisions[_bi2] = "rejected"
                                            st.session_state["batch_decisions"] = _bdecisions
                                            st.rerun()

    # ── TAB 2 (was TAB 3): Call Sheets

    # ── TAB 3: Call Sheets ────────────────────────────────────────────────────────
with tab_callsheet:
    if _has_tab("Call Sheets"):
        hub_ui.section("Call Sheets")
        if not HAS_CALL_SHEET:
            st.error("call_sheet.py not found — check deployment.")
        else:
            _cs1, _cs2, _cs3 = st.columns([3, 2, 1])
            with _cs1:
                if "cs_budget_tabs" not in st.session_state:
                    try:
                        st.session_state["cs_budget_tabs"] = get_budget_tabs()
                    except Exception as e:
                        st.session_state["cs_budget_tabs"] = []
                        st.error(f"Could not read Shoot Budgets sheet: {e}")
                budget_tabs = st.session_state["cs_budget_tabs"]
                selected_month = st.selectbox("Month", budget_tabs, key="cs_month",
                                              label_visibility="collapsed") if budget_tabs else None
            with _cs2:
                _cs_gen_all = st.button(f"Generate All for {selected_month or '...'}", key="cs_all",
                                        type="primary", width="stretch") if selected_month else False
            with _cs3:
                door_code = st.text_input("🔑 Door Code", value="1322", key="cs_door")

            if selected_month:
                _sd_cache_key = f"cs_shoot_dates_{selected_month}"
                if _sd_cache_key not in st.session_state:
                    with st.spinner("Loading shoot dates..."):
                        try:
                            st.session_state[_sd_cache_key] = get_shoot_dates(tab_name=selected_month)
                        except Exception as e:
                            st.session_state[_sd_cache_key] = {}
                            st.error(f"Could not read shoot dates: {e}")
                shoot_dates = st.session_state[_sd_cache_key]

                if not shoot_dates:
                    st.info("No shoots found for this month.")
                else:
                    st.caption(f"{len(shoot_dates)} shoot date(s) found")

                    if _cs_gen_all:
                        results, errors = [], []
                        reload_script_cache()
                        prog = st.progress(0, text="Generating...")
                        date_keys = sorted(shoot_dates.keys())
                        for i, date_key in enumerate(date_keys):
                            scenes = shoot_dates[date_key]
                            prog.progress((i) / len(date_keys), text=f"Generating {date_key}...")
                            try:
                                result = generate_call_sheet(date_key, scenes, door_code=door_code)
                                results.append(result)
                            except Exception as e:
                                errors.append((date_key, str(e)))
                        prog.progress(1.0, text="Done!")
                        for r in results:
                            st.success(f"✅ [{r['title']}]({r['doc_url']})")
                        for date_key, err in errors:
                            st.error(f"❌ {date_key}: {err}")

                    st.divider()
                    for date_key in sorted(shoot_dates.keys()):
                        scenes = shoot_dates[date_key]
                        dt = scenes[0]["date_dt"]
                        date_label = dt.strftime("%B {day}, %Y").replace("{day}", str(dt.day))
                        females = list(dict.fromkeys(s["female"] for s in scenes if s["female"]))
                        studios = list(dict.fromkeys(s["studio"] for s in scenes if s["studio"]))
                        types   = list(dict.fromkeys(s["type"]   for s in scenes if s["type"]))

                        with st.expander(f"**{date_label}** — {', '.join(females)} | {', '.join(studios)} | {', '.join(types)}"):
                            # Scene table
                            scene_rows = [{"Studio": s["studio"], "Type": s["type"], "Female": s["female"], "Male": s["male"], "Agency": s["agency"]} for s in scenes]
                            import pandas as _pd
                            st.dataframe(_pd.DataFrame(scene_rows), width="stretch", hide_index=True)

                            if st.button(f"Generate Call Sheet", key=f"cs_{date_key}", type="primary"):
                                with st.spinner("Creating Google Doc..."):
                                    try:
                                        reload_script_cache()
                                        result = generate_call_sheet(date_key, scenes, door_code=door_code)
                                        st.success("Call sheet created!")
                                        st.markdown(f"**[Open in Google Docs]({result['doc_url']})**")
                                        st.caption(result["title"])
                                    except Exception as e:
                                        st.error(f"Failed: {e}")
                                        import traceback
                                        st.code(traceback.format_exc())


    # ── TAB 4: Titles ─────────────────────────────────────────────────────────────
with tab_titles:
    if _has_tab("Titles"):
        hub_ui.section("Title Card Generator")
        try:
            import cta_generator as _cta
            HAS_CTA = True
        except ImportError as _cta_err:
            HAS_CTA = False
            st.error(f"cta_generator.py not found — copy it to the project folder. ({_cta_err})")

        if HAS_CTA:
            with st.container(border=True):
                # ── Engine + title on one row ─────────────────────────────────────────
                _te1, _te2 = st.columns([1, 3])
                with _te1:
                    _engine = st.segmented_control("Engine", ["☁️ Cloud", "🖥️ Local"],
                                                   default="☁️ Cloud", key="ti_engine",
                                                   label_visibility="collapsed")
                    _engine = _engine or "☁️ Cloud"
                _use_cloud = _engine == "☁️ Cloud"
                with _te2:
                    _ti_title = st.text_input(
                        "Title text",
                        placeholder="e.g. Forbidden Fantasy, Neon Nights, Gold Rush",
                        key="ti_title",
                    )
                if _use_cloud:
                    st.caption("☁️ **Cloud** — Ideogram V3 via fal.ai · photorealistic AI text · ~$0.03/image · 5–15s")
                else:
                    st.caption("🖥️ **Local** — PIL graphic treatments · instant & free · 690+ styles")

            if _use_cloud:
                # Cloud mode — style picker
                try:
                    import cloud_renderer as _cr
                    _cloud_styles = list(_cr.CLOUD_STYLES.keys())
                except Exception:
                    _cloud_styles = []

                with st.container(border=True):
                    col_style, col_n = st.columns([3, 1])
                    with col_style:
                        _cloud_style_mode = st.radio(
                            "Style",
                            ["Random mix", "Pick one"],
                            horizontal=True,
                            key="ti_cloud_style_mode",
                            label_visibility="collapsed",
                        )
                    with col_n:
                        _ti_n = st.number_input("Variations", min_value=1, max_value=20, value=6, key="ti_n_cloud")

                    if _cloud_style_mode == "Pick one" and _cloud_styles:
                        _picked_style = st.selectbox("Style", _cloud_styles, key="ti_cloud_pick")
                    else:
                        _picked_style = None

                    _ti_gen_cloud_btn = st.button("Generate Title PNGs", type="primary", width="stretch", key="ti_gen_cloud")

                if _ti_gen_cloud_btn:
                    if not _ti_title.strip():
                        st.error("Enter a title first.")
                    elif not _cloud_styles:
                        st.error("cloud_renderer.py not found or FAL_KEY not set.")
                    else:
                        import random
                        _n = int(_ti_n)
                        _base_seed = random.randint(1, 99999)
                        if _cloud_style_mode == "Random mix":
                            # Sample styles, wrap around if requesting more than available
                            _shuffled = _cloud_styles.copy()
                            random.shuffle(_shuffled)
                            _style_pool = [_shuffled[i % len(_shuffled)] for i in range(_n)]
                        else:
                            _style_pool = [_picked_style] * _n

                        _imgs_out = []
                        _prog = st.progress(0, text=f"Rendering {_n} styles in parallel…")
                        import concurrent.futures
                        _done_count = 0

                        def _render_one_cloud(_ci_sk):
                            _ci, _sk = _ci_sk
                            return (_ci, _sk, _cr.render_cloud(
                                _ti_title.strip(), _sk,
                                seed=_base_seed + _ci * 137
                            ))

                        with concurrent.futures.ThreadPoolExecutor(max_workers=min(_n, 4)) as _ex:
                            _futs = {_ex.submit(_render_one_cloud, (_ci, _sk)): _sk
                                     for _ci, _sk in enumerate(_style_pool)}
                            for _fut in concurrent.futures.as_completed(_futs):
                                _done_count += 1
                                try:
                                    _ci, _sk, _png = _fut.result()
                                    if _png:
                                        _imgs_out.append((_sk, _png))
                                except Exception as _ce:
                                    st.warning(f"{_futs[_fut]}: {_ce}")
                                _prog.progress(_done_count / _n, text=f"Done {_done_count}/{_n}…")
                        _prog.progress(1.0, text=f"Done! {len(_imgs_out)}/{_n} rendered.")

                        st.session_state["ti_imgs"]  = _imgs_out
                        st.session_state["ti_label"] = _ti_title.strip()
                        st.session_state["ti_seed_used"] = 0

            else:
                # Local PIL mode
                with st.container(border=True):
                    col_mode, col_n, col_seed = st.columns([3, 1, 1])
                    with col_mode:
                        _ti_mode = st.radio(
                            "Mode",
                            ["Random mix", "Auto-match keywords", "Pick one"],
                            horizontal=True,
                            key="ti_mode",
                            label_visibility="collapsed",
                        )
                    with col_n:
                        _ti_n = st.number_input("Variations", min_value=1, max_value=12, value=6, key="ti_n")
                    with col_seed:
                        _ti_seed = st.number_input(
                            "Seed", min_value=0, max_value=999999, value=0, key="ti_seed",
                            help="0 = new random seed each time",
                        )

                    if _ti_mode == "Pick one":
                        _all_treatments = sorted(_cta.TREATMENTS.keys())
                        _feat_treatments = sorted(getattr(_cta, "FEATURED_TREATMENTS", _cta.TREATMENTS).keys())
                        _tp1, _tp2 = st.columns([3, 1])
                        with _tp1:
                            _ti_filt = st.text_input("Filter treatments", placeholder="e.g. gold, neon, chrome…",
                                                      key="ti_filt", label_visibility="collapsed")
                        with _tp2:
                            _feat_only = st.toggle("Featured only", value=True, key="ti_feat_only")
                        _pool_to_show = _feat_treatments if _feat_only else _all_treatments
                        if _ti_filt.strip():
                            _pool_to_show = [t for t in _pool_to_show if _ti_filt.lower() in t.lower()]
                        if _pool_to_show:
                            _ti_treatment = st.selectbox(
                                f"Treatment — {len(_pool_to_show)} shown"
                                + (f" of {len(_feat_treatments)} featured" if _feat_only else f" of {len(_all_treatments)} total"),
                                _pool_to_show, key="ti_treatment",
                            )
                            if _feat_only and not _ti_filt.strip():
                                st.caption(f"{len(_feat_treatments)} featured treatments · toggle off to browse all {len(_all_treatments)}")
                        else:
                            st.caption(f"No treatments match '{_ti_filt}' — try a shorter keyword")
                            _ti_treatment = _all_treatments[0] if _all_treatments else None
                    else:
                        _ti_treatment = None

                    _ti_gen_local_btn = st.button("Generate Title PNGs", type="primary", width="stretch", key="ti_gen")

                if _ti_gen_local_btn:
                    if not _ti_title.strip():
                        st.error("Enter a title first.")
                    else:
                        import random, io

                        _seed = _ti_seed if _ti_seed > 0 else random.randint(1, 999999)
                        _n    = int(_ti_n)
                        _feat = list(getattr(_cta, "FEATURED_TREATMENTS", _cta.TREATMENTS).keys())
                        _keys = list(_cta.TREATMENTS.keys())

                        if _ti_mode == "Random mix":
                            _rng0 = random.Random(_seed)
                            _pool = [_rng0.choice(_feat) for _ in range(_n)]
                        elif _ti_mode == "Auto-match keywords":
                            _kw = _ti_title.lower()
                            _matched = next(
                                (v for k, v in _cta.KEYWORD_TREATMENT.items() if k in _kw),
                                None,
                            )
                            if _matched and _matched in _cta.TREATMENTS:
                                _pool = [_matched] * _n
                                st.caption(f"Matched keyword to treatment: {_matched}")
                            else:
                                _rng0 = random.Random(_seed)
                                _pool = [_rng0.choice(_keys) for _ in range(_n)]
                                st.caption("No keyword match — using random treatments")
                        else:
                            _pool = [_ti_treatment] * _n

                        _imgs_out = []
                        with st.spinner("Rendering..."):
                            for _i, _treatment in enumerate(_pool):
                                try:
                                    _rng_i = random.Random(_seed + _i * 1000)
                                    _img   = _cta.TREATMENTS[_treatment](_ti_title.strip(), _rng_i)
                                    # Sharpen for crisp output
                                    from PIL import ImageFilter as _IF
                                    _img = _img.filter(_IF.UnsharpMask(radius=1.5, percent=60, threshold=2))
                                    _buf   = io.BytesIO()
                                    _img.save(_buf, format="PNG")
                                    _imgs_out.append((_treatment, _buf.getvalue()))
                                except Exception as _ex:
                                    st.warning(f"Treatment '{_treatment}' failed: {_ex}")

                        st.session_state["ti_imgs"]  = _imgs_out
                        st.session_state["ti_label"] = _ti_title.strip()
                        st.session_state["ti_seed_used"] = _seed

            if st.session_state.get("ti_imgs"):
                _label = st.session_state.get("ti_label", "title")
                _ti_count = len(st.session_state["ti_imgs"])
                st.divider()
                _tm1, _tm2, _tm3 = st.columns(3)
                _tm1.metric("Title", _label)
                _tm2.metric("Variations", _ti_count)
                _tm3.metric("Seed", st.session_state.get("ti_seed_used", "random"))

                _cols = st.columns(3)
                for _i, (_treatment, _png) in enumerate(st.session_state["ti_imgs"]):
                    with _cols[_i % 3]:
                        with st.container(border=True):
                            # Show on checkerboard so transparency is visible
                            try:
                                from PIL import Image as _PILImg
                                _ti_img = _PILImg.open(io.BytesIO(_png)).convert("RGBA")
                                _tw, _th = _ti_img.size
                                _checker = _checkerboard_bg(_tw, _th).copy()
                                _checker = _PILImg.alpha_composite(_checker, _ti_img)
                                _prev_buf = io.BytesIO()
                                _checker.save(_prev_buf, format="PNG")
                                st.image(_prev_buf.getvalue(), width="stretch")
                            except Exception:
                                st.image(_png, width="stretch")
                            st.markdown(
                                f"<span style='font-size:0.72rem;font-weight:600;color:{_C['muted']}'>{_treatment}</span>",
                                unsafe_allow_html=True)
                            st.download_button(
                                f"Download",
                                data=_png,
                                file_name=f"{_label.replace(' ', '_')}_{_treatment}.png",
                                mime="image/png",
                                key=f"ti_dl_{_i}",
                                width="stretch",
                            )
                            _ref_prompt = st.text_input(
                                "Refine",
                                placeholder="e.g. gold, darker, glow…",
                                key=f"ti_ref_{_i}",
                                label_visibility="collapsed",
                            )
                            if st.button("Apply", key=f"ti_apply_{_i}", width="stretch"):
                                if _ref_prompt.strip():
                                    with st.spinner("Applying..."):
                                        _seed_used = st.session_state.get("ti_seed_used", 42)
                                        _new_png, _new_name = _refine_treatment(
                                            _treatment, _png, _ref_prompt.strip(),
                                            st.session_state.get("ti_label", "title"),
                                            _seed_used + _i * 1000
                                        )
                                    if _new_png:
                                        _imgs_list = list(st.session_state["ti_imgs"])
                                        _imgs_list[_i] = (_new_name, _new_png)
                                        st.session_state["ti_imgs"] = _imgs_list
                                        st.rerun()
                                    else:
                                        st.warning("Could not apply — try simpler keywords.")

        # ── Model Name Generator ─────────────────────────────────────────────────
        st.divider()
        hub_ui.section("Model Name Generator")
        with st.container(border=True):
            _mn_c1, _mn_c2, _mn_c3 = st.columns([2, 4, 2])
            with _mn_c1:
                _mn_studio = st.selectbox("Studio", ["VRA", "VRH"], key="mn_studio")
            with _mn_c2:
                _mn_name = st.text_input("Model name", placeholder="e.g. Emma Rosie", key="mn_name")
            with _mn_c3:
                st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
                _mn_go = st.button("Generate", type="primary", width="stretch", key="mn_gen")

        if _mn_go:
            if not _mn_name.strip():
                st.error("Enter a model name.")
            else:
                with st.spinner("Rendering model name…"):
                    try:
                        _mn_png = _cta.generate_model_name_png(_mn_name.strip(), _mn_studio)
                        st.session_state["mn_png"] = _mn_png
                        st.session_state["mn_label"] = _mn_name.strip()
                        st.session_state["mn_studio_used"] = _mn_studio
                    except Exception as _mn_ex:
                        st.error(f"Render failed: {_mn_ex}")

        if st.session_state.get("mn_png"):
            _mn_label = st.session_state.get("mn_label", "model")
            _mn_studio_u = st.session_state.get("mn_studio_used", "VRH")
            _mn_png_data = st.session_state["mn_png"]
            st.image(_mn_png_data, width="stretch")

            _mn_safe = _mn_label.replace(" ", "")
            st.download_button(
                f"Download {_mn_studio_u} — {_mn_label}",
                data=_mn_png_data,
                file_name=f"{_mn_studio_u}-{_mn_safe}.png",
                mime="image/png",
                key="mn_dl",
                width="stretch",
            )

    # ── TAB 5: Model Research ─────────────────────────────────────────────────────
with tab_research:
    if _has_tab("Model Research"):
        hub_ui.section("Model Research")
        try:
            from model_research_tab import lookup_model_profile
            _HAS_RESEARCH = True
        except ImportError as _re_err:
            _HAS_RESEARCH = False
            st.error(f"model_research_tab.py not found. ({_re_err})")

        if _HAS_RESEARCH:
            from model_research_tab import fetch_trending_models, fetch_photo_bytes, fetch_performer_photo_url
            try:
                from booking_history import get_booking_history, compute_opportunity_score, get_competitor_scenes
                _HAS_BK_HIST = True
            except ImportError:
                _HAS_BK_HIST = False

            # ── Load trending (disk-cached 6h) ────────────────────────────────────
            if "mr_trending" not in st.session_state:
                with st.spinner("Loading trending models…"):
                    st.session_state["mr_trending"] = fetch_trending_models(10)
            _trending = st.session_state.get("mr_trending", [])

            # ── Fetch portrait headshots for trending models (separate from listing thumbnails)
            if "mr_trending_photos_v1" not in st.session_state and _trending:
                with st.spinner("Loading trending headshots…"):
                    import concurrent.futures as _cf
                    _tnames = [_tm["name"] for _tm in _trending]
                    with _cf.ThreadPoolExecutor(max_workers=5) as _ex:
                        _tfuts = {n: _ex.submit(fetch_performer_photo_url, n) for n in _tnames}
                        _tphotos = {}
                        for n, fut in _tfuts.items():
                            try:
                                _tphotos[n] = fut.result(timeout=10)
                            except Exception:
                                _tphotos[n] = None
                st.session_state["mr_trending_photos_v1"] = _tphotos
            _trending_photos = st.session_state.get("mr_trending_photos_v1", {})

            # ── Priority Outreach list ─────────────────────────────────────────────
            _PRIORITY = [
                {"name": "Leah Gotti",       "agency": "Invision Models",   "mo": None},
                {"name": "Alex Blake",        "agency": "Hussie Models",     "mo": None},
                {"name": "Melissa Stratton",  "agency": "Hussie Models",     "mo": None},
                {"name": "Kenzie Reeves",     "agency": "East Coast Talent", "mo": None},
                {"name": "Kali Roses",        "agency": "The Model Service", "mo": None},
                {"name": "Haley Reed",        "agency": "ATMLA",             "mo": None},
                {"name": "Karma RX",          "agency": "ATMLA",             "mo": None},
                {"name": "Cory Chase",        "agency": "ATMLA",             "mo": None},
                {"name": "Valentina Nappi",   "agency": "Speigler",          "mo": None},
                {"name": "Karlee Grey",       "agency": "ATMLA",             "mo": 104},
            ]

            # Pre-fetch priority photos server-side once per session (v3: VRPorn/SLR, no Babepedia)
            if "mr_priority_photos_v3" not in st.session_state:
                with st.spinner("Loading priority photos…"):
                    import concurrent.futures as _cf
                    with _cf.ThreadPoolExecutor(max_workers=5) as _ex:
                        _futs = {_pm["name"]: _ex.submit(fetch_performer_photo_url, _pm["name"]) for _pm in _PRIORITY}
                        _pp = {}
                        for name, fut in _futs.items():
                            try:
                                _pp[name] = fut.result(timeout=10)
                            except Exception:
                                _pp[name] = None
                st.session_state["mr_priority_photos_v3"] = _pp
            _priority_photos = st.session_state["mr_priority_photos_v3"]

            # Handle card click BEFORE widget renders (can't set widget key after instantiation)
            _mr_auto_name = None
            if "mr_priority_load" in st.session_state:
                _mr_auto_name = st.session_state.pop("mr_priority_load")
                st.session_state["mr_name"] = _mr_auto_name  # set before widget renders

            # ── Search bar — compact two-column ───────────────────────────────────
            _col_search, _col_btn = st.columns([7, 1])
            with _col_search:
                _mr_name = st.text_input(
                    "Performer name",
                    placeholder="Search any performer…",
                    key="mr_name",
                    label_visibility="collapsed",
                )
            with _col_btn:
                _mr_search = st.button("Search", type="primary", width="stretch", key="mr_search")
            _mr_refresh = False  # exposed inside profile via refresh button

            _should_search = (_mr_search and _mr_name.strip()) or bool(_mr_auto_name)
            _search_name   = (_mr_auto_name or _mr_name.strip()) if _mr_auto_name else _mr_name.strip()

            if _should_search and _search_name:
                st.session_state["mr_profile"] = None
                st.session_state["mr_query"]   = _search_name
                with st.spinner(f"Fetching data for {_search_name}…"):
                    _profile = lookup_model_profile(_search_name, force_refresh=False)
                st.session_state["mr_profile"] = _profile

            _profile = st.session_state.get("mr_profile")

            # ── Helper: render one model card ──────────────────────────────────────
            import base64 as _b64

            def _model_card(col, name, photo_src, stat_line, btn_key, score=None):
                """Card: photo fills the card, name+stat as gradient overlay, minimal View button."""
                _img_style = ("width:100%;height:195px;object-fit:cover;"
                              "object-position:50% 15%;display:block;border-radius:8px 8px 0 0")
                if isinstance(photo_src, bytes) and photo_src:
                    _b64_str = _b64.b64encode(photo_src).decode()
                    _media = f"<img src='data:image/jpeg;base64,{_b64_str}' style='{_img_style}'>"
                elif isinstance(photo_src, str) and photo_src:
                    _media = f"<img src='{photo_src}' style='{_img_style}'>"
                else:
                    _initials = "".join(w[0].upper() for w in name.split()[:2])
                    _media = (
                        f"<div style='height:195px;background:{_C['elevated']};display:flex;"
                        f"align-items:center;justify-content:center;font-size:1.8rem;"
                        f"font-weight:700;color:{_C['subtle']};border-radius:8px 8px 0 0'>{_initials}</div>"
                    )
                if score is not None:
                    _sc_bg = _C["green"] if score >= 70 else (_C["amber"] if score >= 50 else _C["muted"])
                    _score_overlay = (
                        f"<div style='position:absolute;top:6px;right:6px;background:{_sc_bg};"
                        f"border-radius:12px;padding:2px 7px;font-size:0.65rem;font-weight:700;"
                        f"color:#fff'>{score}</div>"
                    )
                else:
                    _score_overlay = ""
                with col:
                    st.markdown(
                        f"<div style='border-radius:8px;overflow:hidden;background:{_C['surface']};"
                        f"margin-bottom:3px'>"
                        f"<div style='position:relative'>"
                        f"{_media}{_score_overlay}"
                        f"<div style='position:absolute;bottom:0;left:0;right:0;"
                        f"background:linear-gradient(transparent,rgba(0,0,0,0.82));padding:24px 8px 7px'>"
                        f"<div style='font-size:0.8rem;font-weight:700;color:{_C['text']};"
                        f"white-space:nowrap;overflow:hidden;text-overflow:ellipsis'>{name}</div>"
                        f"<div style='font-size:0.72rem;color:{_C['muted']};margin-top:1px;"
                        f"white-space:nowrap;overflow:hidden;text-overflow:ellipsis'>{stat_line}</div>"
                        f"</div></div></div>",
                        unsafe_allow_html=True,
                    )
                    if st.button("View", key=btn_key, width="stretch"):
                        st.session_state["mr_priority_load"] = name
                        st.rerun()

            # ── If profile loaded: show profile. Otherwise: show card grid. ────────
            if _profile:
                _back_c, _ref_c = st.columns([8, 1])
                with _back_c:
                    if st.button("← Back to list", key="mr_back"):
                        st.session_state.pop("mr_profile", None)
                        st.session_state.pop("mr_query", None)
                        st.rerun()
                with _ref_c:
                    if st.button("🔄 Refresh", key="mr_force_refresh", width="stretch", help="Force re-fetch all data"):
                        with st.spinner("Refreshing…"):
                            _refreshed = lookup_model_profile(_profile.get("name",""), force_refresh=True)
                        st.session_state["mr_profile"] = _refreshed
                        st.rerun()
            else:
                # ── Section header row: label + refresh icon ────────────────────
                _sh1, _sh2 = st.columns([11, 1])
                with _sh1:
                    st.markdown(
                        f"<span style='font-size:0.7rem;font-weight:700;letter-spacing:.08em;"
                        f"color:{_C['amber']};text-transform:uppercase'>🔥 Trending Now</span>",
                        unsafe_allow_html=True)
                with _sh2:
                    if st.button("↺", key="mr_refresh_trending", help="Refresh both sections"):
                        st.session_state.pop("mr_trending", None)
                        st.session_state.pop("mr_trending_photos_v1", None)
                        st.session_state.pop("mr_priority_photos_v3", None)
                        try:
                            import pathlib
                            pathlib.Path(
                                os.path.join(os.path.dirname(__file__), "research_cache_v2", "_trending_models.json")
                            ).unlink(missing_ok=True)
                        except Exception:
                            pass
                        st.rerun()

                # ── Trending grid ─────────────────────────────────────────────
                if _trending:
                    for _trow in [_trending[:5], _trending[5:10]]:
                        if not _trow:
                            continue
                        _tcols = st.columns(5, gap="small")
                        for _ti, _tm in enumerate(_trow):
                            if _ti >= 5:
                                break
                            _tname  = _tm["name"]
                            _tphoto = _trending_photos.get(_tname) or _tm.get("photo_url") or None
                            _tplat  = _tm.get("platform", "")
                            _tsc    = _tm.get("scenes", "")
                            _tfol   = _tm.get("followers", "")
                            _parts  = [p for p in [_tplat, (_tsc + " scenes") if _tsc else "", _tfol] if p]
                            _stat_line = " · ".join(_parts)
                            _model_card(_tcols[_ti], _tname, _tphoto, _stat_line,
                                        f"tr_{_tname.replace(' ', '_')}")
                else:
                    st.caption("Could not load trending models — check connection or click ↺.")

                # ── Priority header ───────────────────────────────────────────
                st.markdown(
                    f"<div style='margin-top:18px'>"
                    f"<span style='font-size:0.7rem;font-weight:700;letter-spacing:.08em;"
                    f"color:{_C['green']};text-transform:uppercase'>⭐ Priority Outreach</span>"
                    f"</div>",
                    unsafe_allow_html=True)

                # ── Priority grid ─────────────────────────────────────────────
                for _prow in [_PRIORITY[:5], _PRIORITY[5:]]:
                    _pcols = st.columns(5, gap="small")
                    for _pi, _pm in enumerate(_prow):
                        _pname = _pm["name"]
                        _pbytes = _priority_photos.get(_pname)
                        _mo     = _pm["mo"]
                        _mo_str = "Never booked" if _mo is None else (f"{_mo} mo ago")
                        _stat_line = f"{_pm['agency']} · {_mo_str}"
                        _pscore = compute_opportunity_score(_pname, {}) if _HAS_BK_HIST else None
                        _model_card(_pcols[_pi], _pname, _pbytes, _stat_line,
                                    f"pr_{_pname.replace(' ', '_')}", score=_pscore)

            if _profile:
                _bio     = _profile.get("bio", {})
                _about   = _profile.get("about", "")
                _booking = _profile.get("booking", {})
                _slr     = _profile.get("slr", [])
                _vrp     = _profile.get("vrp", [])
                _ddg     = _profile.get("ddg_bio", "")
                _photo   = _profile.get("photo_url", "")
                _name    = _profile.get("name", "")

                _bk_agency = _booking.get("agency", "") if _booking else ""
                _bk_data   = _booking.get("data", {}) if _booking else {}
                _rank_raw  = _bk_data.get("rank", "").strip()
                _RANK_BADGE = {
                    "great":    ("🟢", "Great",    "#1a4a1a"),
                    "good":     ("🔵", "Good",     "#0d2d4a"),
                    "moderate": ("🟡", "Moderate", "#4a3d00"),
                    "poor":     ("🔴", "Poor",     "#4a0d0d"),
                }
                _rank_info = _RANK_BADGE.get(_rank_raw.lower())

                # ── Calculate age from birthday ───────────────────────────────────
                def _calc_age(bday_str: str) -> str:
                    if not bday_str:
                        return ""
                    import re as _re2
                    from datetime import datetime as _dt, date as _date
                    # Strip ordinals: 4th → 4, 1st → 1, etc.
                    cleaned = _re2.sub(r'(\d+)(st|nd|rd|th)\b', r'\1', bday_str, flags=_re2.I)
                    for fmt in ("%A %d of %B %Y", "%d of %B %Y", "%B %d %Y",
                                "%B %d, %Y", "%d %B %Y", "%Y-%m-%d", "%m/%d/%Y"):
                        try:
                            bd = _dt.strptime(cleaned.strip(), fmt).date()
                            today = _date.today()
                            age = today.year - bd.year - ((today.month, today.day) < (bd.month, bd.day))
                            return str(age) if 18 <= age <= 80 else ""
                        except ValueError:
                            continue
                    return ""

                _bday_str = _bio.get("birthday", "") or _bk_data.get("born", "")
                _age_str  = _calc_age(_bday_str)

                # ── Booking history data (computed once, used in header) ──────────
                _hist = get_booking_history(_name) if _HAS_BK_HIST else None
                _h_total = 0; _last_fmt = ""; _rate_str = ""; _studios_str = ""
                if _hist:
                    _h_total = _hist.get("total_bookings", 0)
                    _h_last  = _hist.get("last_date", "")
                    _h_avg   = _hist.get("avg_rate")
                    _h_trend = _hist.get("rate_trend", "")
                    try:
                        from datetime import datetime as _dt2
                        _ld = _dt2.fromisoformat(_h_last)
                        _mo_ago = int((_dt2.now() - _ld).days / 30.44)
                        _last_fmt = f"{_ld.strftime('%b %Y')} ({_mo_ago} mo ago)"
                    except Exception:
                        _last_fmt = _h_last
                    _trend_icon = {"up": "↑", "down": "↓", "stable": "→"}.get(_h_trend, "")
                    _rate_str = f"${_h_avg:,}" if _h_avg else "—"
                    if _trend_icon:
                        _rate_str += f" {_trend_icon}"
                    _studios_str = " · ".join(
                        f"{s} ({n}×)" for s, n in
                        sorted(_hist.get("studios", {}).items(), key=lambda x: -x[1])
                    )

                # ── Profile header: photo | name/agency/info | booking history ──
                _hc_photo, _hc_info, _hc_booking = st.columns([2, 4, 4])

                with _hc_photo:
                    if _photo:
                        try:
                            st.image(_photo, width="stretch")
                        except Exception:
                            pass

                with _hc_info:
                    # Name + age + rank badge
                    _age_badge = f" <span style='font-size:1rem;color:{_C['muted']}'>{_age_str}</span>" if _age_str else ""
                    _rank_badge_html = ""
                    if _rank_info:
                        _ri_icon, _ri_label, _ri_bg = _rank_info
                        _rank_badge_html = (f" <span style='background:{_ri_bg};border-radius:4px;"
                                            f"padding:2px 9px;font-size:0.8rem;font-weight:600'>"
                                            f"{_ri_icon} {_ri_label}</span>")
                    st.markdown(
                        f"<div style='font-size:1.6rem;font-weight:700;line-height:1.3;margin-bottom:6px'>"
                        f"{_name}{_age_badge}{_rank_badge_html}</div>",
                        unsafe_allow_html=True)

                    # Agency + profile links
                    if _bk_agency:
                        _agency_url    = _booking.get("agency_url", "")
                        _model_slr_url = _bk_data.get("slr_profile_url", "")
                        _model_vrp_url = _bk_data.get("vrp_profile_url", "")
                        _agency_label  = f"[{_bk_agency}]({_agency_url})" if _agency_url else _bk_agency
                        _plinks = []
                        if _model_slr_url: _plinks.append(f"[SLR]({_model_slr_url})")
                        if _model_vrp_url: _plinks.append(f"[VRPorn]({_model_vrp_url})")
                        _pstr = "  ·  " + "  ·  ".join(_plinks) if _plinks else ""
                        st.markdown(f"{_agency_label}{_pstr}")
                    else:
                        st.caption("Not in booking sheet")

                    # Rate + status + location
                    _info_parts = []
                    _avg_rate_h = _bk_data.get("avg rate", "")
                    _status_h   = _bk_data.get("status", "")
                    _loc_h      = _bk_data.get("location", "")
                    if _avg_rate_h: _info_parts.append(f"💰 {_avg_rate_h}")
                    if _status_h:   _info_parts.append(_status_h)
                    if _loc_h:      _info_parts.append(f"📍 {_loc_h}")
                    if _info_parts:
                        st.caption("  ·  ".join(_info_parts))

                    # Available for tags
                    _avail_h = _bk_data.get("available for", "")
                    if _avail_h:
                        _tags_html = " ".join(
                            f"<span style='background:{_C['blue_dim']};border-radius:4px;padding:2px 7px;"
                            f"font-size:0.75rem;margin:2px;display:inline-block'>{t.strip()}</span>"
                            for t in _avail_h.split(",") if t.strip()
                        )
                        st.markdown(_tags_html, unsafe_allow_html=True)

                with _hc_booking:
                    if _hist:
                        st.markdown(
                            f"<div style='background:{_C['green_dim']};border:1px solid {_C['green']}33;border-radius:8px;"
                            f"padding:12px 14px'>"
                            f"<div style='color:{_C['green']};font-size:1.4rem;font-weight:700'>{_h_total}× booked</div>"
                            f"<div style='color:{_C['text']};font-size:0.82rem;margin-top:5px'>Last: {_last_fmt}</div>"
                            f"<div style='color:{_C['text']};font-size:0.82rem'>Rate: {_rate_str}</div>"
                            f"<div style='color:{_C['muted']};font-size:0.75rem;margin-top:5px'>{_studios_str}</div>"
                            f"</div>",
                            unsafe_allow_html=True)
                    elif _HAS_BK_HIST:
                        st.markdown(
                            f"<div style='background:{_C['red_dim']};border:1px solid {_C['red']}33;border-radius:8px;"
                            f"padding:12px 14px'>"
                            f"<span style='color:{_C['red']};font-size:0.88rem'>🔴 Never booked with your studio</span>"
                            f"</div>",
                            unsafe_allow_html=True)

                st.divider()

                # ── AI Booking Brief ───────────────────────────────────────────────
                _brief_key = f"booking_brief_{_name.replace(' ','_')}"
                with st.expander("✦ Generate Booking Brief", expanded=False):
                    if st.button("Generate", key=f"gen_brief_{_name.replace(' ','_')}"):
                        _claude_key = os.environ.get("ANTHROPIC_API_KEY", "")
                        if _claude_key:
                            import anthropic as _anth
                            # Build context from all available data
                            _ctx_parts = [f"Performer: {_name}"]
                            if _rank_raw: _ctx_parts.append(f"Booking rank: {_rank_raw}")
                            if _bk_agency: _ctx_parts.append(f"Agency: {_bk_agency}")
                            _slr_f = _bk_data.get("slr followers",""); _slr_sc = _bk_data.get("slr scenes","")
                            _vrp_f = _bk_data.get("vrp followers",""); _vrp_v = _bk_data.get("vrp views","")
                            if _slr_f: _ctx_parts.append(f"SLR followers: {_slr_f}")
                            if _slr_sc: _ctx_parts.append(f"SLR scenes: {_slr_sc}")
                            if _vrp_f: _ctx_parts.append(f"VRPorn followers: {_vrp_f}")
                            if _vrp_v: _ctx_parts.append(f"VRPorn views: {_vrp_v}")
                            _avail = _bk_data.get("available for","")
                            if _avail: _ctx_parts.append(f"Available for: {_avail}")
                            _avg_rate = _bk_data.get("avg rate","")
                            if _avg_rate: _ctx_parts.append(f"Avg rate: {_avg_rate}")
                            if _hist:
                                _ctx_parts.append(f"Booked with our studio: {_hist['total_bookings']} times")
                                _ctx_parts.append(f"Last booked: {_hist.get('last_date','')}")
                                _ctx_parts.append(f"Average rate paid: ${_hist.get('avg_rate','N/A')}")
                                _ctx_parts.append(f"Rate trend: {_hist.get('rate_trend','unknown')}")
                            else:
                                _ctx_parts.append("Never booked with our studio")
                            if _age_str: _ctx_parts.append(f"Age: {_age_str}")

                            _brief_prompt = (
                                "You are a talent booking advisor for a VR adult content studio. "
                                "Based on the data below, write a concise 3-sentence booking brief. "
                                "Cover: (1) her current market standing and platform performance, "
                                "(2) your studio's history with her and what that means, "
                                "(3) a clear recommendation — Book Now / Re-book / Monitor / Pass — with one-line reason.\n\n"
                                + "\n".join(_ctx_parts)
                            )
                            try:
                                _bac = _anth.Anthropic(api_key=_claude_key)
                                _bm = _bac.messages.create(
                                    model="claude-sonnet-4-6",
                                    max_tokens=300,
                                    messages=[{"role": "user", "content": _brief_prompt}]
                                )
                                st.session_state[_brief_key] = _bm.content[0].text
                            except Exception as _be:
                                st.session_state[_brief_key] = f"Error: {_be}"
                        else:
                            st.session_state[_brief_key] = "No ANTHROPIC_API_KEY set."

                    _brief_text = st.session_state.get(_brief_key, "")
                    if _brief_text:
                        st.markdown(
                            f"<div class='hub-card' style='font-size:0.88rem;line-height:1.6;"
                            f"color:{_C['text']}'>{_brief_text}</div>",
                            unsafe_allow_html=True
                        )

                # ── Wide two-column layout ────────────────────────────────────────
                _col_bio, _col_scenes = st.columns([5, 7])

                with _col_bio:

                    # ── Notes (prominent warning) ─────────────────────────────────
                    _notes = _bk_data.get("notes", "")
                    if _notes:
                        st.warning(f"**Notes:** {_notes}", icon="⚠️")

                    # ── Compact stats table ───────────────────────────────────────
                    def _stat_row(icon, label, value):
                        return (f"<tr><td style='color:{_C['muted']};padding:3px 10px 3px 0;"
                                f"white-space:nowrap;font-size:0.85rem'>{icon} {label}</td>"
                                f"<td style='font-weight:600;font-size:0.9rem'>{value}</td></tr>")

                    if _bk_agency:
                        # Booking stats
                        _bk_rows = []
                        for _k, _icon, _lbl in [
                            ("avg rate",         "💰", "Rate"),
                            ("bookings",         "📋", "Bookings"),
                            ("last booked date", "📅", "Last Booked"),
                            ("location",         "📍", "Location"),
                            ("status",           "✅", "Status"),
                        ]:
                            _v = _bk_data.get(_k, "")
                            if _v:
                                _bk_rows.append(_stat_row(_icon, _lbl, _v))

                        # Platform stats (skip zeros)
                        _plat_rows = []
                        for _k, _icon, _lbl in [
                            ("slr followers", "👥", "SLR Followers"),
                            ("slr scenes",    "🎬", "SLR Scenes"),
                            ("slr views",     "👁",  "SLR Views"),
                            ("vrp followers", "👥", "VRP Followers"),
                            ("vrp views",     "👁",  "VRP Views"),
                            ("povr views",    "👁",  "POVR Views"),
                        ]:
                            _v = _bk_data.get(_k, "")
                            if _v and _v not in ("0", ""):
                                _plat_rows.append(_stat_row(_icon, _lbl, _v))

                        # Social (skip zeros)
                        _soc_rows = []
                        for _k, _icon, _lbl in [
                            ("onlyfans",  "🔞", "OnlyFans"),
                            ("twitter",   "𝕏",  "Twitter"),
                            ("instagram", "📸", "Instagram"),
                        ]:
                            _v = _bk_data.get(_k, "")
                            if _v and _v not in ("0", ""):
                                _soc_rows.append(_stat_row(_icon, _lbl, _v))

                        all_rows = _bk_rows
                        if _plat_rows:
                            all_rows += [f"<tr><td colspan='2' style='padding-top:8px;"
                                         f"color:{_C['muted']};font-size:0.75rem;text-transform:uppercase;"
                                         f"letter-spacing:.05em'>Platform</td></tr>"] + _plat_rows
                        if _soc_rows:
                            all_rows += [f"<tr><td colspan='2' style='padding-top:8px;"
                                         f"color:{_C['muted']};font-size:0.75rem;text-transform:uppercase;"
                                         f"letter-spacing:.05em'>Social</td></tr>"] + _soc_rows

                        if all_rows:
                            st.markdown(
                                f"<table style='border-collapse:collapse;width:100%'>"
                                + "".join(all_rows) + "</table>",
                                unsafe_allow_html=True
                            )


                    # ── Physical stats ────────────────────────────────────────────
                    _BIO_KEYS = [
                        (None,             "birthday",    "Born"),
                        (None,             "birthplace",  "Birthplace"),
                        (None,             "nationality", "Nationality"),
                        (None,             "ethnicity",   "Ethnicity"),
                        ("height",         "height",      "Height"),
                        ("weight",         "weight",      "Weight"),
                        ("measurements",   "measurements","Measurements"),
                        (None,             "bra/cup size","Bra / Cup"),
                        ("hair",           "hair",        "Hair"),
                        ("eyes",           "eyes",        "Eyes"),
                        ("natural breasts","boobs",       "Natural Breasts"),
                        ("tattoos",        None,          "Tattoos"),
                        ("shoe size",      None,          "Shoe Size"),
                        (None,             "years active","Years Active"),
                    ]
                    _seen = set()
                    _bio_rows = []
                    for _bk_key, _bio_key, _label in _BIO_KEYS:
                        if _label in _seen:
                            continue
                        _val = (_bk_data.get(_bk_key) if _bk_key else "") or \
                               (_bio.get(_bio_key) if _bio_key else "")
                        if _val:
                            _bio_rows.append({"": _label, " ": _val})
                            _seen.add(_label)

                    if _bio_rows:
                        _phys_html = (
                            f"<p style='color:{_C['muted']};font-size:0.75rem;text-transform:uppercase;"
                            f"letter-spacing:.05em;margin:10px 0 4px'>Physical Stats</p>"
                            "<table style='border-collapse:collapse;width:100%'>"
                            + "".join(_stat_row("", r[""], r[" "]) for r in _bio_rows)
                            + "</table>"
                        )
                        st.markdown(_phys_html, unsafe_allow_html=True)

                    # About / bio blurb
                    if _about:
                        st.markdown(
                            f"<p style='font-size:0.85rem;color:{_C['text']};line-height:1.5;"
                            f"margin-top:8px'>{_about}</p>",
                            unsafe_allow_html=True
                        )
                    elif _ddg and not _bio_rows:
                        st.markdown(
                            f"<p style='font-size:0.85rem;color:{_C['text']};line-height:1.5;"
                            f"margin-top:8px'>{_ddg[:400]}</p>",
                            unsafe_allow_html=True
                        )

                    # Source links (subtle, at bottom)
                    _sources = _profile.get("_sources", [])
                    if _sources:
                        st.caption("Sources: " + " · ".join(
                            f"[{_sn}]({_su})" for _sn, _su in _sources
                        ))

                # ── Scenes column ─────────────────────────────────────────────────
                with _col_scenes:
                    _slr_label = f"SexLikeReal ({len(_slr)})" if _slr else "SexLikeReal"
                    _vrp_label = f"VRPorn ({len(_vrp)})"      if _vrp else "VRPorn"
                    _stab_slr, _stab_vrp = st.tabs([_slr_label, _vrp_label])

                    def _render_scenes(scenes, platform_name, base_url=""):
                        if not scenes:
                            st.info(f"No {platform_name} scenes found.", icon="🎬")
                            return
                        for _sc in scenes:
                            _sc_title    = _sc.get("title", "")
                            _sc_date     = _sc.get("date", "")
                            _sc_studio   = _sc.get("studio", "")
                            _sc_url      = _sc.get("url", "")
                            _sc_thumb    = _sc.get("thumb", "")
                            _sc_views    = _sc.get("views", "")
                            _sc_likes    = _sc.get("likes", "")
                            _sc_comments = _sc.get("comments", "")
                            _sc_duration = _sc.get("duration", "")
                            if not _sc_title:
                                continue
                            if _sc_url and not _sc_url.startswith("http") and base_url:
                                _sc_url = base_url + _sc_url

                            # Card-style container
                            with st.container(border=True):
                                _tc, _ic = st.columns([2, 5])
                                with _tc:
                                    if _sc_thumb:
                                        try:
                                            st.image(_sc_thumb, width="stretch")
                                        except Exception:
                                            pass
                                with _ic:
                                    if _sc_url:
                                        st.markdown(f"**[{_sc_title}]({_sc_url})**")
                                    else:
                                        st.markdown(f"**{_sc_title}**")
                                    # Meta: studio · date · duration
                                    _meta = []
                                    if _sc_studio:   _meta.append(f"🎬 {_sc_studio}")
                                    if _sc_date:     _meta.append(f"📅 {_sc_date}")
                                    if _sc_duration: _meta.append(f"⏱️ {_sc_duration}")
                                    if _meta:
                                        st.caption("  ·  ".join(_meta))
                                    # Stats: views · likes · comments
                                    _stats = []
                                    if _sc_views:    _stats.append(f"👁 {_sc_views}")
                                    if _sc_likes:    _stats.append(f"❤️ {_sc_likes}")
                                    if _sc_comments: _stats.append(f"💬 {_sc_comments}")
                                    if _stats:
                                        st.caption("  ·  ".join(_stats))

                    with _stab_slr:
                        _render_scenes(_slr, "SLR", "https://sexlikereal.com")

                    with _stab_vrp:
                        _render_scenes(_vrp, "VRPorn", "https://vrporn.com")

                    # ── Competitor Activity ────────────────────────────────────────
                    _all_scenes = _slr + _vrp
                    _competitor_scenes = get_competitor_scenes(_all_scenes) if _HAS_BK_HIST else []
                    if _competitor_scenes:
                        st.markdown(
                            f"<p style='color:{_C['muted']};font-size:0.72rem;text-transform:uppercase;"
                            f"letter-spacing:.05em;margin:14px 0 6px'>Competitor Activity</p>",
                            unsafe_allow_html=True
                        )
                        for _cs in _competitor_scenes:
                            _cs_date = _cs.get("date","")
                            _cs_title = _cs.get("title","")
                            _cs_studio = _cs.get("studio","")
                            _cs_label = f"**{_cs_studio}** · {_cs_date}"
                            if _cs_title:
                                _cs_label += f" — _{_cs_title[:50]}_"
                            st.caption(_cs_label)


    # ── TAB 6: Description Generator ──────────────────────────────────────────────
with tab_desc:
    if _has_tab("Descriptions"):

        # ── Helper: build .docx bytes ───────────────────────────────────────────────
        @st.cache_data
        def _build_docx(talent_line, title, tags, categories, writeup, meta_title="", meta_desc="", studio="FPVR"):
            """Return bytes of a formatted .docx matching the studio template."""
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import io

            doc = Document()
            for p in doc.paragraphs:
                p._element.getparent().remove(p._element)

            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(10)

            def _add_hyperlink(paragraph, text, url):
                """Add a clickable hyperlink to a paragraph."""
                part = paragraph.part
                r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                run_el = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                c = OxmlElement('w:color')
                c.set(qn('w:val'), '0563C1')
                rPr.append(c)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                run_el.append(rPr)
                txt = OxmlElement('w:t')
                txt.text = text
                txt.set(qn('xml:space'), 'preserve')
                run_el.append(txt)
                hyperlink.append(run_el)
                paragraph._element.append(hyperlink)

            def _add_line(bold_prefix, normal_text=""):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                if bold_prefix:
                    run = p.add_run(bold_prefix)
                    run.bold = True
                if normal_text:
                    p.add_run(normal_text)
                return p

            def _add_linked_line(bold_prefix, items_csv, url_fn):
                """Add a line with bold prefix and comma-separated hyperlinked items."""
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(bold_prefix)
                run.bold = True
                items = [i.strip() for i in items_csv.split(",") if i.strip()]
                for idx, item in enumerate(items):
                    url = url_fn(studio, item)
                    _add_hyperlink(p, item, url)
                    if idx < len(items) - 1:
                        p.add_run(", ")
                return p

            _add_line(talent_line)
            _add_line("Title: ", title)
            _add_linked_line("Tags: ", tags, _tag_url)
            _add_linked_line("Categories: ", categories, _category_url)
            _add_line("Write-up:")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(writeup)
            if meta_title:
                _add_line("")
                _add_line("Meta Title: ", meta_title)
            if meta_desc:
                _add_line("Meta Description: ", meta_desc)

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        def _name_slug(full_name: str) -> str:
            """'First Last' → 'FirstLast'"""
            return full_name.strip().replace(" ", "")

        # ── Batch: Find Missing Descriptions ────────────────────────────────────────
        _SCAN_PATH = os.path.join(os.path.dirname(__file__), "mega_scan.json")

        def _find_scene_data(scene_id: str) -> dict:
            import time, sys
            cache_key = f"_scene_cache_{scene_id}"
            cache_ts_key = f"_scene_cache_ts_{scene_id}"
            now = time.time()
            if cache_key in st.session_state:
                age = now - st.session_state.get(cache_ts_key, 0)
                if age < 300:
                    return st.session_state[cache_key]
            result = {"female": "", "male": "", "title": "", "tags": "", "categories": "", "plot": ""}
            try:
                _sheets_mod_path = os.path.dirname(__file__)
                if _sheets_mod_path not in sys.path:
                    sys.path.insert(0, _sheets_mod_path)
                import sheets_integration as _si
                import re as _re
                _m = _re.match(r'^([A-Za-z]+)(\d+)$', scene_id.strip())
                if not _m:
                    return result
                _scene_studio = _m.group(1).upper()
                _scene_num = int(_m.group(2))
                _scene_num_padded = f"{_scene_num:04d}"
                _scene_full = f"{_scene_studio}{_scene_num_padded}"
                sh = _si.get_spreadsheet()
                for ws in _si.month_tabs(sh):
                    rows = ws.get_all_values()
                    for row in rows[1:]:
                        row = _si._pad(row, 13)
                        _col_studio = row[_si.COL_STUDIO].strip().upper()
                        _col_scene  = row[_si.COL_SCENE].strip()
                        _constructed = f"{_col_studio}{_col_scene.zfill(4)}" if _col_scene.isdigit() else _col_scene.upper()
                        if _constructed == _scene_full or _col_scene.upper() == _scene_full:
                            result["female"]     = row[_si.COL_FEMALE].strip()
                            result["male"]       = row[_si.COL_MALE].strip()
                            result["plot"]       = row[_si.COL_PLOT].strip()
                            result["title"]      = row[10].strip() if len(row) > 10 else ""
                            result["tags"]       = ""
                            result["categories"] = ""
                            st.session_state[cache_key] = result
                            st.session_state[cache_ts_key] = now
                            return result
            except Exception as _e:
                st.warning(f"Sheet lookup failed for {scene_id}: {_e}")
            st.session_state[cache_key] = result
            st.session_state[cache_ts_key] = now
            return result

        _STUDIO_COLORS = {"FPVR": "#3b82f6", "VRH": "#8b5cf6", "VRA": "#ec4899", "NJOI": "#f97316"}
        _GRAIL_ID  = "1Eq5G5FU6A8EqeFZCnZjrEaMYS8F1DiK5vP5tCSINeJk"
        _GRAIL_TAB = {"FPVR": "FPVR", "VRH": "VRH", "VRA": "VRA", "NJOI": "NNJOI"}

        def _batch_prefetch(scene_list: list) -> dict:
            import time, sys, re as _re5
            _ck  = "_desc_batch_names"
            _tsk = "_desc_batch_names_ts"
            _now = time.time()
            if _ck in st.session_state and _now - st.session_state.get(_tsk, 0) < 600:
                return st.session_state[_ck]
            out = {sc["scene_id"]: {"title": "", "categories": "", "tags": "", "plot": ""}
                   for sc in scene_list}
            try:
                _sp = os.path.dirname(__file__)
                if _sp not in sys.path:
                    sys.path.insert(0, _sp)
                import sheets_integration as _si
                import gspread
                from google.oauth2.service_account import Credentials
                _creds = Credentials.from_service_account_file(
                    os.path.join(os.path.dirname(__file__), "service_account.json"),
                    scopes=["https://www.googleapis.com/auth/spreadsheets"])
                _gc = gspread.authorize(_creds)
                _grail_sh = _gc.open_by_key(_GRAIL_ID)
                _by_tab = {}
                for sc in scene_list:
                    _tab = _GRAIL_TAB.get(sc["studio"], sc["studio"])
                    _m   = _re5.match(r'^[A-Za-z]+(\d+)$', sc["scene_id"])
                    if _m:
                        _by_tab.setdefault(_tab, []).append((sc["scene_id"], int(_m.group(1))))
                for _tab, _entries in _by_tab.items():
                    try:
                        _ws   = _grail_sh.worksheet(_tab)
                        _rows = _ws.get_all_values()
                        _id_map = {int(r[1]): r for r in _rows[1:] if len(r) > 1 and r[1].isdigit()}
                        for _sid, _num in _entries:
                            _r = _id_map.get(_num)
                            if _r:
                                out[_sid]["title"]      = _r[3].strip() if len(_r) > 3 else ""
                                out[_sid]["categories"] = _r[5].strip() if len(_r) > 5 else ""
                                out[_sid]["tags"]       = _r[6].strip() if len(_r) > 6 else ""
                    except Exception as _gpf:
                        st.warning(f"Could not load Grail data: {_gpf}")
                _female_map = {}
                for sc in scene_list:
                    _f = sc.get("female", "").strip().lower()
                    if _f and _f not in _female_map:
                        _female_map[_f] = sc["scene_id"]
                _scripts_sh = _gc.open_by_key(_si.SHEET_ID)
                for _ws in _si.month_tabs(_scripts_sh):
                    _rows = _ws.get_all_values()
                    for _row in _rows[1:]:
                        _row = _si._pad(_row, 13)
                        _rf   = _row[_si.COL_FEMALE].strip().lower()
                        _plot = _row[_si.COL_PLOT].strip()
                        if _plot and _rf in _female_map:
                            _sid = _female_map[_rf]
                            if not out[_sid].get("plot"):
                                out[_sid]["plot"] = _plot
            except Exception:
                pass
            st.session_state[_ck]  = out
            st.session_state[_tsk] = _now
            return out

        # Load scan data once
        import json as _json
        if "scan_data" not in st.session_state:
            try:
                with open(_SCAN_PATH, "r", encoding="utf-8") as _f:
                    st.session_state["scan_data"] = _json.load(_f)
            except Exception:
                st.session_state["scan_data"] = None
        _scan_data = st.session_state["scan_data"]
        _missing_scenes = []
        if _scan_data:
            _missing_scenes = [s for s in _scan_data.get("scenes", []) if not s.get("has_description", True)]
        _names_cache = _batch_prefetch(_missing_scenes) if _missing_scenes else {}

        # Handle card load clicks BEFORE widgets render
        if "desc_load_trigger" in st.session_state:
            _trigger    = st.session_state.pop("desc_load_trigger")
            _scene_id_t = _trigger["scene_id"]
            _studio_t   = _trigger["studio"]
            import re as _re3
            _nm3 = _re3.match(r'^[A-Za-z]+(\d+)$', _scene_id_t)
            st.session_state["d_studio"]    = _studio_t
            st.session_state["d_scene_num"] = int(_nm3.group(1)) if _nm3 else 1
            _scan_entry = next((s for s in _missing_scenes if s["scene_id"] == _scene_id_t), {})
            _sheet_data = _names_cache.get(_scene_id_t, {})
            _female_t   = _trigger.get("female") or _scan_entry.get("female", "")
            _male_t     = _trigger.get("male")   or _scan_entry.get("male", "")
            if _female_t:                         st.session_state["d_female"] = _female_t
            if _male_t:                           st.session_state["d_male"]   = _male_t
            if _sheet_data.get("title"):          st.session_state["d_title"]  = _sheet_data["title"]
            if _sheet_data.get("plot"):           st.session_state["d_plot"]   = _sheet_data["plot"]
            if _sheet_data.get("categories"):     st.session_state["d_cats"]   = _sheet_data["categories"]
            if _sheet_data.get("tags"):           st.session_state["d_tags"]   = _sheet_data["tags"]
            st.session_state["desc_active_scene"] = _scene_id_t
            st.session_state["_desc_auto_gen"]    = True
            st.session_state.pop("d_parsed", None)
            st.session_state.pop("d_writeup", None)
            st.session_state.pop("d_writeup_edit", None)
            st.session_state.pop("_desc_saved_mega", None)

        # ── TWO COLUMN LAYOUT ───────────────────────────────────────────────────────
        _col_q, _col_f = st.columns([1, 3], gap="medium")

        # ── LEFT: Batch Queue ───────────────────────────────────────────────────────
        with _col_q:
            if _scan_data:
                _scan_ts = _scan_data.get("scanned_at", "")[:10]
                _has_cnt = len(_scan_data.get("scenes", [])) - len(_missing_scenes)
                st.markdown(
                    f"<div style='margin-bottom:12px'>"
                    f"<div style='font-size:0.9rem;font-weight:700;color:{_C['text']}'>"
                    f"Missing <span style='background:{_C['elevated']};color:{_C['muted']};font-size:0.65rem;"
                    f"padding:2px 8px;border-radius:10px;margin-left:4px'>{len(_missing_scenes)}</span></div>"
                    f"<div style='color:{_C['subtle']};font-size:0.65rem;margin-top:2px'>"
                    f"Scanned {_scan_ts} · {_has_cnt} complete</div></div>",
                    unsafe_allow_html=True
                )
            else:
                st.markdown(f"<div style='font-weight:700;color:{_C['text']};margin-bottom:6px'>Queue</div>",
                            unsafe_allow_html=True)
                st.caption("No scan file. Run `scan_mega.py` on Mac.")

            _active_scene = st.session_state.get("desc_active_scene", "")
            _dq_filter = st.text_input("Filter", key="desc_queue_filter",
                                       placeholder="Search scenes...",
                                       label_visibility="collapsed")
            # Studio filter tabs
            _dq_studios = sorted(set(s["studio"] for s in _missing_scenes))
            _dq_studio_opts = ["All"] + _dq_studios
            _dq_studio_filter = st.segmented_control(
                "Studio", _dq_studio_opts, default="All",
                key="desc_queue_studio", label_visibility="collapsed")
            _filtered_scenes = _missing_scenes
            if _dq_filter:
                _dq_lower = _dq_filter.lower()
                _filtered_scenes = [s for s in _filtered_scenes
                                    if _dq_lower in s.get("scene_id", "").lower()
                                    or _dq_lower in s.get("female", "").lower()
                                    or _dq_lower in s.get("male", "").lower()
                                    or _dq_lower in s.get("studio", "").lower()]
            if _dq_studio_filter and _dq_studio_filter != "All":
                _filtered_scenes = [s for s in _filtered_scenes if s["studio"] == _dq_studio_filter]
            # Pagination — 20 scenes per page
            _DQ_PAGE_SIZE = 20
            _dq_page = st.session_state.get("desc_queue_page", 0)
            _dq_total_pages = max(1, (len(_filtered_scenes) + _DQ_PAGE_SIZE - 1) // _DQ_PAGE_SIZE)
            if _dq_page >= _dq_total_pages:
                _dq_page = 0
            _dq_start = _dq_page * _DQ_PAGE_SIZE
            _dq_page_scenes = _filtered_scenes[_dq_start:_dq_start + _DQ_PAGE_SIZE]
            st.caption(f"Showing {_dq_start + 1}–{min(_dq_start + _DQ_PAGE_SIZE, len(_filtered_scenes))} of {len(_filtered_scenes)}")
            for _ms in _dq_page_scenes:
                _ms_id = _ms["scene_id"]
                _is_active = (_ms_id == _active_scene)
                _female = _ms.get("female", "")
                _male = _ms.get("male", "")
                _talent_display = _female or "—"
                if _male:
                    _talent_display += f" / {_male}"
                _grail_d = _names_cache.get(_ms_id, {})
                _has_plot = bool(_grail_d.get("plot"))
                _has_title = bool(_grail_d.get("title"))
                _ready_dot = "🟢" if (_has_plot and _has_title) else ("🟡" if _has_plot else "🔴")
                _btn_label = f"{_ready_dot} {_ms_id} · {_talent_display}"
                if st.button(_btn_label, key=f"load_{_ms_id}", width="stretch",
                             type="primary" if _is_active else "secondary"):
                    st.session_state["desc_load_trigger"] = {
                        "scene_id": _ms_id, "studio": _ms["studio"],
                        "female": _ms.get("female", ""), "male": _ms.get("male", ""),
                    }
                    st.rerun()
            # Pagination controls
            if _dq_total_pages > 1:
                _pn1, _pn2, _pn3 = st.columns([1, 2, 1])
                with _pn1:
                    if _dq_page > 0 and st.button("Prev", key="dq_prev"):
                        st.session_state["desc_queue_page"] = _dq_page - 1
                        st.rerun()
                with _pn2:
                    st.caption(f"Page {_dq_page + 1} / {_dq_total_pages}")
                with _pn3:
                    if _dq_page < _dq_total_pages - 1 and st.button("Next", key="dq_next"):
                        st.session_state["desc_queue_page"] = _dq_page + 1
                        st.rerun()

        # ── RIGHT: Form + Output ────────────────────────────────────────────────────
        with _col_f:
            hub_ui.section("Description Generator")
            _active_scene = st.session_state.get("desc_active_scene", "")
            if _active_scene:
                _ac_color = _STUDIO_COLORS.get(st.session_state.get("d_studio", ""), _C["muted"])
                st.markdown(
                    f"<div style='display:inline-block;background:{_ac_color}18;border:1px solid {_ac_color}40;"
                    f"color:{_ac_color};font-size:0.75rem;font-weight:700;padding:2px 10px;"
                    f"border-radius:10px;margin-bottom:8px'>Editing: {_active_scene}</div>",
                    unsafe_allow_html=True
                )

            # ── Form ──────────────────────────────────────────────────────────────
            with st.container(border=True):
                _da, _db, _dc = st.columns([3, 2, 2])
                with _da:
                    _d_studio = st.selectbox("Studio", list(_DESC_STUDIO_CONFIG.keys()), key="d_studio")
                with _db:
                    if "d_scene_num" not in st.session_state:
                        st.session_state["d_scene_num"] = 1
                    _d_scene_num = st.number_input("Scene #", min_value=1, max_value=9999,
                                                    step=1, key="d_scene_num")
                with _dc:
                    _d_res = st.selectbox("Resolution", ["8K", "6K", "4K"], key="d_res")

                _de, _df = st.columns(2)
                with _de:
                    _d_female = st.text_input("Female Talent(s)", placeholder="First Last, First Last", key="d_female")
                with _df:
                    _d_male = st.text_input("Male Talent(s)", placeholder="First Last (blank = solo)", key="d_male")

                _d_title = st.text_input("Scene Title", key="d_title")

                # Categories — multiselect from approved list per studio
                _dg, _dh = st.columns(2)
                with _dg:
                    _studio_cats = _STUDIO_CATEGORIES.get(_d_studio)
                    if _studio_cats:
                        _d_cats_list = st.multiselect("Categories", _studio_cats,
                                                       default=[c.strip() for c in st.session_state.get("d_cats", "").split(",") if c.strip() in _studio_cats] if st.session_state.get("d_cats") else [],
                                                       key="d_cats_ms")
                        _d_cats = ", ".join(_d_cats_list)
                    else:
                        _d_cats = st.text_input("Categories", placeholder="8K, Brunette, Natural Tits…", key="d_cats")
                with _dh:
                    _d_tags = st.text_input("Tags", placeholder="Sexy Brunette, POV, Cowgirl, Creampie…", key="d_tags")

                _d_plot = st.text_area("Scene Plot / Setup", placeholder="Paste the scene plot here…",
                                        height=100, key="d_plot")

            # ── Advanced / optional fields ─────────────────────────────────────────
            _adv_expanded = bool(
                st.session_state.get("d_model_props") or st.session_state.get("d_sex_pos") or
                st.session_state.get("d_wardrobe")   or st.session_state.get("d_target_kw")
            )
            with st.expander("📝 Advanced Details (optional)", expanded=_adv_expanded):
                _di, _dj = st.columns(2)
                with _di:
                    _d_model_props = st.text_input("Model Properties",
                                                    placeholder="Long dark hair, big natural tits, smackable ass",
                                                    key="d_model_props")
                with _dj:
                    _d_target_kw = st.text_input("Target Keywords",
                                                  placeholder=f"{_d_studio} VR porn, 8K VR porn",
                                                  key="d_target_kw")
                _d_wardrobe = st.text_input("Wardrobe", placeholder="White lace lingerie, black stockings",
                                             key="d_wardrobe")
                _d_sex_pos = st.text_area("Sex Positions (detailed)", placeholder="She drops to her knees for a BJ, then reverse cowgirl, cowgirl, doggy, missionary. Finishes with a creampie.",
                                           height=80, key="d_sex_pos")
                _studio_tags_ref = _STUDIO_TAGS.get(_d_studio)
                if _studio_tags_ref:
                    with st.expander(f"Approved {_d_studio} Tags Reference", expanded=False):
                        st.caption(_studio_tags_ref)

            # Generate / Regenerate buttons
            _d_gen_c, _d_reg_c = st.columns([1, 1])
            with _d_gen_c:
                _d_generate = st.button("Generate", type="primary",
                                         width="stretch", key="d_generate")
            with _d_reg_c:
                _d_regen = st.button("Regenerate All", width="stretch", key="d_regen",
                                      disabled=("d_parsed" not in st.session_state))

            _auto_gen = st.session_state.pop("_desc_auto_gen", False)

            # ── Generation logic ───────────────────────────────────────────────────
            if (_d_generate or _d_regen or _auto_gen) and _d_plot.strip():  # _d_regen2 triggers via rerun
                _d_cfg = _DESC_STUDIO_CONFIG[_d_studio]
                if _is_compilation(_d_title):
                    _studio_system = _DESC_SYSTEMS_COMPILATION.get(_d_studio, _DESC_SYSTEMS_FULL.get(_d_studio, ""))
                else:
                    _studio_system = _DESC_SYSTEMS_FULL.get(_d_studio, _DESC_SYSTEMS_FULL["VRH"])

                _user_prompt = _build_scene_prompt(
                    _d_studio, _d_cfg,
                    title=_d_title,
                    female=_d_female,
                    male=_d_male,
                    plot=_d_plot,
                    categories=_d_cats,
                    model_props=_d_model_props,
                    sex_positions=_d_sex_pos,
                    target_keywords=_d_target_kw,
                    resolution=_d_res,
                    wardrobe=st.session_state.get("d_wardrobe", ""),
                )

                _max_tok = 1200 if _d_studio == "FPVR" else 600

                _claude_key = os.environ.get("ANTHROPIC_API_KEY", "")
                if _claude_key:
                    import anthropic as _anth
                    with st.spinner("Generating description…"):
                        try:
                            _bac = _anth.Anthropic(api_key=_claude_key)
                            _bm  = _bac.messages.create(
                                model="claude-sonnet-4-6",
                                max_tokens=_max_tok,
                                system=_studio_system,
                                messages=[{"role": "user", "content": _user_prompt}]
                            )
                            _raw_text = _bm.content[0].text.strip()
                            st.session_state["d_writeup"] = _raw_text
                            st.session_state["d_parsed"] = _parse_desc_output(_raw_text)
                            st.session_state.pop("d_editing_para", None)
                        except Exception as _de2:
                            st.error(f"Generation failed: {_de2}")
                else:
                    st.warning("ANTHROPIC_API_KEY not set.")

            # ── Output: Inline paragraph editing ───────────────────────────────────
            if "d_parsed" in st.session_state:
                _parsed = st.session_state["d_parsed"]
                _d_cfg2    = _DESC_STUDIO_CONFIG[_d_studio]
                _f_names   = [n.strip() for n in _d_female.split(",") if n.strip()]
                _m_names   = [n.strip() for n in _d_male.split(",") if n.strip()]
                _talent_ln = ", ".join(_f_names + _m_names)
                _scene_id  = f"{_d_cfg2['prefix']}{int(_d_scene_num):04d}"
                _f_slug    = _name_slug(_f_names[0]) if _f_names else "Unknown"
                _m_slug    = ("-" + _name_slug(_m_names[0])) if _m_names else ""
                _filename_base = f"{_scene_id}-{_f_slug}{_m_slug}"

                st.divider()

                # ── Metrics row ──────────────────────────────────────────────
                _para_count = len(_parsed.get("paragraphs", []))
                _has_meta_t = bool(_parsed.get("meta_title", ""))
                _has_meta_d = bool(_parsed.get("meta_description", ""))
                _dm1, _dm2, _dm3 = st.columns(3)
                _dm1.metric("Scene", _scene_id)
                _dm2.metric("Paragraphs", _para_count)
                _dm3.metric("SEO", "Complete" if (_has_meta_t and _has_meta_d) else "Missing")

                _out_hdr, _out_regen = st.columns([4, 1])
                with _out_hdr:
                    st.markdown(f"<div class='sh'>Output — <code style='font-size:0.75rem'>{_filename_base}</code></div>",
                                unsafe_allow_html=True)
                with _out_regen:
                    _d_regen2 = st.button("Regen All", key="d_regen2", width="stretch")
                if _d_regen2:
                    st.session_state.pop("d_parsed", None)
                    st.session_state["_desc_auto_gen"] = True
                    st.rerun()

                # ── Paragraphs with click-to-edit ─────────────────────────────────
                _editing = st.session_state.get("d_editing_para", None)

                for _pi, _para in enumerate(_parsed.get("paragraphs", [])):
                    _p_title = _para.get("title", "")
                    _p_body  = _para.get("body", "")

                    if _editing == _pi:
                        # ── EDIT MODE ──────────────────────────────────────────
                        if _p_title:
                            _new_title = st.text_input(f"Title", value=_p_title, key=f"d_pt_{_pi}")
                        _new_body = st.text_area(f"Paragraph {_pi + 1}", value=_p_body,
                                                  height=150, key=f"d_pb_{_pi}")
                        _ec1, _ec2, _ec3 = st.columns([1, 1, 1])
                        with _ec1:
                            if st.button("Save Edit", key=f"d_save_{_pi}", width="stretch", type="primary"):
                                _parsed["paragraphs"][_pi]["body"] = _new_body
                                if _p_title:
                                    _parsed["paragraphs"][_pi]["title"] = _new_title
                                st.session_state["d_parsed"] = _parsed
                                st.session_state["d_writeup"] = _reassemble_desc(_parsed)
                                st.session_state.pop("d_editing_para", None)
                                st.rerun()
                        with _ec2:
                            if st.button("Regenerate", key=f"d_regen_p_{_pi}", width="stretch"):
                                _claude_key = os.environ.get("ANTHROPIC_API_KEY", "")
                                if _claude_key:
                                    import anthropic as _anth_r
                                    _regen_prompt = f"""Rewrite ONLY this paragraph (paragraph {_pi + 1}) of the scene description.
Keep the same style, tone, and flow as the rest of the description.

Current paragraph title: {_p_title}
Current paragraph text: {_p_body}

Scene context:
- Studio: {_DESC_STUDIO_CONFIG[_d_studio]['name']}
- Performer: {_d_female}
- Title: {_d_title}
- Plot: {_d_plot[:300]}

Rewrite this paragraph now. Output ONLY the paragraph text (no title, no meta fields). Keep similar length."""
                                _studio_system = _DESC_SYSTEMS_FULL.get(_d_studio, _DESC_SYSTEMS_FULL["VRH"])
                                with st.spinner(f"Regenerating paragraph {_pi + 1}…"):
                                    try:
                                        _bac_r = _anth_r.Anthropic(api_key=_claude_key)
                                        _bm_r = _bac_r.messages.create(
                                            model="claude-sonnet-4-6", max_tokens=500,
                                            system=_studio_system,
                                            messages=[{"role": "user", "content": _regen_prompt}])
                                        _parsed["paragraphs"][_pi]["body"] = _bm_r.content[0].text.strip()
                                        st.session_state["d_parsed"] = _parsed
                                        st.session_state["d_writeup"] = _reassemble_desc(_parsed)
                                        st.rerun()
                                    except Exception as _re2:
                                        st.error(f"Regeneration failed: {_re2}")
                    with _ec3:
                        if st.button("Cancel", key=f"d_cancel_{_pi}", width="stretch"):
                            st.session_state.pop("d_editing_para", None)
                            st.rerun()
                else:
                    # ── READ MODE — ✏️ button inline with title ───────────────
                    _r_lbl, _r_btn = st.columns([6, 1])
                    with _r_lbl:
                        if _p_title:
                            st.markdown(f"**{_p_title}**")
                    with _r_btn:
                        if st.button("✏️", key=f"d_edit_{_pi}", help="Edit this paragraph"):
                            st.session_state["d_editing_para"] = _pi
                            st.rerun()
                    st.markdown(
                        f"<div class='hub-card' style='margin-bottom:8px'>"
                        f"<div style='color:{_C['text']};font-size:0.82rem;line-height:1.6'>{_p_body}</div>"
                        f"</div>",
                        unsafe_allow_html=True
                    )

                st.divider()

                # ── SEO fields (for website publishing) ────────────────────────
                st.markdown("<div class='sh'>SEO — Website Meta Tags</div>", unsafe_allow_html=True)
                with st.container(border=True):
                    st.caption("Title tag and Google search snippet for the scene page.")
                    _meta_t = st.text_input("SEO Page Title", value=_parsed.get("meta_title", ""),
                                             key="d_meta_title")
                    _meta_d = st.text_input("SEO Description (max 160 chars)",
                                             value=_parsed.get("meta_description", ""),
                                             key="d_meta_desc", max_chars=200)
                    _md_len = len(_meta_d)
                    _md_color = _C["green"] if _md_len <= 160 else _C["red"]
                    # Progress-style char counter
                    _md_pct = min(int((_md_len / 160) * 100), 100)
                    st.markdown(
                        f"<div style='display:flex;align-items:center;gap:8px'>"
                        f"<div style='flex:1;height:4px;background:{_C['elevated']};border-radius:2px;overflow:hidden'>"
                        f"<div style='width:{_md_pct}%;height:100%;background:{_md_color};border-radius:2px'></div>"
                        f"</div>"
                        f"<span style='color:{_md_color};font-size:0.7rem;font-weight:600'>{_md_len}/160</span>"
                        f"</div>",
                        unsafe_allow_html=True)

                # Update parsed with edited meta fields
                if _meta_t != _parsed.get("meta_title", ""):
                    _parsed["meta_title"] = _meta_t
                    st.session_state["d_parsed"] = _parsed
                if _meta_d != _parsed.get("meta_description", ""):
                    _parsed["meta_description"] = _meta_d
                    st.session_state["d_parsed"] = _parsed

                st.divider()

                # ── Assemble final text for download ──────────────────────────────
                _final_text = _reassemble_desc(_parsed)

                # Build .txt with HTML-linked categories and tags
                _txt_parts = []
                _txt_parts.append(_talent_ln)
                _txt_parts.append(f"Title: {_d_title}")
                _txt_parts.append(f"Tags: {_tags_as_html(_d_studio, _d_tags)}")
                _txt_parts.append(f"Categories: {_cats_as_html(_d_studio, _d_cats)}")
                _txt_parts.append("")
                _txt_parts.append(_final_text)
                _full_txt = "\n".join(_txt_parts)

                # ── Save to MEGA + Download buttons ───────────────────────────────
                try:
                    _docx_bytes_out = _build_docx(_talent_ln, _d_title, _d_tags,
                                                   _d_cats, _final_text, _meta_t, _meta_d,
                                                   studio=_d_studio)
                except Exception:
                    _docx_bytes_out = None

                st.markdown("<div class='sh'>Actions</div>", unsafe_allow_html=True)
                _dl1, _dl2, _dl3 = st.columns(3)
                with _dl1:
                    _d_tab_key = {"FPVR": "FPVR", "VRH": "VRH", "VRA": "VRA", "NJOI": "NNJOI"}.get(_d_studio, _d_studio)
                    def _on_save_desc_mega(_sid=_scene_id, _fn=_filename_base, _txt=_full_txt,
                                           _docx=_docx_bytes_out, _tab_key=_d_tab_key):
                        _stage = os.path.join(os.path.dirname(__file__), "mega_staging")
                        os.makedirs(_stage, exist_ok=True)
                        _scene_dir = os.path.join(_stage, _tab_key, _sid, "Description")
                        os.makedirs(_scene_dir, exist_ok=True)
                        with open(os.path.join(_scene_dir, f"{_fn}.txt"), "w", encoding="utf-8") as _f:
                            _f.write(_txt)
                        if _docx:
                            with open(os.path.join(_scene_dir, f"{_fn}.docx"), "wb") as _f:
                                _f.write(_docx)
                        import json as _json_save2
                        _scan_path2 = os.path.join(os.path.dirname(__file__), "mega_scan.json")
                        try:
                            with open(_scan_path2, "r", encoding="utf-8") as _sf:
                                _scan2 = _json_save2.load(_sf)
                            _sid_pad2 = re.sub(r'^([A-Za-z]+)(\d+)$', lambda m: m.group(1) + m.group(2).zfill(4), _sid)
                            for _s2 in _scan2.get("scenes", []):
                                if _s2.get("scene_id") in (_sid, _sid_pad2):
                                    _s2["has_description"] = True
                                    break
                            with open(_scan_path2, "w", encoding="utf-8") as _sf:
                                _json_save2.dump(_scan2, _sf, ensure_ascii=False, indent=2)
                        except Exception:
                            pass
                        st.session_state.pop("_missing_data", None)
                        st.session_state.pop("_missing_data_ts", None)
                        st.session_state.pop("scan_data", None)
                        st.session_state["_desc_saved_mega"] = True
                    # Submit description for approval
                    if st.button("Submit for Approval", key="d_submit_approval",
                                 width="stretch", type="primary"):
                        try:
                            import json as _json_d
                            import approval_tools as _apr_d
                            _desc_content = {
                                "title": _d_title,
                                "description": _final_text,
                                "categories": _d_cats,
                                "tags": _d_tags,
                                "talent": _talent_ln,
                                "meta_title": _meta_t,
                                "meta_description": _meta_d,
                            }
                            _d_preview = f"Title: {_d_title}\n{_final_text[:180]}"
                            _d_linked = st.session_state.get("descriptions_linked_ticket", "")
                            _apr_d_id = _apr_d.submit_for_approval(
                                submitted_by=_user_name,
                                content_type="description",
                                scene_id=_scene_id,
                                studio=_d_studio,
                                content_preview=_d_preview,
                                content_json=_json_d.dumps(_desc_content),
                                target_sheet=f"Description:{_scene_id}",
                                linked_ticket=_d_linked,
                            )
                            try:
                                notification_tools.notify_approval_submitted(
                                    _apr_d_id, _scene_id, "description", _user_name)
                                _cached_unread_count.clear()
                            except Exception:
                                pass
                            st.success(f"Submitted for approval: **{_apr_d_id}**")
                        except Exception as _e:
                            st.error(f"Approval submission failed: {_e}")
                with _dl2:
                    if _docx_bytes_out:
                        st.download_button("⬇ Download .docx", data=_docx_bytes_out,
                                           file_name=f"{_filename_base}.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           width="stretch", key="d_dl_docx")
                    else:
                        st.error("docx build failed")
                with _dl3:
                    st.download_button("⬇ Download .txt", data=_full_txt.encode("utf-8"),
                                       file_name=f"{_filename_base}.txt",
                                       mime="text/plain",
                                       width="stretch", key="d_dl_txt")

    # ── TAB 7: Compilations ───────────────────────────────────────────────────────
with tab_comp:
    if _has_tab("Compilations"):
        hub_ui.section("Compilation Planner")
        import sys as _sys2, json as _json_c
        _SW_DIR = os.path.dirname(__file__)  # comp_tools.py lives alongside script_writer_app.py
        if _SW_DIR not in _sys2.path:
            _sys2.path.insert(0, _SW_DIR)

        _COMP_STUDIO_COLORS = {"FPVR": "#3b82f6", "VRH": "#8b5cf6", "VRA": "#ec4899", "NJOI": "#f97316"}

        @st.fragment
        def _comp_fragment():

            # ── Studio selector ────────────────────────────────────────────────────────
            _comp_studios = ["FPVR", "VRH", "VRA"]
            _comp_studio = st.segmented_control(
                "Studio", _comp_studios, default="VRH", key="comp_studio",
                label_visibility="collapsed"
            ) or "VRH"
            _cs_color = _COMP_STUDIO_COLORS.get(_comp_studio, _C["muted"])

            st.divider()

            # ── Load Grail scenes (Streamlit-cached, shared across all sessions) ─────
            @st.cache_data(ttl=3600, show_spinner=False)
            def _load_grail_cached(studio):
                import comp_tools as _ct
                return _ct.load_grail_scenes(studio)

            @st.cache_data(ttl=1800, show_spinner=False)
            def _load_existing_comps_cached(studio):
                import comp_tools as _ct_ec
                return _ct_ec.load_existing_comps(studio)

            # ── Two panels: Generate | Existing ───────────────────────────────────────
            _cp_gen, _cp_exist = st.tabs(["✨ Generate New Comp", "📋 Existing Comps"])

            # ── GENERATE PANEL ─────────────────────────────────────────────────────────
            with _cp_gen:

                # ── Step 1: Suggest ideas ──────────────────────────────────────────────
                st.markdown("<div class='sh'>SUGGEST IDEAS</div>", unsafe_allow_html=True)
                with st.container(border=True):
                    _suggest_col, _n_col = st.columns([3, 1])
                    with _suggest_col:
                        _suggest_btn = st.button(
                            f"Suggest new {_comp_studio} comp ideas",
                            type="primary", width="stretch", key="comp_suggest_btn"
                        )
                    with _n_col:
                        _comp_n_ideas = st.number_input("# ideas", min_value=3, max_value=10, value=6, key="comp_n_ideas")

                if _suggest_btn:
                    _claude_key_c = os.environ.get("ANTHROPIC_API_KEY", "")
                    if not _claude_key_c:
                        st.error("ANTHROPIC_API_KEY not set.")
                    else:
                        with st.spinner(f"Researching existing {_comp_studio} comps + Grail catalogue…"):
                            _grail_for_suggest = _load_grail_cached(_comp_studio)
                            _exist_for_suggest = _load_existing_comps_cached(_comp_studio)
                        if _grail_for_suggest:
                            with st.spinner("Analysing gaps and generating ideas…"):
                                try:
                                    import comp_tools as _ct_sug
                                    _ideas = _ct_sug.suggest_comp_ideas(
                                        _comp_studio, _claude_key_c,
                                        n_ideas=int(_comp_n_ideas),
                                        grail_scenes=_grail_for_suggest,
                                        existing_comps=_exist_for_suggest,
                                    )
                                    st.session_state["comp_ideas"]        = _ideas
                                    st.session_state["comp_ideas_studio"] = _comp_studio
                                except Exception as _ie:
                                    st.error(f"Idea generation failed: {_ie}")

                # ── Show idea cards ────────────────────────────────────────────────────
                if "comp_ideas" in st.session_state and st.session_state.get("comp_ideas_studio") == _comp_studio:
                    _ideas = st.session_state["comp_ideas"]
                    st.markdown("<div class='sh'>💡 Suggested Ideas — click one to build it</div>",
                                unsafe_allow_html=True)

                    _idea_cols = st.columns(3)
                    for _ii, _idea in enumerate(_ideas):
                        with _idea_cols[_ii % 3]:
                            _avail = _idea.get("available_count", 0)
                            _vol   = _idea.get("vol", 1)
                            _vol_badge = f"Vol.{_vol}" if _vol > 1 else "New"
                            _vol_color = _C["amber"] if _vol > 1 else _C["green"]
                            st.markdown(
                                f"<div class='hub-card' style='margin-bottom:8px'>"
                                f"<div style='font-weight:700;font-size:0.82rem;color:{_C['text']};margin-bottom:4px'>"
                                f"{_idea['title']}</div>"
                                f"<div style='font-size:0.68rem;color:{_C['muted']}'>"
                                f"<span style='background:{_vol_color}20;color:{_vol_color};font-size:0.6rem;"
                                f"padding:1px 6px;border-radius:8px;margin-right:4px'>{_vol_badge}</span>"
                                f"~{_avail} scenes available</div>"
                                f"<div style='font-size:0.7rem;color:{_C['muted']};margin-top:6px'>{_idea.get('rationale','')[:100]}</div>"
                                f"</div>",
                                unsafe_allow_html=True
                            )
                            if st.button("Select →", key=f"comp_idea_{_ii}", width="stretch"):
                                st.session_state["comp_theme_prefill"]    = _idea.get("theme", _idea["title"])
                                st.session_state["comp_cands_prefill"]    = _idea.get("candidate_ids", [])
                                st.session_state["comp_title_prefill"]    = _idea["title"]
                                st.session_state["comp_auto_build"]       = True
                                st.session_state.pop("comp_result", None)
                                st.rerun()

                    st.divider()

                # ── Step 2: Theme input (pre-filled if idea was selected) ─────────────
                _theme_default = st.session_state.pop("comp_theme_prefill", "")
                _cands_default = st.session_state.pop("comp_cands_prefill", [])
                _title_default = st.session_state.pop("comp_title_prefill", "")
                if _theme_default:
                    st.session_state["comp_theme"]        = _theme_default
                    st.session_state["comp_cands_hidden"]  = _cands_default
                    st.session_state["comp_title_default"] = _title_default

                st.markdown("<div class='sh'>BUILD SCENE LIST</div>", unsafe_allow_html=True)
                with st.container(border=True):
                    _cg1, _cg2 = st.columns([3, 1])
                    with _cg1:
                        _comp_theme = st.text_input(
                            "Theme / category",
                            placeholder="e.g. Cowgirl, Creampie, Blonde, Blowjob…",
                            key="comp_theme"
                        )
                    with _cg2:
                        _comp_n = st.number_input("# of scenes", min_value=5, max_value=15, value=8, key="comp_n")

                    _cg_btn = st.button("Build scene list", type="primary",
                                         width="stretch", key="comp_gen_btn")

                # Auto-build if user selected an idea card
                _auto_build = st.session_state.pop("comp_auto_build", False)
                if _cg_btn or _auto_build:
                    # On auto-build the widget may not have the prefilled value yet —
                    # fall back to session_state key set during prefill
                    _theme_to_use = _comp_theme.strip() or st.session_state.get("comp_theme", "").strip()
                    if not _theme_to_use:
                        st.error("Enter a theme first.")
                    else:
                        _claude_key_c = os.environ.get("ANTHROPIC_API_KEY", "")
                        if not _claude_key_c:
                            st.error("ANTHROPIC_API_KEY not set.")
                        else:
                            with st.spinner(f"Loading {_comp_studio} Grail data…"):
                                _grail_scenes = _load_grail_cached(_comp_studio)
                                _exist_scenes = _load_existing_comps_cached(_comp_studio)
                            if _grail_scenes:
                                with st.spinner(f"Selecting best {_comp_n} scenes for '{_theme_to_use}'…"):
                                    try:
                                        import comp_tools as _ct2
                                        _result = _ct2.generate_comp_with_ai(
                                            _comp_studio, _theme_to_use, int(_comp_n),
                                            _claude_key_c, _grail_scenes,
                                            existing_comps=_exist_scenes,
                                            preferred_ids=st.session_state.get("comp_cands_hidden", []),
                                        )
                                        if _title_def := st.session_state.get("comp_title_default"):
                                            _result["comp_title"] = _title_def
                                            st.session_state.pop("comp_title_default", None)
                                        st.session_state["comp_result"]        = _result
                                        st.session_state["comp_result_studio"] = _comp_studio
                                        st.session_state.pop("comp_cands_hidden", None)
                                    except Exception as _ce:
                                        st.error(f"Generation failed: {_ce}")

                # ── Show result ────────────────────────────────────────────────────────
                if "comp_result" in st.session_state:
                    _res = st.session_state["comp_result"]
                    _res_studio = st.session_state.get("comp_result_studio", _comp_studio)

                    st.divider()
                    _rt1, _rt2 = st.columns([4, 1])
                    with _rt1:
                        _comp_title_edit = st.text_input(
                            "Compilation title", value=_res.get("comp_title", ""),
                            key="comp_title_edit"
                        )
                    with _rt2:
                        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
                        _comp_clear = st.button("✕ Clear", key="comp_clear")
                        if _comp_clear:
                            st.session_state.pop("comp_result", None)
                            st.rerun()

                    st.markdown(
                        f"<div class='sh'>📽 SCENE LIST"
                        f"<span style='background:{_cs_color}20;color:{_cs_color};font-size:0.6rem;"
                        f"padding:1px 8px;border-radius:8px;margin-left:8px;text-transform:none;"
                        f"letter-spacing:normal;font-weight:600'>{len(_res['scenes'])} scenes</span></div>",
                        unsafe_allow_html=True)

                    # Column headers
                    _ha, _hb, _hc, _hd = st.columns([1.5, 3, 2.5, 1])
                    _ha.markdown(f"<span style='font-size:0.65rem;color:{_C['muted']};font-weight:600'>GRAIL #</span>", unsafe_allow_html=True)
                    _hb.markdown(f"<span style='font-size:0.65rem;color:{_C['muted']};font-weight:600'>SCENE TITLE</span>", unsafe_allow_html=True)
                    _hc.markdown(f"<span style='font-size:0.65rem;color:{_C['muted']};font-weight:600'>PERFORMERS</span>", unsafe_allow_html=True)
                    _hd.markdown("", unsafe_allow_html=True)

                    # Editable scene table
                    _scene_rows = _res.get("scenes", [])
                    _edited_scenes = []
                    for _si, _sc in enumerate(_scene_rows):
                        _sa, _sb, _sc_col, _sd = st.columns([1.5, 3, 2.5, 1])
                        with _sa:
                            _gid = st.text_input("Grail #", value=_sc.get("grail_id",""),
                                                  key=f"comp_gid_{_si}", label_visibility="collapsed")
                        with _sb:
                            _title = st.text_input("Title", value=_sc.get("title",""),
                                                    key=f"comp_ttl_{_si}", label_visibility="collapsed")
                        with _sc_col:
                            _cast = st.text_input("Performers", value=_sc.get("cast",""),
                                                   key=f"comp_cast_{_si}", label_visibility="collapsed")
                        with _sd:
                            _remove = st.button("✕", key=f"comp_rm_{_si}")
                        if not _remove:
                            _edited_scenes.append({
                                "grail_id": _gid, "title": _title, "cast": _cast,
                                "slr_link": _sc.get("slr_link",""),
                            })

                    st.markdown("<div class='sh'>ACTIONS</div>", unsafe_allow_html=True)

                    # Save buttons row
                    _sv1, _sv2 = st.columns(2)
                    with _sv1:
                        _save_sheet_btn = st.button("Save to Planning Sheet", type="primary",
                                                     width="stretch", key="comp_save_sheet")
                    with _sv2:
                        _save_grail_btn = st.button("Add to Grail Sheet", type="secondary",
                                                     width="stretch", key="comp_save_grail")

                    if _save_sheet_btn and _edited_scenes:
                        try:
                            import comp_tools as _ct3
                            _cell = _ct3.write_comp_to_sheet(
                                _res_studio, _comp_title_edit, _edited_scenes
                            )
                            st.success(f"✓ Written to {_res_studio} Compilations sheet starting at {_cell}")
                            st.markdown(
                                f"[Open planning sheet ↗](https://docs.google.com/spreadsheets/d/{_ct3.COMP_SHEET_ID}/edit)",
                                unsafe_allow_html=False
                            )
                        except Exception as _se:
                            st.error(f"Sheet write failed: {_se}")

                    if _save_grail_btn and _edited_scenes:
                        try:
                            import comp_tools as _ct4
                            # Enrich scenes with categories/tags from Grail data for the merge
                            _grail_map = {s["grail_id"]: s for s in _load_grail_cached(_res_studio)}
                            _enriched_for_grail = []
                            for _esc in _edited_scenes:
                                _full = _grail_map.get(_esc["grail_id"], {})
                                _enriched_for_grail.append({
                                    **_esc,
                                    "categories": _full.get("categories", ""),
                                    "tags":       _full.get("tags", ""),
                                })
                            _grail_api_key = os.environ.get("ANTHROPIC_API_KEY", "")
                            _grail_result = _ct4.write_comp_to_grail(
                                _res_studio, _comp_title_edit, _enriched_for_grail,
                                api_key=_grail_api_key,
                            )
                            _new_gid = _grail_result["grail_id"]
                            st.success(f"✓ Added to Grail as **{_new_gid}** (row {_grail_result['row_idx']})")
                            st.markdown(
                                f"<div style='display:inline-block;background:{_cs_color}18;border:1px solid {_cs_color}40;"
                                f"color:{_cs_color};font-size:0.75rem;font-weight:700;padding:2px 10px;"
                                f"border-radius:10px;margin:4px 0'>{_new_gid}</div>",
                                unsafe_allow_html=True
                            )
                            st.markdown(
                                f"[Open Grail sheet ↗](https://docs.google.com/spreadsheets/d/{_ct4.GRAIL_SHEET_ID}/edit)",
                                unsafe_allow_html=False
                            )
                            st.session_state["comp_last_grail_id"] = _new_gid
                        except Exception as _ge:
                            st.error(f"Grail write failed: {_ge}")

                    # Export for photoset builder (Mac script)
                    with st.expander("🖼 Build Photoset (Mac command)", expanded=False):
                        _scene_ids_str = ",".join(s["grail_id"] for s in _edited_scenes if s["grail_id"])
                        _comp_gid = st.session_state.get("comp_last_grail_id", "")
                        if not _comp_gid:
                            _comp_gid = f"{_res_studio}XXXX"
                        _mac_cmd = (
                            f"python3 ~/Scripts/comp_photoset.py "
                            f"--comp-id {_comp_gid} "
                            f"--scenes {_scene_ids_str} "
                            f"--output ~/Desktop/Compilations --zip"
                        )
                        st.code(_mac_cmd, language="bash")
                        if _comp_gid.endswith("XXXX"):
                            st.caption("Add to Grail first to get the real comp ID")

            # ── EXISTING COMPS PANEL ───────────────────────────────────────────────────
            with _cp_exist:
                @st.cache_data(ttl=1800, show_spinner=False)
                def _load_existing_rows(studio):
                    import gspread as _gs2
                    from google.oauth2.service_account import Credentials as _Creds2
                    _sa2 = os.path.join(os.path.dirname(__file__), "service_account.json")
                    _creds2 = _Creds2.from_service_account_file(
                        _sa2, scopes=["https://www.googleapis.com/auth/spreadsheets"])
                    _gc2 = _gs2.authorize(_creds2)
                    _sh2 = _gc2.open_by_key("1i6W4eZ8Bva3HvVmhpAVgjeHwfqbARkwBZ38aUhriaGs")
                    _tab_map = {"FPVR": "FPVR Compilations", "VRH": "VRH Compilations", "VRA": "VRA Compilations"}
                    _ws2 = _sh2.worksheet(_tab_map[studio])
                    return _ws2.get_all_values()

                # Auto-load existing comps
                _ex_needs_load = (
                    "comp_existing" not in st.session_state
                    or st.session_state.get("comp_existing_studio") != _comp_studio
                )
                if _ex_needs_load:
                    try:
                        with st.spinner(f"Loading {_comp_studio} compilations…"):
                            st.session_state["comp_existing"] = _load_existing_rows(_comp_studio)
                            st.session_state["comp_existing_studio"] = _comp_studio
                    except Exception as _ee:
                        st.error(f"Could not load sheet: {_ee}")

                _load_exist = st.button("🔄 Refresh", key="comp_load_exist", width="content")
                if _load_exist:
                    try:
                        _load_existing_rows.clear()
                        st.session_state["comp_existing"] = _load_existing_rows(_comp_studio)
                        st.session_state["comp_existing_studio"] = _comp_studio
                    except Exception as _ee:
                        st.error(f"Could not load sheet: {_ee}")

                if True:

                    _rows2 = st.session_state.get("comp_existing", [])
                    _ex_studio = st.session_state.get("comp_existing_studio", _comp_studio)

                    if _rows2:
                        # Parse: find header row (row index 1), identify comp groups
                        _hdr = _rows2[1] if len(_rows2) > 1 else []
                        # Each group: find columns where row[1] has a comp title
                        _groups = []
                        _ci2 = 0
                        while _ci2 < len(_hdr):
                            if _hdr[_ci2].strip() and _hdr[_ci2].strip() not in ("Link to Scene", "Scene Title", "Performers", "SLR Link", "MEGA Path", "Grail #"):
                                _gname = _hdr[_ci2].strip()
                                _gcol  = _ci2
                                # Collect scene entries for this group
                                _gscenes = []
                                for _row3 in _rows2[2:]:
                                    _cell_val = _row3[_gcol].strip() if _gcol < len(_row3) else ""
                                    if _cell_val:
                                        _link_val = _row3[_gcol + 1].strip() if _gcol + 1 < len(_row3) else ""
                                        _gscenes.append({"scene": _cell_val, "link": _link_val})
                                _groups.append({"title": _gname, "scenes": _gscenes, "col": _gcol})
                            _ci2 += 1

                        if not _groups:
                            st.info("No compilations found in this tab.")
                        else:
                            _sel_comp = st.selectbox(
                                "Select compilation",
                                [g["title"] for g in _groups],
                                key="comp_exist_sel"
                            )
                            _sel_group = next((g for g in _groups if g["title"] == _sel_comp), None)
                            if _sel_group:
                                _gscenes = _sel_group["scenes"]
                                st.markdown(f"**{len(_gscenes)} scenes** — enriching with Grail IDs…")

                                # Try to look up Grail IDs for each scene
                                with st.spinner("Looking up Grail IDs…"):
                                    try:
                                        import comp_tools as _ct4
                                        _grail_s2 = _load_grail_cached(_ex_studio)
                                        _enriched2 = []
                                        for _sc2 in _gscenes:
                                            _match = _ct4.find_grail_id(_sc2["scene"], _grail_s2)
                                            _enriched2.append({
                                                "scene":    _sc2["scene"],
                                                "grail_id": _match["grail_id"] if _match else "—",
                                                "cast":     _match["cast"] if _match else "",
                                                "link":     _sc2["link"],
                                                "mega":     _ct4.mega_path(_match["grail_id"]) if _match else "",
                                            })
                                    except Exception:
                                        _enriched2 = [{"scene": s["scene"], "grail_id":"—", "cast":"", "link": s["link"], "mega":""} for s in _gscenes]

                                # Display enriched table
                                for _e2 in _enriched2:
                                    _ea, _eb, _ec, _ed = st.columns([1.5, 3, 2.5, 2])
                                    with _ea:
                                        st.markdown(f"`{_e2['grail_id']}`")
                                    with _eb:
                                        st.markdown(_e2["scene"])
                                    with _ec:
                                        st.caption(_e2["cast"])
                                    with _ed:
                                        if _e2["mega"]:
                                            st.caption(_e2["mega"])

                                # Export photoset command
                                st.divider()
                                _ids_ex = ",".join(e["grail_id"] for e in _enriched2 if e["grail_id"] != "—")
                                if _ids_ex:
                                    st.markdown("<div class='sh'>🖼 PHOTOSET BUILDER</div>", unsafe_allow_html=True)
                                    st.code(
                                        f"python3 ~/Scripts/comp_photoset.py "
                                        f"--comp-id {_ex_studio}XXXX "
                                        f"--scenes {_ids_ex} "
                                        f"--output ~/Desktop/Compilations --zip",
                                        language="bash"
                                    )
                                    st.caption("Replace XXXX with the actual Grail ID after adding to Grail")

        hub_ui.section("Compilation Planner")
        _comp_fragment()

    # ── TAB 8: Tickets (Asset Tracker + Approvals + Tickets + Submit) ────────────
with tab_tickets:
    if _has_tab("Tickets"):
        import ticket_tools as _tkt
        import approval_tools as _apr
        import asset_tracker as _at

        # ── Sub-view nav ────────────────────────────────────────────────────────
        # Handle pending navigation (set before widget creation to avoid session_state conflict)
        _pending_nav = st.session_state.pop("_tk_nav_to", None)
        if _pending_nav:
            st.session_state["tk_mode"] = _pending_nav

        # Badge count — use stored count (set when Approvals sub-view loads), zero on first visit
        _pending_ct = st.session_state.get("_apr_pending_count", 0)
        _apr_label = f"Approvals ({_pending_ct})" if _pending_ct else "Approvals"

        _tk_options = ["Asset Tracker", _apr_label, "Tickets", "Submit"]
        if _user_can_manage_users:
            _tk_options.append("Users")
        _tk_mode = st.segmented_control(
            "", _tk_options,
            default="Asset Tracker", key="tk_mode", label_visibility="collapsed",
        )

        # Breadcrumb for detail views (clickable back navigation)
        _breadcrumb_detail = ""
        _breadcrumb_back_key = ""
        if _tk_mode == "Asset Tracker" and st.session_state.get("at_selected"):
            _breadcrumb_detail = st.session_state["at_selected"]
            _breadcrumb_back_key = "at_selected"
        elif _tk_mode == "Tickets" and st.session_state.get("tk_selected"):
            _breadcrumb_detail = st.session_state["tk_selected"]
            _breadcrumb_back_key = "tk_selected"
        _mode_name = _tk_mode.split(" (")[0] if _tk_mode and "(" in _tk_mode else (_tk_mode or "")
        if _breadcrumb_detail:
            _bc1, _bc2 = st.columns([6, 1])
            with _bc1:
                st.markdown(
                    f"<div style='font-size:0.78rem;color:{_C['muted']};padding:4px 0'>"
                    f"Tickets &rsaquo; {_mode_name} &rsaquo; "
                    f"<span style='color:{_C['text']};font-weight:500'>{_breadcrumb_detail}</span></div>",
                    unsafe_allow_html=True,
                )
            with _bc2:
                if st.button("Back", key="breadcrumb_back"):
                    st.session_state.pop(_breadcrumb_back_key, None)
                    st.rerun()

        # ── SUB-VIEW 1: Asset Tracker ────────────────────────────────────────────
        if _tk_mode == "Asset Tracker":
            hub_ui.section("Asset Tracker")

            # ── Naming convention spec (from Grail Naming Scheme doc) ────────
            _NAMING_SPEC = {
                "folder": "{PREFIX}{NUM:04d}",          # e.g. FPVR0001
                "description_doc": "{PREFIX}{NUM:04d}-{Names}.doc",
                "description_txt": "{PREFIX}{NUM:04d}-{Names}.txt",
                "thumbnail": "{PREFIX}{NUM:04d}-{Names}_Thumbnail.{ext}",
                "photos_raw": "_{PREFIX}{NUM:04d}.zip",   # underscore prefix
                "photos_web": "{PREFIX}{NUM:04d}.zip",
                "video_full": "{Names}-180-{PREFIX}_{res}",
                "video_trailer_2min": "{Names}-180-{PREFIX}-2min_{res}",
            }

            def _validate_naming(_scene):
                """Check file names against the naming spec. Returns list of issues."""
                issues = []
                sid = _scene["scene_id"]
                prefix = _scene["studio"]
                # Normalize NJOI → NJOI for mega scan paths
                if prefix == "NNJOI":
                    prefix = "NJOI"
                num_str = _scene["scene_num"]
                num_padded = num_str.zfill(4)
                expected_prefix = f"{prefix}{num_padded}"
                mega_files = _scene.get("mega_files", {})

                # Check thumbnail naming
                for f in mega_files.get("thumbnail", []):
                    fname = f.split("/")[-1]
                    if not fname.startswith(expected_prefix):
                        issues.append(("Thumbnail", f, f"Should start with {expected_prefix}"))
                    if "_Thumbnail" not in fname:
                        issues.append(("Thumbnail", f, "Missing _Thumbnail suffix"))

                # Check description naming
                for f in mega_files.get("description", []):
                    fname = f.split("/")[-1]
                    if not fname.startswith(expected_prefix):
                        issues.append(("Description", f, f"Should start with {expected_prefix}"))

                # Check video naming — should contain model names + 180 + studio prefix
                for f in mega_files.get("videos", []):
                    fname = f.split("/")[-1]
                    if f"-{prefix}" not in fname and f"-{prefix.upper()}" not in fname:
                        issues.append(("Video", fname[:50], f"Should contain -{prefix}"))

                return issues

            # Studio color map
            _studio_colors = {"FPVR": _C["fpvr"], "VRH": _C["vrh"], "VRA": _C["vra"], "NNJOI": _C["njoi"]}

            # Load asset data — always needed for both grid and detail
            _at_studio = st.session_state.get("at_studio", "All")
            _at_show = st.session_state.get("at_show", "All")
            _at_limit = 6
            _at_studios_tuple = None if _at_studio == "All" else tuple([_at_studio])
            _asset_data = _cached_load_assets(_at_studios_tuple, _at_limit)
            _scan_date = _asset_data.get("_meta", {}).get("scan_date", "")
            _at_studios = ["FPVR", "VRH", "VRA", "NNJOI"] if _at_studio == "All" else [_at_studio]
            _at_by_studio = {}
            _at_scenes = []
            for _s in _at_studios:
                _studio_scenes = list(_asset_data.get(_s, []))
                if _at_show == "Missing Only":
                    _studio_scenes = [s for s in _studio_scenes if s["completed"] < s["total"]]
                elif _at_show == "Complete":
                    _studio_scenes = [s for s in _studio_scenes if s["completed"] == s["total"]]
                _studio_scenes.sort(key=lambda s: s.get("release_date", "") or "", reverse=True)
                _at_by_studio[_s] = _studio_scenes
                _at_scenes.extend(_studio_scenes)

            _selected_scene_id = st.session_state.get("at_selected")
            _in_detail = _selected_scene_id is not None

            # ── Filters + metrics (hidden when viewing a single scene) ──
            if not _in_detail:
                _at_f1, _at_f2 = st.columns(2)
                with _at_f1:
                    st.selectbox(
                        "Studio", ["All", "FPVR", "VRH", "VRA", "NNJOI"], key="at_studio")
                with _at_f2:
                    st.selectbox(
                        "Show", ["All", "Missing Only", "Complete"], key="at_show")

                if st.button("Refresh", key="at_refresh"):
                    _at.bust_caches()
                    _cached_load_assets.clear()
                    st.rerun()

                _at_total = len(_at_scenes)
                _at_complete = sum(1 for s in _at_scenes if s["completed"] == s["total"])
                _at_partial = _at_total - _at_complete
                _m1, _m2, _m3 = st.columns(3)
                _m1.metric("Total Scenes", _at_total)
                _m2.metric("Complete", _at_complete)
                _m3.metric("In Progress", _at_partial)
                if _scan_date:
                    st.caption(f"MEGA scan: {_scan_date}")

                _issue_scenes = [s for s in _at_scenes if s["completed"] < s["total"]]
                if _issue_scenes:
                    _issue_scenes.sort(key=lambda s: len(s["missing"]), reverse=True)
                    with st.expander(f"Scenes with Issues ({len(_issue_scenes)})", expanded=False):
                        for _is in _issue_scenes:
                            _is_color = _studio_colors.get(_is["studio"], _C["muted"])
                            _is_missing_pills = " ".join(
                                f"<span style='display:inline-block;font-size:0.65rem;padding:1px 6px;"
                                f"border-radius:3px;background:{_C['red_dim']};color:{_C['red']}'>{m}</span>"
                                for m in _is["missing"]
                            )
                            _ic1, _ic2 = st.columns([3, 1])
                            with _ic1:
                                st.markdown(
                                    f"<div style='display:flex;align-items:center;gap:8px;padding:4px 0;flex-wrap:wrap'>"
                                    f"<span style='width:3px;height:16px;border-radius:2px;"
                                    f"background:{_is_color};display:inline-block'></span>"
                                    f"<span style='font-family:DM Mono,monospace;font-size:0.78rem;"
                                    f"font-weight:600;color:{_C['text']}'>{_is['scene_id']}</span>"
                                    f"{_is_missing_pills}"
                                    f"</div>",
                                    unsafe_allow_html=True,
                                )
                            with _ic2:
                                if st.button("View", key=f"issue_{_is['scene_id']}", width="stretch"):
                                    st.session_state["at_selected"] = _is["scene_id"]
                                    st.rerun()

                st.divider()

            # ── Detail view (selected scene) ─────────────────────────────────
            _selected_scene = next((s for s in _at_scenes if s["scene_id"] == _selected_scene_id), None) if _selected_scene_id else None

            if _selected_scene:
                _sc = _selected_scene
                _s_color = _studio_colors.get(_sc["studio"], _C["muted"])
                _pct = int((_sc["completed"] / _sc["total"]) * 100) if _sc["total"] else 0
                _pct_color = _C["green"] if _pct == 100 else (_C["amber"] if _pct >= 50 else _C["red"])

                # Scene header
                _comp_badge = (
                    f"<span style='font-size:0.68rem;background:{_C['accent']};color:#fff;"
                    f"padding:2px 8px;border-radius:4px;font-weight:600'>Compilation</span>"
                ) if _sc.get("is_compilation") else ""
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:12px;margin:8px 0 4px'>"
                    f"<div style='width:4px;height:28px;border-radius:2px;background:{_s_color}'></div>"
                    f"<span style='font-family:DM Mono,monospace;font-size:1.1rem;font-weight:700;"
                    f"color:{_C['text']}'>{_sc['scene_id']}</span>"
                    f"<span style='background:{_s_color}22;color:{_s_color};font-size:0.72rem;"
                    f"font-weight:600;padding:3px 8px;border-radius:4px'>{_sc['studio_name']}</span>"
                    f"{_comp_badge}"
                    f"<span style='font-size:0.82rem;color:{_C['muted']};margin-left:auto'>"
                    f"{_sc.get('release_date','')[:10] or 'No date'}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # Info row
                st.markdown(
                    f"<div style='margin:4px 0 12px 16px;font-size:0.85rem'>"
                    f"<b style='color:{_C['text']}'>{_sc['performers'] or 'TBD'}</b>"
                    f"<span style='color:{_C['muted']};margin-left:12px;font-style:italic'>"
                    f"{_sc['title'] or 'No title'}</span></div>",
                    unsafe_allow_html=True,
                )

                # Progress bar (larger)
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:10px;margin:0 0 16px 0'>"
                    f"<div style='flex:1;height:10px;background:{_C['elevated']};border-radius:5px;overflow:hidden'>"
                    f"<div style='width:{_pct}%;height:100%;background:{_pct_color};border-radius:5px'></div>"
                    f"</div>"
                    f"<span style='font-size:0.85rem;font-weight:700;color:{_pct_color}'>"
                    f"{_sc['completed']}/{_sc['total']} assets</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # ── Asset grid (interactive for editable fields) ────────────
                # Read-only asset status for non-editable items
                _readonly_assets = [
                    ("has_videos", "Videos", "Video files in Videos/"),
                    ("has_thumbnail", "Thumbnail", "In Video Thumbnail/"),
                    ("has_photos", "Photos", "Photo .zip in Photos/"),
                ]
                _ro1, _ro2, _ro3 = st.columns(3)
                for _col_ro, (_ak, _al, _desc) in zip([_ro1, _ro2, _ro3], _readonly_assets):
                    _ok = _sc.get(_ak, False)
                    _icon = f"<span style='color:{_C['green']};font-size:1.1rem'>&#10003;</span>" if _ok else f"<span style='color:{_C['red']};font-size:1.1rem'>&#10007;</span>"
                    _col_ro.markdown(
                        f"<div style='background:{_C['elevated']};border-radius:6px;padding:8px 12px;"
                        f"margin-bottom:6px;display:flex;align-items:center;gap:8px'>"
                        f"{_icon}"
                        f"<div><span style='font-size:0.82rem;font-weight:600;color:{_C['text']}'>{_al}</span>"
                        f"<br><span style='font-size:0.68rem;color:{_C['muted']}'>{_desc}</span></div>"
                        f"</div>",
                        unsafe_allow_html=True,
                    )

                # ── Editable fields: Title, Categories, Tags, Description ────
                _grail_tab = _sc.get("grail_tab", "")
                _grail_row = _sc.get("grail_row", 0)
                _sc_studio = _sc.get("studio", "")
                # Map NNJOI → NJOI for category/tag lookups
                _cat_studio = "NJOI" if _sc_studio == "NNJOI" else _sc_studio

                # — Title —
                hub_ui.section("Title")
                _cur_title = _sc.get("title", "")
                _editing_title = st.session_state.get("at_editing_title", False)

                if _editing_title and _user_can_write_grail:
                    _new_title = st.text_input(
                        "Title", value=_cur_title, key="at_title_input",
                        placeholder="Enter scene title")
                    _tc1, _tc2 = st.columns(2)
                    with _tc1:
                        if st.button("Save Title", key="at_title_save", width="stretch"):
                            if _new_title.strip() and _grail_tab and _grail_row:
                                _ok_t, _msg_t = _write_grail_cell(_grail_tab, _grail_row, 4, _new_title.strip())
                                if _ok_t:
                                    st.session_state.pop("at_editing_title", None)
                                    _cached_load_assets.clear()
                                    st.rerun()
                                else:
                                    st.error(_msg_t)
                            else:
                                st.warning("Title cannot be empty.")
                    with _tc2:
                        if st.button("Cancel", key="at_title_cancel", width="stretch"):
                            st.session_state.pop("at_editing_title", None)
                            st.rerun()
                else:
                    _title_icon = f"<span style='color:{_C['green']}'>&#10003;</span>" if _cur_title else f"<span style='color:{_C['red']}'>&#10007;</span>"
                    st.markdown(
                        f"<div style='display:flex;align-items:center;gap:8px;padding:4px 0'>"
                        f"{_title_icon} <span style='font-size:0.85rem;color:{_C['text']}'>"
                        f"{_cur_title or 'No title'}</span></div>",
                        unsafe_allow_html=True,
                    )
                    if _user_can_write_grail:
                        _title_btn_label = "Edit Title" if _cur_title else "Add Title"
                        if st.button(_title_btn_label, key="at_title_edit"):
                            st.session_state["at_editing_title"] = True
                            st.session_state.pop("at_editing_cats", None)
                            st.session_state.pop("at_editing_tags", None)
                            st.rerun()

                # — Categories —
                hub_ui.section("Categories")
                _cur_cats_raw = ""
                # Categories are in Grail col F — reconstruct from scene data
                # asset_tracker doesn't store raw cats, but we can get from the Grail row
                # For now, use has_categories + look for cats in asset_data
                # Actually the Grail loader stores cats at r[5] — check scene dict
                _editing_cats = st.session_state.get("at_editing_cats", False)
                _approved_cats = _STUDIO_CATEGORIES.get(_cat_studio, [])

                if _editing_cats and _user_can_write_grail:
                    # Need to load current cats from Grail for pre-fill
                    _prefill_cats = [c.strip() for c in st.session_state.get("at_cats_prefill", "").split(",") if c.strip()]
                    if _approved_cats:
                        _new_cats = st.multiselect(
                            "Categories", options=_approved_cats,
                            default=[c for c in _prefill_cats if c in _approved_cats],
                            key="at_cats_input")
                    else:
                        _new_cats_str = st.text_area(
                            "Categories (comma-separated)",
                            value=st.session_state.get("at_cats_prefill", ""),
                            key="at_cats_input_text")
                        _new_cats = [c.strip() for c in _new_cats_str.split(",") if c.strip()]
                    _cc1, _cc2 = st.columns(2)
                    with _cc1:
                        if st.button("Save Categories", key="at_cats_save", width="stretch"):
                            _cats_str = ", ".join(_new_cats)
                            if _grail_tab and _grail_row:
                                _ok_c, _msg_c = _write_grail_cell(_grail_tab, _grail_row, 6, _cats_str)
                                if _ok_c:
                                    st.session_state.pop("at_editing_cats", None)
                                    st.session_state.pop("at_cats_prefill", None)
                                    _cached_load_assets.clear()
                                    st.rerun()
                                else:
                                    st.error(_msg_c)
                    with _cc2:
                        if st.button("Cancel", key="at_cats_cancel", width="stretch"):
                            st.session_state.pop("at_editing_cats", None)
                            st.session_state.pop("at_cats_prefill", None)
                            st.rerun()
                else:
                    _has_cats = _sc.get("has_categories", False)
                    _cats_raw = _sc.get("categories_raw", "")
                    _cats_icon = f"<span style='color:{_C['green']}'>&#10003;</span>" if _has_cats else f"<span style='color:{_C['red']}'>&#10007;</span>"
                    _cats_display = _cats_raw if _cats_raw else "No categories"
                    st.markdown(
                        f"<div style='display:flex;align-items:center;gap:8px;padding:4px 0;flex-wrap:wrap'>"
                        f"{_cats_icon} <span style='font-size:0.85rem;color:{_C['text']}'>"
                        f"{_cats_display}</span></div>",
                        unsafe_allow_html=True,
                    )
                    if _user_can_write_grail:
                        _cats_btn_label = "Edit Categories" if _has_cats else "Add Categories"
                        if st.button(_cats_btn_label, key="at_cats_edit"):
                            # Load current cats from Grail for pre-fill
                            try:
                                import gspread as _gs_cats
                                from google.oauth2.service_account import Credentials as _Creds_cats
                                _creds_c = _Creds_cats.from_service_account_file(
                                    os.path.join(os.path.dirname(__file__), "service_account.json"),
                                    scopes=["https://www.googleapis.com/auth/spreadsheets"])
                                _gc_c = _gs_cats.authorize(_creds_c)
                                _ws_c = _gc_c.open_by_key("1Eq5G5FU6A8EqeFZCnZjrEaMYS8F1DiK5vP5tCSINeJk").worksheet(_grail_tab)
                                _cur_cats_val = _ws_c.cell(_grail_row, 6).value or ""
                                st.session_state["at_cats_prefill"] = _cur_cats_val
                            except Exception:
                                st.session_state["at_cats_prefill"] = ""
                            st.session_state["at_editing_cats"] = True
                            st.session_state.pop("at_editing_title", None)
                            st.session_state.pop("at_editing_tags", None)
                            st.rerun()

                # — Tags —
                hub_ui.section("Tags")
                _editing_tags = st.session_state.get("at_editing_tags", False)
                _approved_tags_ref = _STUDIO_TAGS.get(_cat_studio, "")

                if _editing_tags and _user_can_write_grail:
                    _prefill_tags = st.session_state.get("at_tags_prefill", "")
                    _new_tags = st.text_area(
                        "Tags (comma-separated)", value=_prefill_tags,
                        key="at_tags_input", height=100)
                    if _approved_tags_ref:
                        with st.expander(f"Approved {_cat_studio} Tags Reference", expanded=False):
                            st.caption(_approved_tags_ref)
                    _tc1t, _tc2t = st.columns(2)
                    with _tc1t:
                        if st.button("Save Tags", key="at_tags_save", width="stretch"):
                            if _grail_tab and _grail_row:
                                _ok_tg, _msg_tg = _write_grail_cell(_grail_tab, _grail_row, 7, _new_tags.strip())
                                if _ok_tg:
                                    st.session_state.pop("at_editing_tags", None)
                                    st.session_state.pop("at_tags_prefill", None)
                                    _cached_load_assets.clear()
                                    st.rerun()
                                else:
                                    st.error(_msg_tg)
                    with _tc2t:
                        if st.button("Cancel", key="at_tags_cancel", width="stretch"):
                            st.session_state.pop("at_editing_tags", None)
                            st.session_state.pop("at_tags_prefill", None)
                            st.rerun()
                else:
                    _has_tags = _sc.get("has_tags", False)
                    _tags_raw = _sc.get("tags_raw", "")
                    _tags_icon = f"<span style='color:{_C['green']}'>&#10003;</span>" if _has_tags else f"<span style='color:{_C['red']}'>&#10007;</span>"
                    _tags_display = _tags_raw if _tags_raw else "No tags"
                    st.markdown(
                        f"<div style='display:flex;align-items:center;gap:8px;padding:4px 0;flex-wrap:wrap'>"
                        f"{_tags_icon} <span style='font-size:0.85rem;color:{_C['text']}'>"
                        f"{_tags_display}</span></div>",
                        unsafe_allow_html=True,
                    )
                    if _user_can_write_grail:
                        _tags_btn_label = "Edit Tags" if _has_tags else "Add Tags"
                        if st.button(_tags_btn_label, key="at_tags_edit"):
                            try:
                                import gspread as _gs_tags
                                from google.oauth2.service_account import Credentials as _Creds_tags
                                _creds_t = _Creds_tags.from_service_account_file(
                                    os.path.join(os.path.dirname(__file__), "service_account.json"),
                                    scopes=["https://www.googleapis.com/auth/spreadsheets"])
                                _gc_t = _gs_tags.authorize(_creds_t)
                                _ws_t = _gc_t.open_by_key("1Eq5G5FU6A8EqeFZCnZjrEaMYS8F1DiK5vP5tCSINeJk").worksheet(_grail_tab)
                                _cur_tags_val = _ws_t.cell(_grail_row, 7).value or ""
                                st.session_state["at_tags_prefill"] = _cur_tags_val
                            except Exception:
                                st.session_state["at_tags_prefill"] = ""
                            st.session_state["at_editing_tags"] = True
                            st.session_state.pop("at_editing_title", None)
                            st.session_state.pop("at_editing_cats", None)
                            st.rerun()

                # — Description —
                hub_ui.section("Description")
                _has_desc = _sc.get("has_description", False)
                _desc_files = _sc.get("mega_files", {}).get("description", [])
                _desc_icon = f"<span style='color:{_C['green']}'>&#10003;</span>" if _has_desc else f"<span style='color:{_C['red']}'>&#10007;</span>"

                if _has_desc and _desc_files:
                    _desc_names = ", ".join(f.split("/")[-1] for f in _desc_files)
                    st.markdown(
                        f"<div style='display:flex;align-items:center;gap:8px;padding:4px 0'>"
                        f"{_desc_icon} <span style='font-size:0.85rem;color:{_C['text']}'>"
                        f"{_desc_names}</span></div>",
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f"<div style='display:flex;align-items:center;gap:8px;padding:4px 0'>"
                        f"{_desc_icon} <span style='font-size:0.85rem;color:{_C['text']}'>"
                        f"No description file</span></div>",
                        unsafe_allow_html=True,
                    )
                    # Pre-fill Descriptions tab and prompt user to switch
                    if st.button("Create in Descriptions Tab", key="at_desc_create"):
                        _desc_studio_key = "NJOI" if _sc_studio == "NNJOI" else _sc_studio
                        st.session_state["d_studio"] = _desc_studio_key
                        _sn = _sc.get("scene_num", "")
                        st.session_state["d_scene_num"] = int(_sn) if _sn.isdigit() else 1
                        _female = _sc.get("female", "")
                        if _female:
                            st.session_state["d_female"] = _female
                        st.session_state["desc_active_scene"] = _sc["scene_id"]
                        st.info("Scene pre-filled in the Descriptions tab. Click the Descriptions tab to continue.")

                # ── File listings ─────────────────────────────────────────────
                _mega_files = _sc.get("mega_files", {})
                _has_any_files = any(_mega_files.get(k) for k in _mega_files)

                if _has_any_files:
                    st.divider()
                    hub_ui.section("MEGA Files")

                    # Thumbnail files
                    _thumbs = _mega_files.get("thumbnail", [])
                    if _thumbs:
                        st.markdown(f"**Video Thumbnail** ({len(_thumbs)} file{'s' if len(_thumbs)>1 else ''})")
                        for _tf in _thumbs:
                            _fname = _tf.split("/")[-1] if "/" in _tf else _tf
                            st.markdown(
                                f"<div style='font-family:DM Mono,monospace;font-size:0.75rem;"
                                f"color:{_C['text']};background:{_C['elevated']};padding:4px 8px;"
                                f"border-radius:4px;margin:2px 0'>{_fname}</div>",
                                unsafe_allow_html=True,
                            )

                    # Video files
                    _vids = _mega_files.get("videos", [])
                    if _vids:
                        st.markdown(f"**Videos** ({len(_vids)} file{'s' if len(_vids)>1 else ''})")
                        # Group by type: full, 2min, 45sec, 6min, rollover
                        _vid_groups = {"Full": [], "2min": [], "45sec": [], "6min": [], "Rollover": [], "Other": []}
                        for _vf in _vids:
                            _vn = _vf.split("/")[-1] if "/" in _vf else _vf
                            if "-2min" in _vn: _vid_groups["2min"].append(_vn)
                            elif "-45sec" in _vn: _vid_groups["45sec"].append(_vn)
                            elif "-6min" in _vn: _vid_groups["6min"].append(_vn)
                            elif "rollover" in _vn.lower(): _vid_groups["Rollover"].append(_vn)
                            else: _vid_groups["Full"].append(_vn)
                        for _gname, _gfiles in _vid_groups.items():
                            if _gfiles:
                                st.caption(f"{_gname} ({len(_gfiles)})")
                                for _vn in sorted(_gfiles):
                                    st.markdown(
                                        f"<div style='font-family:DM Mono,monospace;font-size:0.7rem;"
                                        f"color:{_C['text']};padding:1px 8px'>{_vn}</div>",
                                        unsafe_allow_html=True,
                                    )

                    # Description files
                    _descs = _mega_files.get("description", [])
                    if _descs:
                        st.markdown(f"**Description** ({len(_descs)} file{'s' if len(_descs)>1 else ''})")
                        for _df in _descs:
                            _fname = _df.split("/")[-1] if "/" in _df else _df
                            st.markdown(
                                f"<div style='font-family:DM Mono,monospace;font-size:0.75rem;"
                                f"color:{_C['text']};padding:1px 8px'>{_fname}</div>",
                                unsafe_allow_html=True,
                            )

                    # Photos + storyboard
                    _photos = _mega_files.get("photos", [])
                    _story = _mega_files.get("storyboard", [])
                    if _photos:
                        st.markdown(f"**Photos** ({len(_photos)} file{'s' if len(_photos)>1 else ''})")
                        for _pf in _photos:
                            _fname = _pf.split("/")[-1] if "/" in _pf else _pf
                            st.markdown(
                                f"<div style='font-family:DM Mono,monospace;font-size:0.75rem;"
                                f"color:{_C['text']};padding:1px 8px'>{_fname}</div>",
                                unsafe_allow_html=True,
                            )
                    if _story:
                        st.markdown(f"**Storyboard** ({len(_story)} file{'s' if len(_story)>1 else ''})")

                elif not _has_any_files:
                    st.info("No MEGA files found for this scene. Run the MEGA scanner to update.")

                # ── Naming validation ─────────────────────────────────────────
                _naming_issues = _validate_naming(_sc)
                if _naming_issues:
                    st.divider()
                    hub_ui.section("Naming Issues")
                    for _cat, _file, _issue in _naming_issues:
                        st.markdown(
                            f"<div style='background:{_C['red_dim']};border-radius:4px;padding:6px 10px;"
                            f"margin:3px 0;font-size:0.78rem'>"
                            f"<b style='color:{_C['red']}'>{_cat}</b> "
                            f"<code style='font-size:0.72rem'>{_file}</code>"
                            f"<br><span style='color:{_C['muted']}'>{_issue}</span></div>",
                            unsafe_allow_html=True,
                        )
                elif _has_any_files:
                    st.success("All file names match naming conventions.")

                # ── Actions ───────────────────────────────────────────────────
                st.divider()
                if _sc["missing"]:
                    if st.button("Create Ticket for Missing Assets", key="at_detail_tkt", width="stretch"):
                        _missing_str = ", ".join(_sc["missing"])
                        st.session_state["_tk_nav_to"] = "Submit"
                        st.session_state["tk_prefill"] = {
                            "project": "Content Pipeline",
                            "type": "Missing Content",
                            "title": f"Missing assets: {_sc['scene_id']}",
                            "desc": f"Scene {_sc['scene_id']} ({_sc['studio_name']}) — {_sc['performers']}\n\nMissing: {_missing_str}",
                        }
                        st.rerun()
                if st.button("Report Issue", key="at_detail_report", width="stretch"):
                    st.session_state["_tk_nav_to"] = "Submit"
                    st.session_state["tk_prefill"] = {
                        "project": "Content Pipeline",
                        "type": "Bug",
                        "title": f"Issue with {_sc['scene_id']}",
                        "desc": f"Scene {_sc['scene_id']} ({_sc['studio_name']}) — {_sc['performers']}\n\nDescribe the issue:\n",
                    }
                    st.rerun()

            # ── Grid view grouped by studio ──────────────────────────────────
            else:
                if not _at_scenes:
                    st.info("No scenes found matching your filters.")
                else:
                    _studio_names = {"FPVR": "FuckPassVR", "VRH": "VRHush", "VRA": "VRAllure", "NNJOI": "NaughtyJOI"}
                    for _grp_studio in _at_studios:
                        _grp_scenes = _at_by_studio.get(_grp_studio, [])
                        if not _grp_scenes:
                            continue
                        _s_color = _studio_colors.get(_grp_studio, _C["muted"])
                        st.markdown(
                            f"<div style='display:flex;align-items:center;gap:10px;margin:16px 0 8px'>"
                            f"<div style='width:4px;height:20px;border-radius:2px;background:{_s_color}'></div>"
                            f"<span style='font-family:Syne,sans-serif;font-size:1rem;font-weight:700;"
                            f"color:{_C['text']}'>{_studio_names.get(_grp_studio, _grp_studio)}</span>"
                            f"<span style='font-size:0.72rem;color:{_C['muted']}'>{len(_grp_scenes)} scenes</span>"
                            f"</div>",
                            unsafe_allow_html=True,
                        )
                        for _ri in range(0, len(_grp_scenes), 3):
                            _row_scenes = _grp_scenes[_ri:_ri + 3]
                            _cols = st.columns(3)
                            for _ci, _scene in enumerate(_row_scenes):
                                with _cols[_ci]:
                                    _pct = int((_scene["completed"] / _scene["total"]) * 100) if _scene["total"] else 0
                                    _pct_color = _C["green"] if _pct == 100 else (_C["amber"] if _pct >= 50 else _C["red"])

                                    _missing_names = _scene.get("missing", [])
                                    if _missing_names:
                                        _checks_html = " ".join(
                                            f"<span style='display:inline-block;font-size:0.62rem;padding:1px 5px;"
                                            f"border-radius:3px;background:{_C['red_dim']};color:{_C['red']}'>{m}</span>"
                                            for m in _missing_names
                                        )
                                    else:
                                        _checks_html = (
                                            f"<span style='font-size:0.72rem;color:{_C['green']}'>"
                                            f"&#10003; All assets complete</span>"
                                        )

                                    _perf_display = _scene["performers"] or "TBD"
                                    _title_display = _scene["title"][:40] + "..." if len(_scene.get("title", "")) > 40 else (_scene.get("title") or "\u2014")
                                    _date_display = _scene.get("release_date", "")[:10] or ""
                                    _card_comp = (
                                        f"<span style='font-size:0.6rem;background:{_C['accent']};color:#fff;"
                                        f"padding:1px 6px;border-radius:3px;margin-left:6px;font-family:inherit'>COMP</span>"
                                    ) if _scene.get("is_compilation") else ""

                                    _date_html = f"<span style='font-size:0.68rem;color:{_C['subtle']}'>{_date_display}</span>" if _date_display else ""
                                    st.markdown(
                                        f"<div style='background:{_C['surface']};border:1px solid {_C['border']};"
                                        f"border-left:3px solid {_s_color};border-radius:8px;padding:14px 16px;"
                                        f"min-height:200px;display:flex;flex-direction:column'>"
                                        f"<div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:6px'>"
                                        f"<span style='font-family:DM Mono,monospace;font-size:0.82rem;font-weight:600;"
                                        f"color:{_C['text']}'>{_scene['scene_id']}{_card_comp}</span>"
                                        f"{_date_html}"
                                        f"</div>"
                                        f"<div style='font-size:0.82rem;color:{_C['text']};font-weight:500;"
                                        f"margin-bottom:2px;overflow:hidden;display:-webkit-box;"
                                        f"-webkit-line-clamp:2;-webkit-box-orient:vertical'>{_perf_display}</div>"
                                        f"<div style='font-size:0.72rem;color:{_C['muted']};margin-bottom:auto;"
                                        f"font-style:italic'>{_title_display}</div>"
                                        f"<div style='display:flex;align-items:center;gap:8px;margin:8px 0 6px'>"
                                        f"<div style='flex:1;height:6px;background:{_C['elevated']};border-radius:3px;overflow:hidden'>"
                                        f"<div style='width:{_pct}%;height:100%;background:{_pct_color};border-radius:3px'></div>"
                                        f"</div>"
                                        f"<span style='font-size:0.72rem;font-weight:600;color:{_pct_color}'>"
                                        f"{_scene['completed']}/{_scene['total']}</span>"
                                        f"</div>"
                                        f"<div style='display:flex;flex-wrap:wrap;gap:3px'>"
                                        f"{_checks_html}</div>"
                                        f"</div>",
                                        unsafe_allow_html=True,
                                    )
                                    if st.button("Open", key=f"at_open_{_scene['scene_id']}", width="stretch"):
                                        st.session_state["at_selected"] = _scene["scene_id"]
                                        st.rerun()

        # ── SUB-VIEW 2: Approvals ────────────────────────────────────────────────
        elif _tk_mode == _apr_label:
            hub_ui.section("Approval Queue")

            # Load approvals (globally cached)
            _all_approvals = _cached_load_approvals()
            st.session_state["_apr_pending_count"] = sum(1 for a in _all_approvals if a["status"] == "Pending")

            # Refresh button
            if st.button("Refresh", key="apr_refresh"):
                _cached_load_approvals.clear()
                st.rerun()

            # Filter controls
            _ap_f1, _ap_f2, _ap_f3 = st.columns(3)
            with _ap_f1:
                _ap_filt = st.selectbox("Status", ["Pending", "All", "Approved", "Rejected", "Superseded"],
                                        key="ap_filt_status")
            with _ap_f2:
                _ap_type = st.selectbox("Content Type", ["All"] + _apr.CONTENT_TYPES, key="ap_filt_type")
            with _ap_f3:
                _ap_submitter = st.selectbox(
                    "Submitter", ["All", "Mine"],
                    index=0 if _user_is_admin else 1,
                    key="ap_filt_submitter",
                )

            _filt_approvals = _all_approvals
            if _ap_filt != "All":
                _filt_approvals = [a for a in _filt_approvals if a["status"] == _ap_filt]
            if _ap_type != "All":
                _filt_approvals = [a for a in _filt_approvals if a["content_type"] == _ap_type]
            if _ap_submitter == "Mine":
                _filt_approvals = [a for a in _filt_approvals if a["submitted_by"] == _user_name]

            # Summary
            _ap_pending = sum(1 for a in _all_approvals if a["status"] == "Pending")
            _ap_approved = sum(1 for a in _all_approvals if a["status"] == "Approved")
            _ap_rejected = sum(1 for a in _all_approvals if a["status"] == "Rejected")
            _apm1, _apm2, _apm3 = st.columns(3)
            _apm1.metric("Pending", _ap_pending)
            _apm2.metric("Approved", _ap_approved)
            _apm3.metric("Rejected", _ap_rejected)

            st.divider()

            _type_colors = {
                "script": _C["blue"], "description": _C["green"],
                "title_text": _C["amber"], "title_card": _C["vra"],
            }

            if not _filt_approvals:
                st.info("No approvals found matching filters.")
            else:
                for _appr in reversed(_filt_approvals):
                    _tc = _type_colors.get(_appr["content_type"], _C["muted"])
                    _is_pending = _appr["status"] == "Pending"
                    _status_badge = hub_ui.badge(_appr["status"], {
                        "Pending": "amber", "Approved": "green",
                        "Rejected": "red", "Superseded": "gray",
                    }.get(_appr["status"], "gray"))
                    _type_badge = hub_ui.badge(_appr["content_type"], "blue")
                    _linked = f" &middot; {_appr['linked_ticket']}" if _appr.get("linked_ticket") else ""

                    with st.expander(
                        f"{_appr['id']} — {_appr['content_type']} — {_appr['scene_id']} ({_appr['studio']})",
                        expanded=_is_pending,
                    ):
                        st.markdown(
                            f"<div style='display:flex;align-items:center;gap:8px;margin-bottom:8px'>"
                            f"{_status_badge} {_type_badge}"
                            f"<span style='font-size:0.75rem;color:{_C['muted']}'>"
                            f"by {_appr['submitted_by']} &middot; {_appr['date']}{_linked}</span>"
                            f"</div>",
                            unsafe_allow_html=True,
                        )
                        # Rejection feedback (prominent, before preview)
                        if _appr["status"] == "Rejected" and _appr["notes"]:
                            hub_ui.card(
                                f"<div style='font-size:0.82rem;color:{_C['red']}'>"
                                f"<b>Rejection Reason:</b> {_appr['notes']}</div>",
                                accent=_C["red"],
                            )

                        # Content preview
                        hub_ui.card(
                            f"<div style='font-size:0.82rem;color:{_C['text']};white-space:pre-wrap;"
                            f"line-height:1.6'>{_appr['preview']}</div>",
                        )
                        # Truncation indicator
                        if len(_appr.get("preview", "")) >= 200 and _appr["content_json"]:
                            st.caption("Content truncated — click below for full text")
                        # Full content JSON toggle
                        if _appr["content_json"]:
                            with st.popover("View Full Content"):
                                try:
                                    import json as _json_mod
                                    _parsed = _json_mod.loads(_appr["content_json"])
                                    st.json(_parsed)
                                except Exception:
                                    st.code(_appr["content_json"])

                        # Target sheet (where content was/will be written)
                        if _appr.get("target_sheet"):
                            _target_display = _appr["target_sheet"].replace(":", " > ")
                            _target_label = "Written to" if _appr["status"] == "Approved" else "Target"
                            st.caption(f"{_target_label}: {_target_display}")

                        # Admin notes (skip for Rejected — already shown above)
                        if _appr["notes"] and _appr["status"] != "Rejected":
                            st.markdown(
                                f"<div style='font-size:0.78rem;color:{_C['muted']};margin-top:4px'>"
                                f"<b>Notes:</b> {_appr['notes']}</div>",
                                unsafe_allow_html=True,
                            )

                        # Admin actions for pending items
                        if _is_pending and _user_is_admin:
                            _ap_c1, _ap_c2 = st.columns(2)
                            with _ap_c1:
                                if st.button("Approve", key=f"apr_ok_{_appr['id']}", width="stretch"):
                                    with st.spinner("Approving..."):
                                        try:
                                            _apr.approve_item(_appr["row_index"], approved_by=_user_name)
                                            try:
                                                notification_tools.notify_approval_decided(
                                                    _appr["id"], _appr.get("scene_id", ""),
                                                    _appr.get("content_type", ""), "Approved",
                                                    _user_name, _appr.get("submitted_by", ""))
                                                _cached_unread_count.clear()
                                            except Exception:
                                                pass
                                            st.success(f"{_appr['id']} approved and written to target.")
                                            _cached_load_approvals.clear()
                                            st.rerun()
                                        except Exception as _e:
                                            st.error(f"Failed: {_e}")
                            with _ap_c2:
                                _rej_notes = st.text_input("Rejection notes", key=f"apr_rej_notes_{_appr['id']}")
                                if st.button("Reject", key=f"apr_rej_{_appr['id']}", width="stretch"):
                                    if not _rej_notes.strip():
                                        st.error("Notes are required for rejection.")
                                    else:
                                        with st.spinner("Rejecting..."):
                                            try:
                                                _apr.reject_item(_appr["row_index"], rejected_by=_user_name,
                                                                 notes=_rej_notes.strip())
                                                try:
                                                    notification_tools.notify_approval_decided(
                                                        _appr["id"], _appr.get("scene_id", ""),
                                                        _appr.get("content_type", ""), "Rejected",
                                                        _user_name, _appr.get("submitted_by", ""))
                                                    _cached_unread_count.clear()
                                                except Exception:
                                                    pass
                                                st.warning(f"{_appr['id']} rejected.")
                                                _cached_load_approvals.clear()
                                                st.rerun()
                                            except Exception as _e:
                                                st.error(f"Failed: {_e}")

                        # Non-admin: read-only status
                        elif _is_pending and not _user_is_admin:
                            st.markdown(
                                f"<div style='font-size:0.78rem;color:{_C['amber']};margin-top:4px'>"
                                f"Awaiting admin review</div>",
                                unsafe_allow_html=True,
                            )

                        # Link to Asset Tracker scene
                        if _appr.get("scene_id"):
                            if st.button(
                                f"View {_appr['scene_id']} in Asset Tracker",
                                key=f"apr_goto_at_{_appr['id']}",
                            ):
                                st.session_state["_tk_nav_to"] = "Asset Tracker"
                                st.session_state["at_selected"] = _appr["scene_id"]
                                st.rerun()

        # ── SUB-VIEW 3: Tickets ──────────────────────────────────────────────────
        elif _tk_mode == "Tickets":
            hub_ui.section("Ticket Dashboard")

            # Filters
            _tf1, _tf2 = st.columns(2)
            with _tf1:
                _filt_status = st.selectbox("Status", ["All"] + _tkt.STATUSES, key="tk_filt_status")
            with _tf2:
                _filt_project = st.selectbox("Project", ["All"] + _tkt.PROJECTS, key="tk_filt_project")
            _tf3, _tf4 = st.columns(2)
            with _tf3:
                _filt_priority = st.selectbox("Priority", ["All"] + _tkt.PRIORITIES, key="tk_filt_priority")
            with _tf4:
                _filt_assignee = st.selectbox(
                    "Assigned To",
                    ["All", "Mine", "Unassigned"] + _tkt.EMPLOYEES,
                    key="tk_filt_assignee",
                )

            # Load tickets (globally cached)
            _all_tickets = _cached_load_tickets()

            # Refresh button
            if st.button("Refresh", key="tk_refresh"):
                _cached_load_tickets.clear()
                st.rerun()

            # Apply filters
            _filtered = _all_tickets
            if _filt_status != "All":
                _filtered = [t for t in _filtered if t["status"] == _filt_status]
            if _filt_project != "All":
                _filtered = [t for t in _filtered if t["project"] == _filt_project]
            if _filt_priority != "All":
                _filtered = [t for t in _filtered if t["priority"] == _filt_priority]
            if _filt_assignee == "Mine":
                _filtered = [t for t in _filtered
                             if t.get("assigned_to") == _user_name or t.get("submitted_by") == _user_name]
            elif _filt_assignee == "Unassigned":
                _filtered = [t for t in _filtered if not t.get("assigned_to")]
            elif _filt_assignee not in ("All",):
                _filtered = [t for t in _filtered if t.get("assigned_to") == _filt_assignee]

            # Summary stats — consolidated 3 metrics
            _active_ct = sum(1 for t in _all_tickets if t["status"] in ("New", "Approved", "In Progress"))
            _review_ct = sum(1 for t in _all_tickets if t["status"] == "In Review")
            _resolved_ct = sum(1 for t in _all_tickets if t["status"] in ("Closed", "Rejected"))
            _sm1, _sm2, _sm3 = st.columns(3)
            _sm1.metric("Active", _active_ct)
            _sm2.metric("In Review", _review_ct)
            _sm3.metric("Resolved", _resolved_ct)

            st.divider()

            # Hide Closed/Rejected from grid unless explicitly filtered
            _hidden_count = 0
            if _filt_status == "All":
                _hidden_count = sum(1 for t in _filtered if t["status"] in ("Closed", "Rejected"))
                _filtered = [t for t in _filtered if t["status"] not in ("Closed", "Rejected")]
            if _hidden_count:
                st.caption(f"{_hidden_count} closed/rejected ticket{'s' if _hidden_count != 1 else ''} hidden — filter by status to view")

            _pri_colors = {"Critical": _C["red"], "High": _C["amber"], "Medium": _C["blue"], "Low": _C["green"]}
            _st_colors = {
                "New": _C["blue"], "Approved": _C["green"], "In Progress": _C["amber"],
                "In Review": _C["accent"], "Closed": _C["green"], "Rejected": _C["red"],
            }
            _st_bg = {
                "New": _C["blue_dim"], "Approved": _C["green_dim"], "In Progress": _C["amber_dim"],
                "In Review": _C["accent_dim"], "Closed": _C["green_dim"], "Rejected": _C["red_dim"],
            }

            # ── Detail view (selected ticket) ────────────────────────────────
            _selected_tid = st.session_state.get("tk_selected")
            _selected_ticket = next((t for t in _all_tickets if t["id"] == _selected_tid), None) if _selected_tid else None

            if _selected_ticket:
                _t = _selected_ticket
                _pc = _pri_colors.get(_t["priority"], _C["muted"])
                _sc = _st_colors.get(_t["status"], _C["muted"])
                _sb = _st_bg.get(_t["status"], "rgba(255,255,255,0.06)")

                if st.button("Back to all tickets", key="tk_back"):
                    st.session_state.pop("tk_selected", None)
                    st.rerun()

                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:12px;margin:8px 0 4px'>"
                    f"<div style='width:4px;height:28px;border-radius:2px;background:{_pc}'></div>"
                    f"<span style='font-family:DM Mono,monospace;font-size:0.82rem;color:{_C['subtle']}'>{_t['id']}</span>"
                    f"<span style='font-family:Syne,sans-serif;font-size:1.2rem;font-weight:700;color:{_C['text']}'>{_t['title']}</span>"
                    f"<span style='background:{_sb};color:{_sc};font-size:0.72rem;font-weight:600;"
                    f"padding:3px 10px;border-radius:9999px;margin-left:auto'>{_t['status']}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                _b_pri = hub_ui.badge(_t['priority'], {"Critical": "red", "High": "amber", "Medium": "blue", "Low": "green"}.get(_t["priority"], "gray"))
                _b_proj = hub_ui.badge(_t['project'], 'blue')
                _b_type = hub_ui.badge(_t['type'], 'gray')
                _b_assign = hub_ui.badge(_t.get('assigned_to') or 'Unassigned', 'accent' if _t.get('assigned_to') else 'gray')
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:10px;margin:0 0 12px 16px'>"
                    f"{_b_pri} {_b_proj} {_b_type} {_b_assign}"
                    f"<span style='font-size:0.75rem;color:{_C['muted']}'>by {_t['submitted_by']} &middot; {_t['date']}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                hub_ui.section("Description")
                st.markdown(
                    f"<div class='hub-card' style='white-space:pre-wrap;font-size:0.85rem;"
                    f"line-height:1.7;color:{_C['text']}'>{_t['description']}</div>",
                    unsafe_allow_html=True,
                )

                if _t["admin_notes"]:
                    hub_ui.section("Admin Notes")
                    # Render as timeline if timestamped entries exist
                    _notes_raw = _t["admin_notes"]
                    import re as _re_mod
                    _note_entries = _re_mod.split(r'(?=\[\d{4}-\d{2}-\d{2})', _notes_raw)
                    _note_entries = [n.strip() for n in _note_entries if n.strip()]
                    if len(_note_entries) > 1 or _note_entries[0].startswith("["):
                        _notes_html = ""
                        for _ne in _note_entries:
                            _ts_match = _re_mod.match(r'\[(\d{4}-\d{2}-\d{2}\s*\d{0,2}:?\d{0,2})\s*([^\]]*)\]\s*(.*)', _ne, _re_mod.DOTALL)
                            if _ts_match:
                                _notes_html += (
                                    f"<div style='padding:6px 0;border-bottom:1px solid {_C['border']}'>"
                                    f"<span style='font-size:0.68rem;color:{_C['subtle']}'>{_ts_match.group(1)}</span>"
                                    f" <span style='font-size:0.72rem;font-weight:600;color:{_C['accent']}'>{_ts_match.group(2)}</span>"
                                    f"<div style='font-size:0.82rem;color:{_C['text']};margin-top:2px'>{_ts_match.group(3)}</div>"
                                    f"</div>"
                                )
                            else:
                                _notes_html += f"<div style='font-size:0.82rem;color:{_C['text']};padding:4px 0'>{_ne}</div>"
                        st.markdown(
                            f"<div class='hub-card hub-card-accent' style='border-left-color:{_C['accent']}'>"
                            f"{_notes_html}</div>",
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(
                            f"<div class='hub-card hub-card-accent' style='border-left-color:{_C['accent']}'>"
                            f"<span style='font-size:0.85rem'>{_notes_raw}</span></div>",
                            unsafe_allow_html=True,
                        )

                if _t["date_resolved"]:
                    st.caption(f"Resolved: {_t['date_resolved']}")

                # QC Feedback — anyone can report test results on active tickets
                if _t["status"] in ("New", "Approved", "In Progress"):
                    hub_ui.section("QC Feedback")
                    _qc_note = st.text_input("Notes (optional)", key="tk_qc_note",
                                             placeholder="What did you find?")
                    _qc1, _qc2 = st.columns(2)
                    with _qc1:
                        if st.button("Fixed", key="tk_qc_fixed", width="stretch"):
                            with st.spinner("Updating..."):
                                try:
                                    _qc_text = f"QC passed" + (f": {_qc_note.strip()}" if _qc_note.strip() else "")
                                    _tkt.update_ticket(
                                        _t["row_index"], status="In Review",
                                        admin_notes=_tkt.append_note(_t["admin_notes"], _user_name, _qc_text),
                                    )
                                    try:
                                        notification_tools.notify_ticket_status(
                                            _t["id"], _t["title"], "In Review", _user_name,
                                            submitted_by=_t.get("submitted_by", ""),
                                            assigned_to=_t.get("assigned_to", ""))
                                        _cached_unread_count.clear()
                                    except Exception:
                                        pass
                                    st.success(f"Marked as fixed — moved to In Review.")
                                    _cached_load_tickets.clear()
                                    st.rerun()
                                except Exception as _e:
                                    st.error(f"Failed: {_e}")
                    with _qc2:
                        if st.button("Still Broken", key="tk_qc_broken", width="stretch"):
                            with st.spinner("Updating..."):
                                try:
                                    _qc_text = f"QC failed" + (f": {_qc_note.strip()}" if _qc_note.strip() else "")
                                    _tkt.update_ticket(
                                        _t["row_index"],
                                        admin_notes=_tkt.append_note(_t["admin_notes"], _user_name, _qc_text),
                                    )
                                    st.warning(f"Feedback recorded.")
                                    _cached_load_tickets.clear()
                                    st.rerun()
                                except Exception as _e:
                                    st.error(f"Failed: {_e}")

                # Verify & Close for In Review tickets
                if _t["status"] == "In Review":
                    hub_ui.section("Verify This Change")
                    _last_note = _t["admin_notes"].strip().split("\n")[-1] if _t["admin_notes"] else ""
                    _verify_ctx = f"Change reported: **{_last_note}**" if _last_note else "No notes on what was changed."
                    st.markdown(
                        f"<div style='font-size:0.82rem;color:{_C['muted']};margin-bottom:8px'>"
                        f"This ticket was marked as done. {_verify_ctx}</div>",
                        unsafe_allow_html=True,
                    )
                    _vc1, _vc2 = st.columns(2)
                    with _vc1:
                        if st.button("Verified — Close", key="tk_verify_close", width="stretch"):
                            with st.spinner("Closing..."):
                                try:
                                    _tkt.update_ticket(
                                        _t["row_index"], status="Closed",
                                        approved_by=_user_name,
                                        admin_notes=(_t["admin_notes"] + f"\nVerified by {_user_name}").strip(),
                                    )
                                    try:
                                        notification_tools.notify_ticket_status(
                                            _t["id"], _t["title"], "Closed", _user_name,
                                            submitted_by=_t.get("submitted_by", ""),
                                            assigned_to=_t.get("assigned_to", ""))
                                        _cached_unread_count.clear()
                                    except Exception:
                                        pass
                                    st.success(f"Ticket {_t['id']} verified and closed.")
                                    _cached_load_tickets.clear()
                                    st.session_state.pop("tk_selected", None)
                                    st.rerun()
                                except Exception as _e:
                                    st.error(f"Failed: {_e}")
                    with _vc2:
                        if st.button("Not Fixed — Reopen", key="tk_verify_reopen", width="stretch"):
                            with st.spinner("Reopening..."):
                                try:
                                    _tkt.update_ticket(
                                        _t["row_index"], status="In Progress",
                                        admin_notes=(_t["admin_notes"] + f"\nReopened by {_user_name} — not fixed").strip(),
                                    )
                                    st.warning(f"Ticket {_t['id']} reopened.")
                                    _cached_load_tickets.clear()
                                    st.rerun()
                                except Exception as _e:
                                    st.error(f"Failed: {_e}")

                # Quick approve/reject for New tickets
                if _t["status"] == "New" and _user_is_admin:
                    hub_ui.section("Review New Ticket")
                    _qa1, _qa2 = st.columns(2)
                    with _qa1:
                        if st.button("Approve", key="tk_quick_approve", width="stretch"):
                            with st.spinner("Approving..."):
                                try:
                                    _tkt.update_ticket(
                                        _t["row_index"], status="Approved",
                                        approved_by=_user_name,
                                        admin_notes=_tkt.append_note(_t["admin_notes"], _user_name, "Approved"),
                                    )
                                    try:
                                        notification_tools.notify_ticket_status(
                                            _t["id"], _t["title"], "Approved", _user_name,
                                            submitted_by=_t.get("submitted_by", ""))
                                        _cached_unread_count.clear()
                                    except Exception:
                                        pass
                                    st.success(f"Ticket {_t['id']} approved.")
                                    _cached_load_tickets.clear()
                                    st.rerun()
                                except Exception as _e:
                                    st.error(f"Failed: {_e}")
                    with _qa2:
                        _reject_reason = st.text_input("Rejection reason", key="tk_quick_rej_reason")
                        if st.button("Reject", key="tk_quick_reject", width="stretch"):
                            if not _reject_reason.strip():
                                st.error("Please provide a reason.")
                            else:
                                with st.spinner("Rejecting..."):
                                    try:
                                        _tkt.update_ticket(
                                            _t["row_index"], status="Rejected",
                                            approved_by=_user_name,
                                            admin_notes=_tkt.append_note(_t["admin_notes"], _user_name, f"Rejected: {_reject_reason.strip()}"),
                                        )
                                        try:
                                            notification_tools.notify_ticket_status(
                                                _t["id"], _t["title"], "Rejected", _user_name,
                                                submitted_by=_t.get("submitted_by", ""))
                                            _cached_unread_count.clear()
                                        except Exception:
                                            pass
                                        st.warning(f"Ticket {_t['id']} rejected.")
                                        _cached_load_tickets.clear()
                                        st.rerun()
                                    except Exception as _e:
                                        st.error(f"Failed: {_e}")

                # Admin actions
                if _user_is_admin:
                    hub_ui.section("Admin Actions")
                    # Valid status transitions
                    _TRANSITIONS = {
                        "New":         ["New", "Approved", "Rejected"],
                        "Approved":    ["Approved", "In Progress", "Rejected"],
                        "In Progress": ["In Progress", "In Review", "Closed"],
                        "In Review":   ["In Review", "In Progress", "Closed"],
                        "Closed":      ["Closed", "In Progress"],
                        "Rejected":    ["Rejected", "New"],
                    }
                    _valid_statuses = _TRANSITIONS.get(_t["status"], _tkt.STATUSES)
                    _ac1, _ac2 = st.columns(2)
                    with _ac1:
                        _new_status = st.selectbox(
                            "Update Status", _valid_statuses,
                            index=0,
                            key="tk_detail_status",
                        )
                    with _ac2:
                        _assign_opts = ["Unassigned"] + _tkt.EMPLOYEES
                        _cur_assign = _t.get("assigned_to", "") or "Unassigned"
                        _assign_idx = _assign_opts.index(_cur_assign) if _cur_assign in _assign_opts else 0
                        _new_assignee = st.selectbox(
                            "Assign To", _assign_opts, index=_assign_idx,
                            key="tk_detail_assignee",
                        )
                    _add_note = st.text_input("Add Note", key="tk_detail_add_note", placeholder="Add a timestamped note...")
                    if st.button("Update Ticket", key="tk_detail_save", width="stretch"):
                        _final_notes = None
                        if _add_note.strip():
                            _final_notes = _tkt.append_note(_t["admin_notes"], _user_name, _add_note.strip())
                        _final_assignee = _new_assignee if _new_assignee != "Unassigned" else ""
                        _status_changed = _new_status != _t["status"]
                        _assignee_changed = _final_assignee != (_t.get("assigned_to") or "")
                        _has_note = _final_notes is not None

                        if not _status_changed and not _assignee_changed and not _has_note:
                            st.warning("Nothing changed.")
                        else:
                            with st.spinner("Updating..."):
                                try:
                                    _tkt.update_ticket(
                                        _t["row_index"],
                                        status=_new_status if _status_changed else None,
                                        approved_by=_user_name if _status_changed and _new_status in ("Approved", "Rejected", "Closed") else None,
                                        admin_notes=_final_notes,
                                        assigned_to=_final_assignee if _assignee_changed else None,
                                    )
                                    try:
                                        if _status_changed:
                                            notification_tools.notify_ticket_status(
                                                _t["id"], _t["title"], _new_status, _user_name,
                                                submitted_by=_t.get("submitted_by", ""),
                                                assigned_to=_t.get("assigned_to", ""))
                                        if _assignee_changed and _final_assignee:
                                            notification_tools.notify_ticket_assigned(
                                                _t["id"], _t["title"], _final_assignee, _user_name)
                                        _cached_unread_count.clear()
                                    except Exception:
                                        pass
                                    st.success(f"Ticket {_t['id']} updated.")
                                    _cached_load_tickets.clear()
                                    st.rerun()
                                except Exception as _e:
                                    st.error(f"Failed to update: {_e}")

            # ── Grid view (3 columns) ────────────────────────────────────────
            else:
                if not _filtered:
                    st.info("No tickets found matching your filters.")
                else:
                    _sorted_tickets = list(reversed(_filtered))
                    for _ri in range(0, len(_sorted_tickets), 3):
                        _row_tickets = _sorted_tickets[_ri:_ri + 3]
                        _cols = st.columns(3)
                        for _ci, _ticket in enumerate(_row_tickets):
                            with _cols[_ci]:
                                _pc = _pri_colors.get(_ticket["priority"], _C["muted"])
                                _sc = _st_colors.get(_ticket["status"], _C["muted"])
                                _sb = _st_bg.get(_ticket["status"], "rgba(255,255,255,0.06)")
                                _trunc_title = _ticket["title"][:60] + ("..." if len(_ticket["title"]) > 60 else "")
                                _trunc_desc = _ticket["description"][:120] + ("..." if len(_ticket["description"]) > 120 else "")
                                _ticket["_assignee_html"] = (
                                    f"<span style='color:{_C['border']}'>&middot;</span>"
                                    f"<span style='font-size:0.68rem;color:{_C['accent']}'>{_ticket.get('assigned_to', '')}</span>"
                                ) if _ticket.get("assigned_to") else ""
                                st.markdown(
                                    f"<div style='background:{_C['surface']};border:1px solid {_C['border']};"
                                    f"border-top:3px solid {_pc};border-radius:8px;padding:14px 16px;"
                                    f"min-height:160px;display:flex;flex-direction:column'>"
                                    f"<div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:8px'>"
                                    f"<span style='font-family:DM Mono,monospace;font-size:0.7rem;color:{_C['subtle']}'>{_ticket['id']}</span>"
                                    f"<span style='background:{_sb};color:{_sc};font-size:0.65rem;font-weight:600;"
                                    f"padding:2px 8px;border-radius:9999px'>{_ticket['status']}</span>"
                                    f"</div>"
                                    f"<div style='font-size:0.88rem;font-weight:600;color:{_C['text']};"
                                    f"margin-bottom:6px;line-height:1.3'>{_trunc_title}</div>"
                                    f"<div style='font-size:0.75rem;color:{_C['muted']};line-height:1.4;"
                                    f"margin-bottom:auto;overflow:hidden;display:-webkit-box;"
                                    f"-webkit-line-clamp:3;-webkit-box-orient:vertical'>{_trunc_desc}</div>"
                                    f"<div style='display:flex;gap:6px;align-items:center;margin-top:10px;"
                                    f"padding-top:8px;border-top:1px solid {_C['border']}'>"
                                    f"<span style='font-size:0.68rem;color:{_C['muted']}'>{_ticket['project']}</span>"
                                    f"<span style='color:{_C['border']}'>&middot;</span>"
                                    f"<span style='font-size:0.68rem;color:{_C['subtle']}'>{_ticket['submitted_by']}</span>"
                                    f"{_ticket.get('_assignee_html', '')}"
                                    f"<span style='font-size:0.68rem;color:{_C['subtle']};margin-left:auto'>{_ticket['date'][:10]}</span>"
                                    f"</div>"
                                    f"</div>",
                                    unsafe_allow_html=True,
                                )
                                if st.button("Open", key=f"tk_open_{_ticket['id']}", width="stretch"):
                                    st.session_state["tk_selected"] = _ticket["id"]
                                    st.rerun()

        # ── SUB-VIEW 4: Submit ───────────────────────────────────────────────────
        elif _tk_mode == "Submit":
            hub_ui.section("Submit a New Ticket")

            # Persistent success banner
            _last_submitted = st.session_state.get("_last_submitted_ticket")
            if _last_submitted:
                st.info(f"Last submitted: **{_last_submitted}** — view it in the Tickets tab")
                if st.button("Dismiss", key="dismiss_last_ticket"):
                    st.session_state.pop("_last_submitted_ticket", None)
                    st.rerun()

            # Check for pre-fill from Asset Tracker
            _prefill = st.session_state.pop("tk_prefill", {})
            if _prefill:
                if st.button("Cancel — return to Asset Tracker", key="tk_cancel_prefill"):
                    st.session_state["_tk_nav_to"] = "Asset Tracker"
                    st.rerun()

            with st.form("ticket_submit_form", clear_on_submit=True):
                _tk_c1, _tk_c2 = st.columns(2)
                with _tk_c1:
                    st.text_input("Submitted By", value=_user_name, disabled=True, key="tk_who_display")
                with _tk_c2:
                    _pf_proj = _prefill.get("project", "")
                    _pf_proj_idx = _tkt.PROJECTS.index(_pf_proj) if _pf_proj in _tkt.PROJECTS else 0
                    _tk_project = st.selectbox("Project", _tkt.PROJECTS, index=_pf_proj_idx, key="tk_project")
                _tk_c3, _tk_c4 = st.columns(2)
                with _tk_c3:
                    _pf_type = _prefill.get("type", "")
                    _pf_type_idx = _tkt.TICKET_TYPES.index(_pf_type) if _pf_type in _tkt.TICKET_TYPES else 0
                    _tk_type = st.selectbox("Type", _tkt.TICKET_TYPES, index=_pf_type_idx, key="tk_type")
                with _tk_c4:
                    _tk_priority = st.selectbox("Priority", _tkt.PRIORITIES,
                                                index=1, key="tk_priority")
                _tk_title = st.text_input("Title",
                                          value=_prefill.get("title", ""),
                                          placeholder="Short summary of the issue or request",
                                          key="tk_title")
                _tk_desc = st.text_area("Description",
                                        value=_prefill.get("desc", ""),
                                        placeholder="Detailed description — steps to reproduce, expected behavior, screenshots, etc.",
                                        height=200, key="tk_desc")
                _tk_submit = st.form_submit_button("Submit Ticket", width="stretch")

            if _tk_submit:
                if not _tk_title.strip():
                    st.error("Title is required.")
                elif not _tk_desc.strip():
                    st.error("Description is required.")
                else:
                    with st.spinner("Submitting ticket..."):
                        try:
                            _new_id = _tkt.create_ticket(
                                _user_name, _tk_project, _tk_type, _tk_priority,
                                _tk_title.strip(), _tk_desc.strip(),
                            )
                            try:
                                notification_tools.notify_ticket_created(_new_id, _tk_title.strip(), _user_name)
                                _cached_unread_count.clear()
                            except Exception:
                                pass
                            st.session_state["_last_submitted_ticket"] = _new_id
                            _cached_load_tickets.clear()
                            st.rerun()
                        except Exception as _e:
                            st.error(f"Failed to submit ticket: {_e}")

        # ── SUB-VIEW 5: Users (Drew + David only) ────────────────────────────
        elif _tk_mode == "Users" and _user_can_manage_users:
            hub_ui.section("User Management")

            # Load current users from sheet
            _um_users = auth_config.load_users_config()
            _um_roles = ["admin", "editor"]
            _um_editing = st.session_state.get("um_editing")

            # Collapse duplicate emails into unique people (Drew has 2 emails)
            _um_people = {}
            for _um_email, _um_info in _um_users.items():
                _um_name = _um_info["name"]
                if _um_name not in _um_people:
                    _um_people[_um_name] = {"emails": [], "role": _um_info["role"], "tabs": _um_info["allowed_tabs"]}
                _um_people[_um_name]["emails"].append(_um_email)

            # Summary metrics
            _um_total = len(_um_people)
            _um_admin_ct = sum(1 for p in _um_people.values() if p["role"] == "admin")
            _um_editor_ct = _um_total - _um_admin_ct
            _um1, _um2, _um3 = st.columns(3)
            _um1.metric("Total Users", _um_total)
            _um2.metric("Admins", _um_admin_ct)
            _um3.metric("Editors", _um_editor_ct)

            st.divider()

            # ── User cards — 2 per row ───────────────────────────────────────
            _um_sorted = sorted(_um_people.items(), key=lambda x: (0 if x[1]["role"] == "admin" else 1, x[0]))
            for _um_ri in range(0, len(_um_sorted), 2):
                _um_row = _um_sorted[_um_ri:_um_ri + 2]
                _um_cols = st.columns(2)
                for _um_ci, (_um_name, _um_p) in enumerate(_um_row):
                    with _um_cols[_um_ci]:
                        _um_role = _um_p["role"]
                        _um_tabs = _um_p["tabs"]
                        _um_emails = _um_p["emails"]
                        _um_key = _um_name.lower().replace(" ", "_")
                        _um_is_all = set(_um_tabs) >= set(auth_config.ALL_TAB_KEYS)
                        _um_tabs_display = "All Tabs" if _um_is_all else ", ".join(t for t in _um_tabs if t in auth_config.ALL_TAB_KEYS)

                        # Role badge
                        _role_badge_var = "accent" if _um_role == "admin" else "blue"
                        _role_badge = hub_ui.badge(_um_role.capitalize(), _role_badge_var)

                        # Grail writer indicator
                        _writer_badge = ""
                        if _um_name in auth_config._GRAIL_WRITERS:
                            _writer_badge = f" {hub_ui.badge('Grail Writer', 'green')}"

                        # Email(s)
                        _emails_html = "<br>".join(
                            f"<span style='font-size:0.72rem;color:{_C['muted']}'>{e}</span>"
                            for e in _um_emails
                        )

                        # Tabs as small pills (consistent style for all)
                        if _um_is_all:
                            _tabs_html = (
                                f"<span style='display:inline-block;font-size:0.65rem;padding:2px 8px;"
                                f"border-radius:3px;background:{_C['green_dim']};color:{_C['green']};"
                                f"font-weight:500'>All tabs</span>"
                            )
                        else:
                            _valid = [t for t in _um_tabs if t in auth_config.ALL_TAB_KEYS]
                            _tabs_html = " ".join(
                                f"<span style='display:inline-block;font-size:0.65rem;padding:2px 6px;"
                                f"border-radius:3px;background:{_C['elevated']};color:{_C['muted']};"
                                f"margin:1px'>{t}</span>"
                                for t in _valid
                            ) if _valid else f"<span style='font-size:0.7rem;color:{_C['red']}'>No tabs</span>"

                        hub_ui.card(
                            f"<div style='display:flex;align-items:center;gap:8px;margin-bottom:6px'>"
                            f"<span style='font-family:Syne,sans-serif;font-size:0.95rem;font-weight:700;"
                            f"color:{_C['text']}'>{_um_name}</span>"
                            f"{_role_badge}{_writer_badge}"
                            f"</div>"
                            f"{_emails_html}"
                            f"<div style='margin-top:8px'>{_tabs_html}</div>"
                        )

                        # Edit / Remove buttons
                        _eb1, _eb2 = st.columns(2)
                        with _eb1:
                            if st.button("Edit", key=f"um_edit_{_um_key}", width="stretch"):
                                st.session_state["um_editing"] = _um_name
                                st.rerun()
                        with _eb2:
                            with st.container():
                                st.markdown("<div class='danger-btn'>", unsafe_allow_html=True)
                                if st.button("Remove", key=f"um_del_{_um_key}", width="stretch"):
                                    st.session_state[f"um_confirm_del_{_um_key}"] = True
                                    st.rerun()
                                st.markdown("</div>", unsafe_allow_html=True)

            # ── Edit panel (slides in below cards when editing) ──────────────
            if _um_editing and _um_editing in _um_people:
                st.divider()
                _ep = _um_people[_um_editing]
                _ep_key = _um_editing.lower().replace(" ", "_")
                hub_ui.section(f"Editing {_um_editing}")

                _ep_r1, _ep_r2 = st.columns(2)
                with _ep_r1:
                    _ep_new_role = st.selectbox(
                        "Role", _um_roles,
                        index=_um_roles.index(_ep["role"]) if _ep["role"] in _um_roles else 1,
                        key=f"um_erole_{_ep_key}")
                with _ep_r2:
                    _ep_valid_tabs = [t for t in _ep["tabs"] if t in auth_config.ALL_TAB_KEYS]
                    _ep_is_all = set(_ep["tabs"]) >= set(auth_config.ALL_TAB_KEYS)
                    _ep_new_tabs = st.multiselect(
                        "Allowed Tabs", auth_config.ALL_TAB_KEYS,
                        default=auth_config.ALL_TAB_KEYS if _ep_is_all else _ep_valid_tabs,
                        key=f"um_etabs_{_ep_key}")

                _ep_b1, _ep_b2 = st.columns(2)
                with _ep_b1:
                    if st.button("Save Changes", key=f"um_esave_{_ep_key}", width="stretch"):
                        try:
                            _gc_um = auth_config._get_client()
                            _sh_um = _gc_um.open_by_key(auth_config.USERS_SHEET_ID)
                            _ws_um = _sh_um.worksheet(auth_config.USERS_TAB_NAME)
                            _um_rows = _ws_um.get_all_values()
                            _tabs_val = "ALL" if set(_ep_new_tabs) == set(auth_config.ALL_TAB_KEYS) else ", ".join(_ep_new_tabs)
                            _updated = False
                            for _ep_email in _ep["emails"]:
                                for _ri_um, _row_um in enumerate(_um_rows[1:], start=2):
                                    if _row_um[0].strip().lower() == _ep_email:
                                        _ws_um.update_cell(_ri_um, 3, _ep_new_role)
                                        _ws_um.update_cell(_ri_um, 4, _tabs_val)
                                        _updated = True
                            if _updated:
                                auth_config.invalidate_cache()
                                st.session_state.pop("um_editing", None)
                                st.rerun()
                        except Exception as _e_um:
                            st.error(f"Failed to update: {_e_um}")
                with _ep_b2:
                    if st.button("Cancel", key=f"um_ecancel_{_ep_key}", width="stretch"):
                        st.session_state.pop("um_editing", None)
                        st.rerun()

            # ── Delete confirmation (appears inline) ─────────────────────────
            for _um_name_d, _um_p_d in _um_people.items():
                _um_key_d = _um_name_d.lower().replace(" ", "_")
                if st.session_state.get(f"um_confirm_del_{_um_key_d}"):
                    st.divider()
                    st.warning(f"Remove **{_um_name_d}** ({', '.join(_um_p_d['emails'])})? This cannot be undone.")
                    _d1, _d2 = st.columns(2)
                    with _d1:
                        if st.button("Yes, Remove", key=f"um_cyes_{_um_key_d}", width="stretch"):
                            try:
                                _gc_um2 = auth_config._get_client()
                                _sh_um2 = _gc_um2.open_by_key(auth_config.USERS_SHEET_ID)
                                _ws_um2 = _sh_um2.worksheet(auth_config.USERS_TAB_NAME)
                                _um_rows2 = _ws_um2.get_all_values()
                                for _ep_email_d in _um_p_d["emails"]:
                                    for _ri_um2, _row_um2 in enumerate(_um_rows2[1:], start=2):
                                        if _row_um2[0].strip().lower() == _ep_email_d:
                                            _ws_um2.delete_rows(_ri_um2)
                                            break
                                auth_config.invalidate_cache()
                                st.session_state.pop(f"um_confirm_del_{_um_key_d}", None)
                                st.rerun()
                            except Exception as _e_um2:
                                st.error(f"Failed to remove: {_e_um2}")
                    with _d2:
                        if st.button("Cancel", key=f"um_cno_{_um_key_d}", width="stretch"):
                            st.session_state.pop(f"um_confirm_del_{_um_key_d}", None)
                            st.rerun()

            # ── Add user ─────────────────────────────────────────────────────
            st.divider()
            hub_ui.section("Add User")
            with st.form("um_add_form", clear_on_submit=True):
                _um_ac1, _um_ac2 = st.columns(2)
                with _um_ac1:
                    _um_new_email = st.text_input("Email", placeholder="user@example.com", key="um_new_email")
                with _um_ac2:
                    _um_new_name = st.text_input("Name", placeholder="First name", key="um_new_name")
                _um_ac3, _um_ac4 = st.columns(2)
                with _um_ac3:
                    _um_new_r = st.selectbox("Role", _um_roles, index=1, key="um_new_role")
                with _um_ac4:
                    _um_new_t = st.multiselect("Tabs", auth_config.ALL_TAB_KEYS,
                                                default=auth_config.ALL_TAB_KEYS, key="um_new_tabs")
                _um_add_btn = st.form_submit_button("Add User", width="stretch")

            if _um_add_btn:
                if not _um_new_email.strip() or not _um_new_name.strip():
                    st.error("Email and name are required.")
                elif _um_new_email.strip().lower() in _um_users:
                    st.error("User already exists.")
                else:
                    try:
                        _gc_uma = auth_config._get_client()
                        _sh_uma = _gc_uma.open_by_key(auth_config.USERS_SHEET_ID)
                        _ws_uma = _sh_uma.worksheet(auth_config.USERS_TAB_NAME)
                        _tabs_val_new = "ALL" if set(_um_new_t) == set(auth_config.ALL_TAB_KEYS) else ", ".join(_um_new_t)
                        _ws_uma.append_row(
                            [_um_new_email.strip().lower(), _um_new_name.strip(), _um_new_r, _tabs_val_new],
                            value_input_option="USER_ENTERED")
                        auth_config.invalidate_cache()
                        st.success(f"Added {_um_new_name}")
                        st.rerun()
                    except Exception as _e_uma:
                        st.error(f"Failed to add user: {_e_uma}")
